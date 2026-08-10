from __future__ import annotations

import json
from pathlib import Path

import pdfplumber
from docx import Document


ROOT = Path(r"D:\Projects\amora_ai")
SRS = Path(r"C:\Users\Yash Andrapiya\Downloads\Amora_SRS.pdf")
FEATURES = Path(r"C:\Users\Yash Andrapiya\Downloads\Amora_App_Feature_Note.docx")
OUT = ROOT / "tmp" / "project_audit" / "requirements"


def extract_pdf() -> dict:
    pages = []
    with pdfplumber.open(SRS) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            pages.append({"page": index, "text": page.extract_text(x_tolerance=2, y_tolerance=3) or ""})
    return {"path": str(SRS), "page_count": len(pages), "pages": pages}


def extract_docx() -> dict:
    doc = Document(FEATURES)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            paragraphs.append({"index": index, "style": paragraph.style.name, "text": text})

    tables = []
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    return {
        "path": str(FEATURES),
        "paragraphs": paragraphs,
        "tables": tables,
        "section_count": len(doc.sections),
    }


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    srs = extract_pdf()
    features = extract_docx()
    (OUT / "srs.json").write_text(json.dumps(srs, indent=2, ensure_ascii=False), encoding="utf-8")
    (OUT / "feature_note.json").write_text(json.dumps(features, indent=2, ensure_ascii=False), encoding="utf-8")
    (OUT / "srs.txt").write_text(
        "\n\n".join(f"=== PAGE {p['page']} ===\n{p['text']}" for p in srs["pages"]),
        encoding="utf-8",
    )
    feature_lines = []
    for p in features["paragraphs"]:
        feature_lines.append(f"[{p['style']}] {p['text']}")
    for table in features["tables"]:
        feature_lines.append(f"\n[TABLE {table['index']}]")
        feature_lines.extend(" | ".join(row) for row in table["rows"])
    (OUT / "feature_note.txt").write_text("\n".join(feature_lines), encoding="utf-8")
    print(json.dumps({"srs_pages": srs["page_count"], "feature_paragraphs": len(features["paragraphs"]), "feature_tables": len(features["tables"])}))


if __name__ == "__main__":
    main()
