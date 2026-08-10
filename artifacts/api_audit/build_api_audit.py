from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = Path(__file__).resolve().parent / "AMORAA_API_Audit_2026-08-10.docx"

BLUE = "2E74B5"
NAVY = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FA"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"
GRAY = "F2F2F2"
DARK = "243242"
WHITE = "FFFFFF"


existing = [
    (1, "GET", "/health", "No", "Implemented", "Backend-only", "Process health response; not a Flutter product API."),
    (2, "POST", "/api/auth/signup", "No", "Implemented", "Connected", "Creates local account and sends verification OTP."),
    (3, "POST", "/api/auth/verify-account", "No", "Implemented", "Connected", "Verifies phone OTP and issues access/refresh tokens."),
    (4, "POST", "/api/auth/resend-verification-code", "No", "Implemented", "Connected", "Generates another account-verification OTP; should verify account eligibility before issuing."),
    (5, "POST", "/api/auth/login", "No", "Implemented", "Connected", "Email/password sign-in for verified local accounts."),
    (6, "POST", "/api/auth/google", "No", "Incomplete", "Connected", "Returns 503 until GOOGLE_CLIENT_IDS is configured."),
    (7, "POST", "/api/auth/forgot-password", "No", "Implemented", "Connected", "Non-enumerating password-reset OTP request."),
    (8, "POST", "/api/auth/verify-reset-code", "No", "Implemented", "Connected", "Consumes reset OTP and returns a short-lived recovery token."),
    (9, "POST", "/api/auth/reset-password", "No", "Implemented", "Connected", "Changes password and revokes refresh tokens."),
    (10, "POST", "/api/auth/refresh-token", "No", "Implemented", "Connected", "Rotates refresh token and issues new tokens."),
    (11, "POST", "/api/auth/logout", "Bearer", "Implemented", "Connected", "Revokes the supplied refresh token."),
    (12, "GET", "/api/auth/me", "Bearer", "Implemented", "Connected", "Returns the authenticated account summary."),
    (13, "GET", "/api/onboarding/status", "Bearer", "Implemented", "Connected", "Loads or creates the user's onboarding profile."),
    (14, "PUT", "/api/onboarding/age", "Bearer", "Implemented", "Connected", "Validates ISO birth date and 18+ restriction."),
    (15, "PUT", "/api/onboarding/gender", "Bearer", "Implemented", "Connected", "Stores gender and optional custom value."),
    (16, "PUT", "/api/onboarding/interested-in", "Bearer", "Implemented", "Connected", "Stores one or more interested-in values."),
    (17, "PUT", "/api/onboarding/relationship-goal", "Bearer", "Implemented", "Connected", "Stores relationship-goal array."),
    (18, "PUT", "/api/onboarding/location", "Bearer", "Implemented", "Connected", "Stores city and preferred distance."),
    (19, "PUT", "/api/onboarding/starter-profile", "Bearer", "Implemented", "Connected", "Stores profession, education, and company."),
    (20, "PUT", "/api/onboarding/profile-completion", "Bearer", "Implemented", "Connected", "Stores bio, interests, prompts, communication style, and related fields."),
    (21, "POST", "/api/onboarding/photos", "Bearer", "Implemented", "Connected", "Multipart upload; six-photo and file constraints."),
    (22, "DELETE", "/api/onboarding/photos/:index", "Bearer", "Implemented", "Client only", "Dart client method exists, but no production UI call site was found."),
    (23, "PUT", "/api/onboarding/photos/primary", "Bearer", "Implemented", "Connected", "Sets primary photo during onboarding sync."),
    (24, "POST", "/api/onboarding/complete", "Bearer", "Implemented", "Connected", "Validates required onboarding fields and marks completion."),
    (25, "GET", "/api/discover/feed", "Bearer", "Incomplete", "Connected", "Loads candidates, but most filtering/pagination is in memory; distance/presence/premium data are placeholders."),
    (26, "POST", "/api/discover/swipe", "Bearer", "Implemented", "Connected", "Persists pass/like/super-like and creates reciprocal match."),
    (27, "POST", "/api/discover/rewind", "Bearer", "Implemented", "Connected", "Deletes the latest discover action."),
    (28, "POST", "/api/discover/boost", "Bearer", "Incomplete", "Connected", "Creates a free 30-minute boost without inventory, entitlement, or payment enforcement."),
    (29, "GET", "/api/discover/filters", "Bearer", "Implemented", "Connected", "Reads persisted DiscoverFilterPreference."),
    (30, "PUT", "/api/discover/filters", "Bearer", "Implemented", "Connected", "Validates and persists provided filter fields."),
]

feature_gaps = [
    ("Profile editing & bio", "Partial", "LocalProfileRepository / SharedPreferences", "Reuse onboarding starter-profile and profile-completion endpoints from edit flows; add a stable post-onboarding profile contract."),
    ("Photo manager & media prompts", "Partial", "Session/local photo operations", "Wire existing upload/delete/primary APIs; add reorder and binary voice/video upload operations."),
    ("Public profile detail", "Not connected", "DummyProfile or feed payload", "Add an authenticated profile-detail read endpoint with relationship visibility rules."),
    ("KYC / identity verification", "Not connected", "UI reports secure verification unavailable", "Add provider-backed verification submission and status APIs."),
    ("Compatibility answers & dealbreakers", "Not connected", "Auto-saved locally", "Persist answers and dealbreakers; feed them into matching/filter logic."),
    ("Matches & why-we-matched", "Not connected", "Dummy profile recommendations", "Expose match collection/detail/unmatch and explanation APIs."),
    ("Likes, super-likes & liked-you", "Not connected", "Session-scoped relationship controller", "Expose sent/received likes, revoke, entitlement, and paywall-aware results."),
    ("Saved & blocked profiles", "Not connected", "Session-scoped relationship controller", "Add list/create/delete APIs and enforce blocks in Discover and messaging."),
    ("Chat & messaging", "Not connected", "LocalChatRepository with seeded conversations", "Add conversation/message/media/read APIs plus authenticated realtime delivery."),
    ("Events catalog & detail", "Not connected", "events_dummy_data / ImageRepository", "Add event browse, detail, search, and eligibility APIs."),
    ("Event participation", "Not connected", "In-memory EventParticipationController", "Add registration, cancellation, waitlist, check-in, and My Events APIs."),
    ("Event group chat & feedback", "Not connected", "Local messages and session-only feedback", "Add group messages, attendee authorization, and post-event feedback APIs."),
    ("Notifications hub", "Not connected", "Seed notification list", "Add inbox, read state, deletion, and pagination."),
    ("Notification preferences & push", "Not connected", "Local UI state", "Persist preferences and register/revoke device push tokens."),
    ("Subscriptions & payments", "Not connected", "Static plans / test membership controller", "Add plan, subscription, checkout, verification, webhook, cancellation, and restore APIs."),
    ("Wallet & redemption", "Not connected", "Local ledger; backend deduction explicitly pending", "Add balance, transaction, top-up, and redemption APIs."),
    ("Gift catalog & sending", "Not connected", "Static catalog", "Add catalog and transactional gift-send APIs."),
    ("Referrals & leaderboard", "Not connected", "Static/local screens", "Add referral identity, invite tracking, and ranked leaderboard APIs."),
    ("AI coach, icebreakers & bio", "Not connected", "Static suggestions / local tools", "Add authenticated, moderated AI generation APIs with quotas."),
    ("Match explanations & dating recap", "Not connected", "Local/dummy insight presentation", "Add explanation and aggregate insight APIs."),
    ("Date spots", "Not connected", "ImageRepository venue list", "Add curated venue browse/detail APIs; map service remains separate."),
    ("Safety, reports & SOS", "Not connected", "Local/unavailable controls", "Add reports, evidence, emergency contacts, check-ins, and enforcement integration."),
    ("Account deactivation & deletion", "Not connected", "UI-only callbacks", "Add authenticated lifecycle APIs, data-retention handling, and session revocation."),
    ("Data export & privacy", "Not connected", "UI-only request flow", "Add asynchronous export request/status/download workflow."),
    ("Admin & host dashboards", "Not connected", "Dummy admin/host records", "Add role-protected metrics, moderation, user, report, and hosted-event APIs."),
    ("Success stories & support submissions", "Not connected", "Static stories/FAQ", "Add moderated story submission and support-ticket APIs where product scope requires server handling."),
]

missing_groups = {
    "Profile, account, privacy": [
        ("P01", "GET", "/api/profiles/:userId", "Public profile detail", "P0"),
        ("P02", "PUT", "/api/onboarding/photos/order", "Photo reorder", "P1"),
        ("P03", "POST", "/api/onboarding/voice-prompt", "Voice prompt upload", "P1"),
        ("P04", "DELETE", "/api/onboarding/voice-prompt", "Voice prompt removal", "P1"),
        ("P05", "POST", "/api/onboarding/video-prompt", "Video prompt upload", "P1"),
        ("P06", "DELETE", "/api/onboarding/video-prompt", "Video prompt removal", "P1"),
        ("P07", "POST", "/api/profile/verification", "KYC submission", "P1"),
        ("P08", "GET", "/api/profile/verification/status", "KYC status", "P1"),
        ("P09", "POST", "/api/account/deactivate", "Deactivate account", "P0"),
        ("P10", "POST", "/api/account/reactivate", "Reactivate account", "P1"),
        ("P11", "DELETE", "/api/account", "Delete account", "P0"),
        ("P12", "POST", "/api/privacy/data-exports", "Request data export", "P1"),
        ("P13", "GET", "/api/privacy/data-exports/:exportId", "Export status/download", "P1"),
        ("P14", "PUT", "/api/profile/compatibility-answers", "Compatibility answers", "P1"),
        ("P15", "PUT", "/api/profile/dealbreakers", "Dealbreakers", "P1"),
    ],
    "Matches and relationships": [
        ("R01", "GET", "/api/matches", "Match list", "P0"),
        ("R02", "GET", "/api/matches/:matchId", "Match detail", "P0"),
        ("R03", "DELETE", "/api/matches/:matchId", "Unmatch", "P0"),
        ("R04", "GET", "/api/likes/received", "Liked-you list", "P1"),
        ("R05", "GET", "/api/likes/sent", "Sent likes", "P1"),
        ("R06", "DELETE", "/api/likes/:targetUserId", "Revoke like", "P2"),
        ("R07", "GET", "/api/super-likes/entitlement", "Super-like balance", "P1"),
        ("R08", "GET", "/api/saved-profiles", "Saved profiles", "P1"),
        ("R09", "POST", "/api/saved-profiles/:userId", "Save profile", "P1"),
        ("R10", "DELETE", "/api/saved-profiles/:userId", "Remove saved profile", "P1"),
        ("R11", "GET", "/api/blocks", "Blocked profiles", "P0"),
        ("R12", "POST", "/api/blocks/:userId", "Block user", "P0"),
        ("R13", "DELETE", "/api/blocks/:userId", "Unblock user", "P0"),
    ],
    "Chat and realtime": [
        ("C01", "GET", "/api/conversations", "Conversation list", "P0"),
        ("C02", "POST", "/api/conversations", "Create/resolve conversation", "P0"),
        ("C03", "GET", "/api/conversations/:conversationId/messages", "Message history", "P0"),
        ("C04", "POST", "/api/conversations/:conversationId/messages", "Send message", "P0"),
        ("C05", "PUT", "/api/conversations/:conversationId/read", "Read receipt", "P0"),
        ("C06", "POST", "/api/conversations/:conversationId/media", "Message media upload", "P1"),
        ("C07", "DELETE", "/api/messages/:messageId", "Delete message", "P1"),
        ("C08", "PUT", "/api/conversations/:conversationId/draft", "Persist draft", "P2"),
        ("C09", "DELETE", "/api/conversations/:conversationId/draft", "Clear draft", "P2"),
        ("C10", "POST", "/api/realtime/token", "Authenticated realtime session", "P0"),
    ],
    "Events and hosting": [
        ("E01", "GET", "/api/events", "Event browse/search", "P0"),
        ("E02", "GET", "/api/events/:eventId", "Event detail", "P0"),
        ("E03", "POST", "/api/events/:eventId/registration", "Register", "P0"),
        ("E04", "DELETE", "/api/events/:eventId/registration", "Cancel registration", "P0"),
        ("E05", "POST", "/api/events/:eventId/waitlist", "Join waitlist", "P0"),
        ("E06", "DELETE", "/api/events/:eventId/waitlist", "Leave waitlist", "P0"),
        ("E07", "GET", "/api/events/me", "My Events", "P0"),
        ("E08", "POST", "/api/events/:eventId/check-in", "Event check-in", "P1"),
        ("E09", "POST", "/api/events/:eventId/feedback", "Post-event feedback", "P1"),
        ("E10", "GET", "/api/events/:eventId/group-chat/messages", "Event group messages", "P1"),
        ("E11", "POST", "/api/events/:eventId/group-chat/messages", "Send event group message", "P1"),
        ("E12", "GET", "/api/host/dashboard", "Host dashboard", "P2"),
        ("E13", "POST", "/api/host/events", "Create hosted event", "P2"),
        ("E14", "PUT", "/api/host/events/:eventId", "Update hosted event", "P2"),
    ],
    "Notifications": [
        ("N01", "GET", "/api/notifications", "Notification inbox", "P1"),
        ("N02", "PUT", "/api/notifications/:notificationId/read", "Mark read", "P1"),
        ("N03", "PUT", "/api/notifications/read-all", "Mark all read", "P1"),
        ("N04", "DELETE", "/api/notifications/:notificationId", "Delete notification", "P2"),
        ("N05", "GET", "/api/notification-preferences", "Read preferences", "P1"),
        ("N06", "PUT", "/api/notification-preferences", "Save preferences", "P1"),
        ("N07", "POST", "/api/devices/push-token", "Register push token", "P1"),
        ("N08", "DELETE", "/api/devices/push-token", "Revoke push token", "P1"),
    ],
    "Subscriptions, payments, wallet, gifts, referrals": [
        ("M01", "GET", "/api/subscriptions/plans", "Subscription plans", "P1"),
        ("M02", "GET", "/api/subscriptions/me", "Current membership", "P1"),
        ("M03", "POST", "/api/payments/orders", "Create checkout order", "P1"),
        ("M04", "POST", "/api/payments/verify", "Verify payment", "P1"),
        ("M05", "POST", "/api/payments/webhook", "Provider webhook", "P1"),
        ("M06", "POST", "/api/subscriptions/cancel", "Cancel renewal", "P1"),
        ("M07", "POST", "/api/subscriptions/restore", "Restore purchase", "P1"),
        ("M08", "GET", "/api/wallet", "Wallet balance", "P1"),
        ("M09", "GET", "/api/wallet/transactions", "Wallet ledger", "P1"),
        ("M10", "POST", "/api/wallet/top-up/orders", "Wallet top-up", "P2"),
        ("M11", "POST", "/api/wallet/redemptions", "Redeem wallet value", "P1"),
        ("M12", "GET", "/api/gifts", "Gift catalog", "P2"),
        ("M13", "POST", "/api/gifts/send", "Send gift", "P2"),
        ("M14", "GET", "/api/referrals/me", "Referral status", "P2"),
        ("M15", "POST", "/api/referrals/invitations", "Track invitation", "P2"),
        ("M16", "GET", "/api/referrals/leaderboard", "Referral leaderboard", "P2"),
        ("M17", "POST", "/api/boosts/purchase", "Purchase/consume boost", "P1"),
    ],
    "AI, insights, venues, social proof": [
        ("A01", "POST", "/api/ai/icebreakers", "Generate icebreakers", "P2"),
        ("A02", "POST", "/api/ai/dating-coach", "Dating coach", "P2"),
        ("A03", "POST", "/api/ai/bio-suggestions", "Bio suggestions", "P2"),
        ("A04", "GET", "/api/ai/matches/:matchId/explanation", "Why-we-matched", "P2"),
        ("A05", "GET", "/api/insights/dating-recap", "Dating recap", "P2"),
        ("A06", "GET", "/api/date-spots", "Date spot browse", "P2"),
        ("A07", "GET", "/api/success-stories", "Success-story browse", "P2"),
        ("A08", "POST", "/api/success-stories", "Submit success story", "P2"),
    ],
    "Safety, support, moderation": [
        ("S01", "POST", "/api/reports", "Report user/content", "P0"),
        ("S02", "POST", "/api/reports/:reportId/evidence", "Upload report evidence", "P0"),
        ("S03", "GET", "/api/safety/emergency-contacts", "Emergency contacts", "P1"),
        ("S04", "POST", "/api/safety/emergency-contacts", "Add emergency contact", "P1"),
        ("S05", "DELETE", "/api/safety/emergency-contacts/:contactId", "Remove emergency contact", "P1"),
        ("S06", "POST", "/api/safety/check-ins", "Start SOS check-in", "P1"),
        ("S07", "PUT", "/api/safety/check-ins/:checkInId", "Update/close check-in", "P1"),
        ("S08", "POST", "/api/support/tickets", "Create support ticket", "P2"),
        ("S09", "GET", "/api/support/tickets", "Support ticket history", "P2"),
        ("S10", "GET", "/api/admin/dashboard", "Admin metrics", "P2"),
        ("S11", "GET", "/api/admin/users", "Admin user search", "P2"),
        ("S12", "PUT", "/api/admin/users/:userId/status", "Moderate account", "P1"),
        ("S13", "GET", "/api/admin/reports", "Moderation queue", "P1"),
        ("S14", "PUT", "/api/admin/reports/:reportId", "Resolve report", "P1"),
    ],
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("AMORAA API Audit  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("6B7280")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_table_borders(table, color: str = "B7C5D5", size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths, font_size=8.1, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, label in enumerate(headers):
        cell = header.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(label))
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if status_col is not None and i == status_col:
                text = str(value)
                set_cell_shading(cell, GREEN if text == "Implemented" else AMBER if text in {"Incomplete", "Partial", "Client only"} else RED)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(DARK)
        if len(table.rows) % 2 == 1:
            for cell in row.cells:
                if status_col is None or cell is not row.cells[status_col]:
                    set_cell_shading(cell, "FAFBFC")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=BLUE, size=8)
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(NAVY)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("AMORAA  |  ENGINEERING AUDIT")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    add_page_number(section.footer.paragraphs[0])


def build_document() -> None:
    missing_count = sum(len(items) for items in missing_groups.values())
    assert len(existing) == 30
    assert sum(1 for row in existing if row[4] == "Implemented") == 27
    assert sum(1 for row in existing if row[4] == "Incomplete") == 3
    assert sum(1 for row in existing if row[5] == "Connected") == 28
    assert missing_count == 99, missing_count
    assert len(feature_gaps) == 26

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "AMORAA Complete API Audit"
    props.subject = "Backend API inventory, gaps, and Flutter integration status"
    props.author = "Codex"
    props.keywords = "AMORAA, Flutter, Express, Sequelize, MySQL, API audit"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AMORAA")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("COMPLETE API AUDIT")
    r2.bold = True
    r2.font.size = Pt(29)
    r2.font.color.rgb = RGBColor.from_string(BLUE)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run("Backend inventory • implementation status • Flutter integration • missing API backlog")
    r3.font.size = Pt(13)
    r3.font.color.rgb = RGBColor.from_string(NAVY)

    add_callout(
        doc,
        "Audit conclusion",
        "The repository contains a real, coherent API foundation for authentication, onboarding, and discovery. It is not yet a full-stack implementation of the wider product: only three Dart source files make HTTP requests, while 26 active feature areas remain partly local, seeded, unavailable, or UI-only.",
    )
    add_table(
        doc,
        ["Measure", "Count", "Meaning"],
        [
            ("Registered backend endpoints", "30", "Express routes mounted in server.js"),
            ("Implemented", "27", "Concrete handler and persistence behavior present"),
            ("Incomplete / conditional", "3", "Route exists, but behavior is materially unfinished or deployment-gated"),
            ("Flutter-consumed endpoints", "28", "A production Flutter call site invokes the endpoint"),
            ("Missing API operations", str(missing_count), "Minimum operations inferred from active frontend capabilities"),
            ("Frontend feature areas with a backend gap", "26", "Active product areas that are partial, local, seeded, or disconnected"),
        ],
        [2.15, 0.75, 3.60],
        font_size=9.2,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Static audit date: {date(2026, 8, 10).strftime('%d %B %Y')}  |  Workspace: D:\\Projects\\amora_ai")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("6B7280")
    doc.add_page_break()

    add_heading(doc, "1. Scope and counting rules", 1)
    doc.add_paragraph(
        "This audit traces the active Express route tree through controllers, validation, Sequelize models, MySQL persistence, Flutter API clients, repositories/controllers, and production UI call sites. It also inventories local, seeded, mock, session-only, and explicitly unavailable frontend behavior."
    )
    add_heading(doc, "Status definitions", 2)
    add_table(
        doc,
        ["Status", "Definition"],
        [
            ("Implemented", "Route is registered, validation/authorization is present where required, and the handler performs the intended operation."),
            ("Incomplete", "Route is registered, but a key behavior is conditional, placeholder-based, or lacks required business enforcement."),
            ("Connected", "At least one production Flutter call site invokes the route through an API client."),
            ("Client only", "A Dart client method exists, but no production feature call site invokes it."),
            ("Missing", "No registered backend route implements the frontend capability; the proposed path is a design recommendation, not existing code."),
        ],
        [1.2, 5.3],
        font_size=9.1,
    )
    add_heading(doc, "Repository coverage", 2)
    add_bullets(doc, [
        "202 Dart files under lib, organized into 37 feature directories.",
        "67 Flutter route-map entries in lib/main.dart; 74 route-name declarations overall, including aliases and unregistered roadmap prototypes.",
        "Only three Dart files contain the app's HTTP implementations: auth_service.dart, onboarding_api_service.dart, and discover_api_service.dart.",
        "32 backend JavaScript files, three Express route modules, and eight Sequelize models.",
        "The audit is source- and test-evidence based. It did not mutate or inspect production data and did not execute live Postman/MySQL integration tests against a running environment.",
    ])

    add_heading(doc, "2. Executive findings", 1)
    add_callout(doc, "Coverage is concentrated", "Authentication, onboarding, and Discover are the only backend-connected feature families. The rest of the visible product surface substantially exceeds the current server surface.", fill=AMBER)
    add_bullets(doc, [
        "The 30 registered endpoints divide into health (1), authentication (11), onboarding (12), and discovery (6).",
        "Twenty-eight endpoint routes are called by Flutter. The onboarding delete-photo method exists in the Dart client but has no production call site; /health is intentionally backend-only.",
        "The Discover feed is the highest-risk existing endpoint: most filters, scoring, sorting, and pagination happen after loading eligible users into Node memory. Communication Style alone is pushed into the Sequelize WHERE clause.",
        "Distance is currently a city-based placeholder (5 km for same city, otherwise 80 km); presence is always Offline; premium is always false; online-now and event-interest filters produce no matches.",
        "The Boost endpoint creates a free 30-minute boost and is not tied to wallet, subscription, entitlement, inventory, or payment state.",
        "Database bootstrapping still uses sequelize.sync(), and OnboardingProfile adds many columns dynamically in an afterSync hook. Only the Communication Style change has a reviewed migration-style file.",
        "The frontend's profile editor persists primarily to SharedPreferences. Existing onboarding update endpoints can be reused, but the edit screens do not enqueue the same remote sync path.",
        "Blocks are not modeled or enforced in Discover, messaging, or matches because no Block model or API exists.",
    ])
    add_heading(doc, "Priority order", 2)
    add_table(
        doc,
        ["Priority", "Workstream", "Why first"],
        [
            ("P0", "Safety and relationship integrity", "Blocks/reports, match APIs, account lifecycle, and enforcement must precede social scale."),
            ("P0", "Discover query hardening", "Move filtering, exclusions, ordering, and pagination into SQL; replace placeholder distance/presence."),
            ("P0", "Chat foundation", "The current chat experience is entirely local and seeded."),
            ("P0", "Events core", "Catalog, registration, waitlist, and My Events are visible but local/dummy."),
            ("P1", "Commerce and entitlements", "Payment, subscription, wallet, gifts, and boosts need transactional server state."),
            ("P1", "Production schema migrations", "Eliminate sync/afterSync schema mutation before production deployment."),
        ],
        [0.7, 2.0, 3.8],
        font_size=8.8,
    )

    doc.add_page_break()
    add_heading(doc, "3. Complete existing API catalogue", 1)
    doc.add_paragraph("All paths below are registered by the active server. “Frontend” reflects production call-site evidence, not merely the existence of a Dart method.")
    domain_ranges = [
        ("Health", 1, 1),
        ("Authentication", 2, 12),
        ("Onboarding and profile capture", 13, 24),
        ("Discover", 25, 30),
    ]
    for name, start, end in domain_ranges:
        add_heading(doc, name, 2)
        rows = [row for row in existing if start <= row[0] <= end]
        add_table(doc, ["#", "Method", "Path", "Auth", "Backend", "Frontend", "Finding"], rows, [0.32, 0.52, 1.75, 0.55, 0.78, 0.75, 1.83], font_size=7.4, status_col=4)

    add_heading(doc, "4. Incomplete or conditional endpoints", 1)
    add_table(
        doc,
        ["Endpoint", "Observed behavior", "Completion requirement"],
        [
            ("POST /api/auth/google", "Returns GOOGLE_AUTH_NOT_CONFIGURED (503) when GOOGLE_CLIENT_IDS is absent.", "Configure approved client IDs in each environment and add success/failure integration tests."),
            ("GET /api/discover/feed", "Uses in-memory filtering/sort/slice for most criteria; placeholder distance, presence, and premium; online/event filters are nonfunctional.", "Build SQL/Sequelize predicates and joins before LIMIT/OFFSET; add geospatial/presence/event data and complete exclusions."),
            ("POST /api/discover/boost", "Always grants a new free 30-minute boost when none is active.", "Consume a server-owned entitlement or verified purchase in a transaction; define idempotency and abuse limits."),
        ],
        [1.65, 2.35, 2.5],
        font_size=8.6,
    )
    add_heading(doc, "Implemented endpoints with important findings", 2)
    add_bullets(doc, [
        "POST /api/auth/resend-verification-code creates an OTP for any valid-looking phone number without first confirming a matching unverified account.",
        "Backend/README.md shows verify-account with an email field, but the active validator/controller require phoneNumber. The documentation example is stale.",
        "GET/PUT /api/discover/filters correctly persist values, but onlineNow and hasEventInterest cannot yield valid feed matches until Discover data sources are implemented.",
        "Refresh-token lookup compares the presented token against every unexpired stored hash. A token identifier would avoid an O(n) scan.",
        "The /uploads static route exposes uploaded assets; production requires an explicit access/CDN policy and malware/content checks appropriate for dating media.",
    ])

    doc.add_page_break()
    add_heading(doc, "5. Flutter-to-backend integration status", 1)
    add_table(
        doc,
        ["Integration measure", "Count", "Detail"],
        [
            ("Connected endpoint routes", "28", "11 auth + 11 onboarding + 6 Discover"),
            ("Client method without UI caller", "1", "DELETE /api/onboarding/photos/:index"),
            ("Operational-only endpoint", "1", "GET /health"),
            ("HTTP implementation files", "3", "AuthService, OnboardingApiService, DiscoverApiService"),
        ],
        [2.2, 0.75, 3.55],
        font_size=9,
    )
    add_heading(doc, "Connected flows", 2)
    add_bullets(doc, [
        "Authentication: signup, verification/resend, login, Google sign-in request, password recovery/reset, session restore, refresh, /me, and logout.",
        "Onboarding: status hydration, age, gender, interested-in, relationship goals, location, starter profile, profile completion, photo upload, primary photo, and completion.",
        "Discover: paginated feed requests, swipe, rewind, boost, filter load, and filter save/apply/reset.",
    ])
    add_heading(doc, "Active frontend feature areas with a backend gap", 2)
    add_table(doc, ["Feature area", "Status", "Current source of truth", "Required connection"], feature_gaps, [1.35, 0.72, 1.65, 2.78], font_size=7.8)
    add_callout(doc, "Local-by-design settings", "Theme, language, accessibility presentation, and offline-mode preferences may legitimately remain device-local. They are not counted as missing product APIs unless the product requires cross-device synchronization.")
    add_callout(doc, "Roadmap screens", "The roadmap/premium prototype files declare additional routes, but those screens are not registered in the active route table. They are excluded from the 26 active feature-gap count and from the 99-operation minimum backlog.", fill=GRAY)

    doc.add_page_break()
    add_heading(doc, "6. Missing API backlog", 1)
    doc.add_paragraph(
        "The following 99 operations are the minimum explicit server surface inferred from active UI actions and missing production capabilities. Paths are recommendations and should be normalized during API design; they do not exist in the repository today. Existing onboarding endpoints should be reused for ordinary profile field updates rather than duplicated."
    )
    add_table(
        doc,
        ["Priority", "Meaning"],
        [("P0", "Safety, identity, integrity, or core product blocker"), ("P1", "Required for a production-complete active feature"), ("P2", "Secondary, premium, administrative, or later-stage capability")],
        [0.8, 5.7],
        font_size=9,
    )
    for group, items in missing_groups.items():
        add_heading(doc, f"{group} ({len(items)})", 2)
        add_table(doc, ["ID", "Method", "Proposed path", "Frontend capability", "Pri"], items, [0.45, 0.62, 2.55, 2.35, 0.53], font_size=7.7)

    add_heading(doc, "7. Database and Sequelize audit", 1)
    add_table(
        doc,
        ["Model / table", "Purpose", "Audit status"],
        [
            ("User / Users", "Account identity and auth provider", "Present"),
            ("OtpToken / OtpTokens", "Hashed verification/reset OTPs", "Present"),
            ("RefreshToken / RefreshTokens", "Hashed refresh sessions", "Present"),
            ("OnboardingProfile / OnboardingProfiles", "Profile and onboarding data", "Present; several columns added dynamically after sync"),
            ("DiscoverAction / DiscoverActions", "Pass/like/super-like", "Present"),
            ("Match / Matches", "Reciprocal matches", "Present; no list/detail API"),
            ("Boost / Boosts", "Active boost periods", "Present; no entitlement/payment relation"),
            ("DiscoverFilterPreference / DiscoverFilterPreferences", "Persisted filters", "Present"),
        ],
        [1.75, 2.2, 2.55],
        font_size=8.6,
    )
    add_heading(doc, "Missing persistence domains", 2)
    doc.add_paragraph("No Sequelize models were found for conversations/messages, blocks, events/registrations/waitlists, notifications/device tokens, subscriptions/payments, wallet transactions, gifts, referrals, KYC, reports, emergency contacts/check-ins, support tickets, admin roles/permissions, or success-story submissions.")
    add_heading(doc, "Schema readiness", 2)
    add_bullets(doc, [
        "initializeDatabase creates the database, runs the Communication Style migration's up function, then calls sequelize.sync().",
        "OnboardingProfile.afterSync describes the table and adds a long list of missing fields dynamically.",
        "This is convenient for development but is not a controlled production migration strategy. Replace all implicit schema mutation with ordered, reversible Sequelize migrations before deployment.",
        "Communication Style is correctly modeled as a nullable single-select ENUM on OnboardingProfiles and as a JSON array on DiscoverFilterPreferences.",
    ])

    add_heading(doc, "8. Security, authorization, and data integrity", 1)
    add_table(
        doc,
        ["Area", "Finding", "Action"],
        [
            ("Authentication", "Protected onboarding/discover routes use bearer auth and derive identity from the token.", "Keep this pattern for every new API; never accept client userId as identity."),
            ("Blocks", "No block table/API or Discover exclusion exists.", "Implement before chat/matches expansion and enforce bidirectionally."),
            ("Media", "Multer enforces type/size/count, but no content scan/moderation workflow is visible.", "Add quarantine, malware scan, moderation, and signed/CDN access policy."),
            ("OTP", "Rate limiters exist; resend does not verify an eligible account.", "Bind resend to an unverified account and preserve non-enumeration."),
            ("Commerce", "Boost activation is not transactionally tied to entitlement.", "Use provider verification, idempotency keys, ledger records, and server-side consumption."),
            ("Admin/host", "Visible dashboards have no role-protected server surface.", "Introduce roles/permissions and audit logging before endpoints."),
        ],
        [1.05, 2.65, 2.8],
        font_size=8.3,
    )

    add_heading(doc, "9. Test and quality evidence", 1)
    add_table(
        doc,
        ["Check", "Result", "Evidence"],
        [
            ("Backend npm test", "PASS", "5/5 Node tests pass; all cover Communication Style parsing/query semantics."),
            ("Flutter analyze", "FAIL", "5 issues: two auth_experience_test compile errors, one http_parser dependency warning, two signup brace notices."),
            ("Flutter test", "INCONCLUSIVE", "Full suite exceeded the 120-second audit timeout without a final result."),
            ("Test inventory", "Observed", "66 Flutter test files / 490 test declarations; one backend test file / 5 test cases."),
            ("API integration coverage", "LOW", "No route-level integration suite for auth/onboarding/discovery, auth middleware, MySQL transactions, or pagination was found."),
            ("Live API/MySQL", "NOT RUN", "No running-environment mutation was required for this static audit."),
        ],
        [1.25, 0.9, 4.35],
        font_size=8.7,
    )
    add_heading(doc, "Current Flutter analyzer issues", 2)
    add_bullets(doc, [
        "lib/features/auth/presentation/signup_screen.dart:264 and :287 — curly_braces_in_flow_control_structures.",
        "lib/features/onboarding/data/onboarding_api_service.dart:7 — http_parser is imported directly but is not declared in pubspec.yaml.",
        "test/auth_experience_test.dart:838-839 — AccountVerificationScreen test uses obsolete destination and omits required phoneNumber.",
    ])

    doc.add_page_break()
    add_heading(doc, "10. Recommended implementation sequence", 1)
    add_table(
        doc,
        ["Phase", "Outcome", "Primary deliverables"],
        [
            ("1", "Make current foundation production-safe", "Migration baseline; Discover SQL pagination; real distance/presence; Google environment; API integration tests."),
            ("2", "Enforce relationship safety", "Blocks, reports, account lifecycle, match list/detail/unmatch, profile detail, eligibility exclusions."),
            ("3", "Enable real communication", "Conversation/message schema, APIs, realtime authentication, media, read state, moderation."),
            ("4", "Enable real events", "Event/registration/waitlist/My Events/group chat/check-in/feedback plus host role."),
            ("5", "Enable monetization", "Plans, subscriptions, payment verification/webhooks, wallet ledger, boost entitlement, gifts."),
            ("6", "Connect secondary product areas", "Notifications/push, referrals, AI services, insights, venues, success stories, support/admin."),
        ],
        [0.55, 1.85, 4.1],
        font_size=8.7,
    )
    add_heading(doc, "Acceptance criteria for each new API", 2)
    add_bullets(doc, [
        "Authenticated identity, authorization, request validation, stable error codes, and consistent response envelope.",
        "Sequelize query or transaction with a reviewed reversible migration; no production sync/alter behavior.",
        "Filtering before pagination and no client-side substitute for server eligibility rules.",
        "Flutter repository/controller integration, loading/error/empty behavior, and removal of local/dummy production paths.",
        "Backend unit/integration tests, Flutter serialization/integration tests, and representative authenticated Postman/cURL evidence.",
    ])

    add_heading(doc, "11. Source evidence index", 1)
    add_table(
        doc,
        ["Area", "Primary files inspected"],
        [
            ("Server registration, routes & tests", "Backend/src/server.js; Backend/src/routes/*.js; lib/main.dart and routeName declarations under lib/features; Backend/test/communicationStyleFilter.test.js; test/*.dart"),
            ("Handlers", "Backend/src/controllers/authController.js; onboardingController.js; discoverController.js"),
            ("Persistence", "Backend/src/models/*.js; Backend/src/config/db.js; Backend/src/migrations/202608100001-add-communication-style.js"),
            ("Flutter HTTP", "lib/core/auth/auth_service.dart; lib/features/onboarding/data/onboarding_api_service.dart; lib/features/discover/data/discover_api_service.dart"),
            ("Flutter integration", "lib/features/onboarding/data/local_onboarding_repository.dart; discover_action_controller.dart; browse_grid_screen.dart; advanced_filters_screen.dart"),
            ("Local/dummy gaps", "local_chat_repository.dart; local_profile_repository.dart; events_dummy_data.dart; event_participation_controller.dart; amora_dummy_data.dart; profile_relationship_controller.dart"),
        ],
        [1.35, 5.15],
        font_size=7.7,
    )
    add_heading(doc, "12. Conclusion", 1)
    add_callout(doc, "Bottom line", "AMORAA currently has 30 registered backend endpoints: 27 implemented and 3 incomplete/conditional. Flutter invokes 28 of those routes, but only three feature families are network-backed. Completing the active product requires an estimated 99 additional API operations across eight backend domains, plus production migration and integration-test work.", fill=LIGHT_BLUE)
    add_bullets(doc, [
        "Protect the existing foundation first: production migrations, Discover SQL pagination, blocks/reports, and route-level integration tests.",
        "Then connect the highest-value visible gaps: matches, chat, events, account lifecycle, and notifications.",
        "Treat commerce, AI, admin/host, and roadmap capabilities as separate bounded programs with explicit authorization and data models.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
