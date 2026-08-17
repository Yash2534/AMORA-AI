from __future__ import annotations

from pathlib import Path
from datetime import date
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs"
WORK_DIR = ROOT / "artifacts" / "third_party_api_report"
OUT_FILE = WORK_DIR / "AMORAA_Third-Party_API_Implementation_Report_draft.docx"
DIAGRAM_FILE = WORK_DIR / "amoraa_target_integration_architecture.png"

NAVY = "19283F"
ROSE = "B54A66"
ROSE_LIGHT = "F8E9EE"
BLUE = "2E74B5"
BLUE_LIGHT = "E8EEF5"
INK = "26313D"
MUTED = "667085"
GRAY = "F2F4F7"
MID_GRAY = "D7DCE3"
GREEN = "217A57"
GREEN_LIGHT = "E8F5EE"
AMBER = "946200"
AMBER_LIGHT = "FFF4D6"
RED = "A33A3A"
RED_LIGHT = "FCEBEC"
WHITE = "FFFFFF"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def shade(element, fill: str):
    props = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            cell.width = Inches(widths_dxa[min(idx, len(widths_dxa) - 1)] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_font(run, name="Aptos", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(c)
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_picture_with_alt(doc: Document, image_path: Path, width, title: str, description: str):
    shape = doc.add_picture(str(image_path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)
    return shape


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, fld_text, fld_end):
        run._r.append(node)


def add_toc(doc: Document):
    """Insert an updateable Word table-of-contents field."""
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Open in Word and update fields if page numbers are not refreshed automatically."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(node)
    doc.add_page_break()


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_callout(doc, title: str, body: str, kind="info"):
    colors = {
        "info": (BLUE_LIGHT, BLUE),
        "decision": (AMBER_LIGHT, AMBER),
        "risk": (RED_LIGHT, RED),
        "success": (GREEN_LIGHT, GREEN),
    }
    fill, accent = colors[kind]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    shade(p._p, fill)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(f"{title}: ")
    set_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(body)
    set_font(r, size=10.5, color=INK)
    return p


def add_bullets(doc, items: Iterable[str], level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_numbered(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_label_value(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=NAVY)
    p.add_run(value)
    return p


def add_key_table(doc, rows, widths=(2160, 7200)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        set_cant_split(row)
        cells = row.cells
        cells[0].text = label
        cells[1].text = value
        shade(cells[0]._tc, GRAY)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, size=9.5, color=NAVY, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, size=9.5, color=INK)
    set_table_geometry(table, list(widths))
    table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_meta_table(doc, integration):
    rows = [
        ("Priority", integration["priority"]),
        ("Disposition", integration["disposition"]),
        ("Current AMORAA state", integration["current"]),
        ("Client dependency", integration["dependency"]),
        ("Engineering size", integration["effort"]),
        ("Primary layer", integration["layer"]),
    ]
    return add_key_table(doc, rows)


def add_integration_section(doc, item):
    h = doc.add_paragraph(style="Heading 1")
    h.add_run(f'{item["num"]}. {item["name"]}')
    keep_with_next(h)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(item["purpose"])
    set_font(r, size=11.5, color=MUTED, italic=True)
    add_meta_table(doc, item)
    add_callout(doc, "Recommendation", item["recommendation"], "decision" if "Decision" in item["disposition"] or "Choose" in item["disposition"] else "info")

    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Implementation and connection to AMORAA")
    keep_with_next(h2)
    add_label_value(doc, "SDK/API", item["sdk"])
    add_label_value(doc, "Flutter/web changes", item["frontend"])
    add_label_value(doc, "Node.js/Express changes", item["backend"])
    add_label_value(doc, "MySQL requirements", item["database"])
    add_label_value(doc, "Configuration/environment", item["env"])
    add_label_value(doc, "Webhooks/callbacks", item["webhooks"])

    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Development team will implement")
    keep_with_next(h2)
    add_bullets(doc, item["team"])

    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Client must provide or complete")
    keep_with_next(h2)
    add_bullets(doc, item["client"])

    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Implementation sequence")
    keep_with_next(h2)
    add_numbered(doc, item["steps"])

    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Testing and acceptance")
    keep_with_next(h2)
    add_bullets(doc, item["tests"])

    h2 = doc.add_paragraph(style="Heading 2")
    h2.add_run("Security and privacy controls")
    keep_with_next(h2)
    add_bullets(doc, item["security"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Official implementation reference: ")
    set_font(r, size=9, color=MUTED, bold=True)
    add_hyperlink(p, item["source_label"], item["source"])
    doc.add_page_break()


def integration_records():
    return [
        dict(num=1, name="Firebase Analytics", purpose="Canonical first-party mobile/app analytics, funnels, cohorts, retention, and user properties.", priority="P0", disposition="Implement", current="Not installed; no Firebase config files or FlutterFire initialization.", dependency="High - final app IDs, Firebase/GCP ownership, privacy decisions; 1-3 client days.", effort="M", layer="Flutter (Android/iOS/Web) + reporting", recommendation="Use Firebase Analytics as the canonical app event collector. The same event dictionary must feed GA4 and the chosen MMP; do not create competing names for the same signup, match, subscription, or retention events.", sdk="FlutterFire firebase_core and firebase_analytics; Firebase/GA4 app data streams.", frontend="Initialize Firebase before AuthService, add consent-aware analytics service, screen-view observer, stable user ID after login, and typed event wrappers. Exclude chat text, exact location, Aadhaar/KYC media, and other sensitive attributes.", backend="No Firebase SDK is required for normal client events. Backend will emit authoritative server events (payment, entitlement, verification) through a shared event outbox to GA4 Measurement Protocol/server routing only after consent rules are satisfied.", database="Add AnalyticsEventOutbox (eventId, userId nullable, eventName, payload JSON, consentState, destinationState, occurredAt) or reuse a generic IntegrationOutbox. Store no advertising identifiers in Users.", env="FIREBASE_* platform configuration via generated firebase_options.dart; release identifiers are non-secret. Server measurement secrets live only in the secret manager, never Flutter.", webhooks="None for the client SDK. Server event delivery uses provider HTTP responses and retry/dead-letter tracking.", team=["Define a versioned AMORAA event and user-property taxonomy with owners and PII classification.", "Instrument acquisition, signup, verification, onboarding, discover, match, chat-open, event registration, paywall, purchase, cancel, and retention milestones.", "Add debug-mode verification and a data-quality dashboard; document which events are client-observed versus server-authoritative.", "Add consent/opt-out behavior and account deletion/reset handling."], client=["Provide a company-controlled Google account and create/invite the team to a Firebase project and linked GA4 property.", "Approve the event taxonomy, conversion definitions, retention period, data-sharing settings, and whether advertising features are enabled.", "Provide the final Android application ID, iOS bundle ID, App Store ID when available, and production web domain.", "Provide approved Privacy Policy/Cookie Policy wording and analytics consent requirements by launch market."], steps=["Finalize identifiers and create dev/staging/prod Firebase apps.", "Install FlutterFire, generate per-platform config, and initialize consent defaults.", "Implement typed analytics facade and event schema validation.", "Instrument critical flows and server-authoritative event forwarding.", "Validate DebugView, user isolation, opt-out, and dashboard funnels before release."], tests=["Firebase DebugView confirms exact names/parameters once per action on Android, iOS, and web.", "Automated unit tests verify event mapping and block prohibited PII fields.", "Logout/account-switch tests clear user identity; consent-denied tests send no disallowed analytics.", "Reconcile signup/purchase counts against MySQL within an agreed tolerance."], security=["Use pseudonymous internal user IDs only after authentication; never send phone, email, Aadhaar, chat content, photos, or precise coordinates.", "Restrict console roles and enable MFA; separate production from non-production.", "Document deletion/retention behavior and consent evidence."], source_label="Firebase for Flutter setup", source="https://firebase.google.com/docs/flutter/setup"),

        dict(num=2, name="Google Analytics 4 (Web)", purpose="Web acquisition, audience, conversion, journey, and funnel reporting for the Flutter web property.", priority="P0", disposition="Implement", current="No GA4 web stream, Google tag, consent layer, or measurement code in web/index.html.", dependency="High - GA ownership, domain, consent/legal decisions; 1-3 client days.", effort="M", layer="Flutter Web + GA4", recommendation="Create one GA4 property with distinct production web and app streams where governance permits. Treat Firebase Analytics as the app collector and GA4/GTM as the web collector, with a shared event dictionary and deduplication rules.", sdk="Google tag/GA4 through GTM for web; Firebase Analytics for Flutter app streams; Measurement Protocol only for authoritative server events.", frontend="Add consent default before tag load, route/screen tracking for Flutter navigation, virtual page views, campaign parameters, and typed dataLayer events. Add canonical URL/referrer handling for SPA navigation.", backend="Issue server events only when client identifiers/consent are valid. Generate a stable event_id for purchase/signup deduplication and keep secrets server-side.", database="Use AnalyticsEventOutbox and optional WebAttributionTouch (userId, clientIdHash, campaign fields, consent, first/last touch) with strict retention; do not persist raw cookies unless approved.", env="GA4_MEASUREMENT_ID is public web config; GA4_API_SECRET and measurement endpoint configuration are backend secrets. Configure allowed domains and referral exclusions in GA.", webhooks="No inbound webhook. Measurement Protocol delivery must record provider response/retry state.", team=["Create web data stream and GTM-delivered GA configuration.", "Instrument virtual page views and conversion events in the Flutter router/dataLayer bridge.", "Configure cross-domain/referral rules, internal-traffic filters, retention, and BigQuery export decision.", "Build server-event reconciliation and duplicate suppression."], client=["Provide GA4 property owner access or authorize creation under the corporate Google organization.", "Provide final production/staging domains and DNS/hosting owner.", "Approve key events/conversions, audiences, retention, data-sharing, and consent-mode choices.", "Provide cookie/privacy notices and the marketing/legal owner who can approve analytics usage."], steps=["Create property/data stream and naming standard.", "Implement consent-first loader and Flutter web dataLayer bridge.", "Configure conversions and referral/domain rules.", "Verify in Realtime/DebugView and reconcile against backend records."], tests=["Tag Assistant and GA DebugView show SPA route changes without duplicates.", "UTM/gclid persistence survives login only within approved consent scope.", "Consent denied/granted/withdrawn scenarios behave as documented.", "Purchase and signup totals reconcile to MySQL/outbox records."], security=["API secrets never appear in web assets.", "Default consent state follows the approved jurisdictional policy.", "No sensitive dating preferences, messages, verification details, or raw identifiers enter GA."], source_label="GA4 consent types", source="https://support.google.com/analytics/answer/12334711"),

        dict(num=3, name="Google Tag Manager (Web)", purpose="Governed deployment of GA4, ad pixels, consent signals, and marketing tags without rebuilding the Flutter app.", priority="P0", disposition="Implement", current="No GTM container or dataLayer bootstrap in the web shell.", dependency="High - GTM admin access and marketing governance; 1-2 client days.", effort="M", layer="Flutter Web shell + marketing operations", recommendation="Use GTM only for web tags. Keep business-event generation in typed Flutter/backend code and expose a controlled dataLayer contract; do not allow ad-hoc tags to read the entire page or application state.", sdk="Google Tag Manager web container and dataLayer API.", frontend="Insert the approved GTM loader and consent defaults in web/index.html, then dispatch sanitized route and conversion objects from Flutter through a JS interop adapter.", backend="No direct runtime dependency. Backend supplies authoritative event IDs and, where approved, first-party campaign context to the web client/server-side tag path.", database="No GTM-specific table. Reuse event outbox and consent/audit records.", env="GTM_CONTAINER_ID per environment; container environments/auth strings for staging preview; CSP allow-list updates.", webhooks="None. Tags may call approved endpoints only; all vendor endpoints are documented in the tag register.", team=["Create a versioned dataLayer schema with a field allow-list and examples.", "Implement consent initialization before GTM and SPA route events after Flutter navigation.", "Configure workspaces, environments, approvals, naming, folders, and publishing workflow.", "Add Content Security Policy and automated container export backups."], client=["Provide corporate GTM account/container owner access and nominate marketing publishers/approvers.", "Approve the tag register, consent categories, and who may publish production changes.", "Provide production/staging domains and access to modify the web host/CSP.", "Provide all vendor pixel IDs only after the corresponding business accounts are verified."], steps=["Create account/containers and least-privilege roles.", "Define consent and dataLayer contracts.", "Install/verify on staging.", "Configure tags/triggers/variables and approvals.", "Publish production container with rollback/export."], tests=["GTM Preview shows only approved fields and triggers.", "Tag Assistant confirms tags are suppressed until the relevant consent.", "SPA navigation and duplicate-fire regression tests pass.", "CSP and page-load performance remain within agreed thresholds."], security=["No secrets in GTM variables; public pixel IDs only.", "Limit custom HTML tags and require review for templates/network destinations.", "Maintain container export, change log, and emergency rollback owner."], source_label="Google Tag Manager web help", source="https://support.google.com/tagmanager/answer/6102821"),

        dict(num=4, name="Server-side Google Tag Manager", purpose="First-party server event routing, transformation, validation, and controlled forwarding to analytics/advertising destinations.", priority="P1", disposition="Implement after web GTM and consent design", current="Absent; Express sends no marketing events and no tagging domain exists.", dependency="High - GCP/billing/DNS/consent ownership; typically 3-10 client days plus DNS propagation.", effort="L", layer="Backend/web + managed tagging infrastructure", recommendation="Deploy only after the client approves hosting cost and data governance. Use a dedicated first-party subdomain and retain the Express outbox as the system of record; sGTM is a router, not the entitlement/payment authority.", sdk="GTM server container on Cloud Run/App Engine or approved managed host; GA4 client and vendor server tags.", frontend="Point approved GA4/GTM traffic to the first-party tagging domain and preserve consent/event_id parameters.", backend="Add signed/allow-listed event ingress or direct server-container requests from the integration worker; validate schema, consent, timestamps, and retry behavior.", database="AnalyticsEventOutbox/IntegrationDeliveryAttempt plus provider event IDs and dead-letter status.", env="SGTM_BASE_URL, SGTM_AUTH_SECRET or network identity, GTM server container config, GCP project/region; DNS CNAME for analytics.<domain>.", webhooks="Outbound routing to GA4, Google Ads, Meta, TikTok as approved. No anonymous inbound business events.", team=["Provision a production-grade multi-instance server container and custom domain.", "Implement transformation/redaction templates and destination allow-list.", "Connect backend outbox with idempotent delivery and monitoring.", "Document operating cost, scaling, logging, rollback, and ownership."], client=["Approve server-side tagging scope, monthly hosting budget, GCP billing project, and data-processing destinations.", "Provide DNS access for the first-party tagging subdomain and TLS validation.", "Provide GTM/GA/Ads/Meta/TikTok administrator access and legal approval for server event sharing.", "Nominate a marketing-operations owner for ongoing container publishing."], steps=["Approve architecture and hosting model.", "Provision server container and custom domain.", "Implement redaction/client templates and Express outbox delivery.", "Connect one destination at a time, validate, then expand.", "Load/failover test and document runbook."], tests=["Preview/debug confirms field transformations and consent propagation.", "Duplicate, late, and malformed events are rejected or deduplicated.", "Outage/retry/dead-letter tests preserve events without double conversions.", "DNS, TLS, autoscaling, log redaction, and monitoring checks pass."], security=["Authenticate backend-to-container traffic and restrict admin roles.", "Redact IP/PII and drop prohibited dating/KYC/chat fields before vendors.", "Set short log retention and alert on new destinations or template changes."], source_label="Server-side GTM overview", source="https://developers.google.com/tag-platform/tag-manager/server-side"),

        dict(num=5, name="Meta Pixel + Conversions API", purpose="Meta campaign measurement, optimization, retargeting, and deduplicated browser/server conversion signals.", priority="P0", disposition="Implement for active Meta campaigns", current="Absent; no Meta Business assets, pixel, CAPI route, or consent mapping.", dependency="Very High - verified Meta business, dataset ownership, legal consent, ad account; often 1-4 weeks.", effort="L", layer="Web + backend", recommendation="Use Pixel for consented browser events and CAPI for authoritative signup/subscription events with the same event_name and event_id. Do not place Meta SDK/pixel in the mobile app unless the MMP and privacy plan explicitly require it.", sdk="Meta Pixel through GTM; Meta Conversions API/Marketing API from backend or sGTM.", frontend="Generate consented PageView/ViewContent/Lead events from the web dataLayer and capture only approved click/cookie parameters.", backend="Create CAPI adapter for CompleteRegistration, Subscribe, Purchase, Cancel/eligible lifecycle events; hash normalized contact fields only when legally approved and needed for match quality.", database="MarketingEventOutbox (or shared outbox), eventId, sourceUrl, consent snapshot, pixel/dataset destination status; store access tokens only in secrets manager.", env="META_PIXEL_ID, META_DATASET_ID, META_CAPI_ACCESS_TOKEN, META_TEST_EVENT_CODE (non-prod), META_API_VERSION.", webhooks="Outbound CAPI. Optional Meta lead/webhook ingestion is out of scope unless the client later requests it.", team=["Map AMORAA events to Meta standard/custom events and define deduplication.", "Implement Pixel via GTM and backend CAPI adapter with test-event mode.", "Configure domain verification, Aggregated Event Measurement priorities where applicable, and diagnostics.", "Add consent gating, retry, match-quality monitoring, and deletion workflow."], client=["Create/verify Meta Business Portfolio, ad account, dataset/pixel, and corporate domain; invite development and marketing roles.", "Provide payment method/ad billing and nominate business/admin owners.", "Approve the exact event list, audience/retargeting policy, contact-data hashing, retention, and consent language.", "Provide Privacy Policy/Cookie Policy and domain DNS access for verification."], steps=["Verify business/domain and create dataset.", "Approve event/data-sharing matrix.", "Install Pixel in staging and implement CAPI test events.", "Validate deduplication/diagnostics and consent behavior.", "Enable production campaigns after marketing sign-off."], tests=["Meta Test Events receives browser and server events with one conversion after deduplication.", "Event Match Quality is reviewed without adding prohibited data.", "Consent denied/withdrawn blocks future ad signals.", "Refund/cancel and duplicate webhook cases do not inflate purchase conversions."], security=["Keep access token server-side and rotate through the client-owned business account.", "Hash only normalized approved fields; hashing does not remove the need for consent/legal basis.", "Never transmit sexual orientation, religion, dating preferences, KYC, messages, or precise location."], source_label="Meta Conversions API", source="https://developers.facebook.com/docs/marketing-api/conversions-api/"),

        dict(num=6, name="Google Ads conversions", purpose="Google Ads conversion measurement, campaign optimization, and audience activation for approved acquisition programs.", priority="P0", disposition="Implement for active Google Ads campaigns", current="Absent; no conversion actions, linked GA4 property, tag, enhanced conversions, or backend upload.", dependency="Very High - Ads account/billing, GA/GTM access, consent/legal decisions; 2-10 client days.", effort="M/L", layer="Web + backend", recommendation="Prefer GA4 key-event import for broad web measurement and direct Google Ads conversion/server tags for revenue-critical events where deduplication and enhanced conversions are approved.", sdk="Google Ads conversion tags through GTM/sGTM; GA4 linking/import; optional Google Ads API for offline/enhanced conversion uploads.", frontend="Send consented gclid/gbraid/wbraid and conversion event IDs through the web dataLayer without exposing backend secrets.", backend="Persist approved click identifiers with attribution expiry, upload authoritative conversions/refunds, and record Google response/job status.", database="AttributionTouches and MarketingEventOutbox with click ID, event ID, conversion time/value/currency, consent state, destination status, expiry.", env="GOOGLE_ADS_CUSTOMER_ID, CONVERSION_ID/LABEL or action resource name, developer/OAuth credentials if API uploads are used; GA/GTM identifiers.", webhooks="Outbound tag/API uploads. Google Ads does not drive AMORAA entitlement state.", team=["Create conversion mapping and value rules for lead, verified signup, subscription, and renewal/cancel events.", "Link GA4/Ads and implement tags or API upload with one authoritative source per conversion.", "Implement click-ID capture/retention and consent-mode signals.", "Build reconciliation dashboard and campaign launch checklist."], client=["Provide Google Ads admin access, advertiser verification details, billing profile/payment method, and campaign owner.", "Provide GA4/GTM admin access and approve linked-product data sharing.", "Approve conversion values, attribution windows, customer-data usage, audiences, and consent requirements.", "Provide domain ownership/DNS access and legal/privacy documents."], steps=["Complete advertiser/account verification and link products.", "Approve conversion source/value model.", "Implement web and server signals with deduplication.", "Validate in Tag Assistant/Ads diagnostics.", "Launch campaigns only after reconciliation sign-off."], tests=["Test conversions appear with correct action, value, currency, timestamp, and transaction ID.", "Duplicate client/server events count once.", "Consent-mode and click-ID expiry tests pass.", "Refund/cancel corrections and backend reconciliation are verified."], security=["OAuth/API credentials remain server-side; use least-privilege client-owned accounts.", "Obtain required consent before enhanced-conversion/customer-data use.", "Prohibit sensitive profile/KYC/chat data in ads or audiences."], source_label="Server-side Google Ads conversions", source="https://developers.google.com/tag-platform/tag-manager/server-side/ads-setup"),

        dict(num=7, name="TikTok Events API", purpose="TikTok campaign measurement and optimization when TikTok is an approved acquisition channel.", priority="P1", disposition="Conditional - implement only when channel is funded", current="Absent; no pixel, event source, access token, or event mapping.", dependency="Very High - TikTok Business/Ads account, billing, legal approval; 1-4 weeks.", effort="M/L", layer="Web + backend", recommendation="Defer until the client confirms an active TikTok campaign plan. When used, implement Pixel plus Events API with shared event_id deduplication and the same consent/data-minimization controls as Meta.", sdk="TikTok Pixel through GTM and TikTok Events API 2.0 direct/sGTM integration.", frontend="Emit approved web events and click/cookie context only after consent.", backend="Send authoritative registration/subscription events with event_id, event_time, context, and approved match keys; record retries/diagnostics.", database="Reuse MarketingEventOutbox and AttributionTouches; no TikTok-specific user fields.", env="TIKTOK_PIXEL_CODE, TIKTOK_EVENTS_ACCESS_TOKEN, TIKTOK_API_BASE/VERSION, TIKTOK_TEST_EVENT_CODE.", webhooks="Outbound Events API only for this scope.", team=["Define the TikTok event/match-key matrix with marketing/legal.", "Configure Pixel and direct/server events with deduplication.", "Implement retry, diagnostics, consent enforcement, and destination kill switch.", "Document campaign QA and ongoing signal-quality ownership."], client=["Provide verified TikTok Business Center, Ads Manager account, pixel/event source, billing, and admin invitations.", "Approve that TikTok is an active channel and define campaign/conversion objectives.", "Approve match keys, audience policy, consent text, and privacy/data-sharing terms.", "Provide domain verification access if requested by the platform."], steps=["Complete business/ad setup.", "Approve data-sharing matrix.", "Implement Pixel and Events API in test mode.", "Validate deduplication/diagnostics.", "Enable production destination for campaign launch."], tests=["TikTok Events Manager diagnostics show browser/server events and deduplication.", "Consent-denied and data-field allow-list tests pass.", "Retries/duplicates/late events do not inflate conversions.", "Counts reconcile against MySQL and MMP reporting."], security=["Access token is a backend secret; disable destination when no campaigns are active.", "No sensitive dating, KYC, chat, or precise location data.", "Legal review is required before match keys or custom audiences."], source_label="TikTok Events API overview", source="https://ads.tiktok.com/help/article/events-api"),

        dict(num=8, name="AppsFlyer OR Adjust (Mobile Measurement Partner)", purpose="Install attribution, campaign source, deep-link attribution, in-app event measurement, and mobile marketing postbacks.", priority="P0", disposition="Choose exactly one MMP", current="Neither SDK exists; no ATT/advertising attribution flow, install table, or partner mappings.", dependency="Very High - vendor contract, store listings/IDs, ad accounts, privacy decisions; 1-4 weeks.", effort="L", layer="Flutter Android/iOS + backend", recommendation="Run a commercial/technical selection and implement one MMP only. Prefer the selected MMP's native deep-link product first; add Branch only if it supplies a documented capability the MMP cannot meet.", sdk="appsflyer_sdk plus OneLink, or adjust_sdk plus TrueLink/deep links; vendor S2S API/postbacks.", frontend="Initialize after consent policy, set internal customer ID, record standardized lifecycle/revenue events, handle attribution/deferred links, ATT/SKAdNetwork flow, and logout identity reset.", backend="Receive verified postbacks where needed, send authoritative revenue/subscription events, reconcile store/Razorpay transactions, and handle deletion/forget requests.", database="AttributionInstalls, AttributionTouches, ProviderPostbackEvents, and event outbox with provider IDs, campaign dimensions, consent, first/last touch, and expiry.", env="MMP_DEV_KEY/APP_TOKEN, iOS App ID, Android package, environment, deep-link domain/template IDs, S2S auth token; public app token may be in app only per vendor guidance.", webhooks="Attribution/install postbacks, partner postbacks, and data-deletion callbacks as contracted; all signed/authenticated and idempotent.", team=["Lead an AppsFlyer-vs-Adjust decision against pricing, channel coverage, deep linking, fraud, privacy, raw data, and support.", "Create the unified mobile event map and revenue-source rules.", "Implement SDK, consent/ATT, deep links, S2S events, and postback ingestion.", "Build attribution QA, discrepancy dashboard, deletion, and account-switch controls."], client=["Select and purchase one MMP plan; provide organization owner/admin invitations and billing contact.", "Provide final Android package, iOS bundle/App Store ID, store access, ad-network accounts, and campaign naming conventions.", "Approve ATT prompt strategy, consent wording, data retention/export, raw-data access, partner postbacks, and fraud settings.", "Provide the marketing attribution owner and acceptable discrepancy thresholds."], steps=["Vendor selection and contract/DPA.", "Create apps, event map, partner connections, and deep-link domain.", "Implement sandbox/dev SDK and consent/ATT behavior.", "Connect S2S revenue/postbacks and reconcile.", "Run live test campaigns before production sign-off."], tests=["Fresh install, reinstall, organic, paid, deferred deep link, and account-switch cases on physical devices.", "Sandbox/live-test campaign attribution appears under correct source/campaign.", "Revenue equals verified server transactions; SDK/S2S duplicates are suppressed.", "Opt-out/deletion and ATT-denied flows behave per policy."], security=["Keep S2S tokens and reporting credentials server-side.", "Treat advertising IDs and attribution data as personal data; apply retention and access controls.", "Do not send sensitive profile traits, KYC, messages, or photos as event parameters."], source_label="AppsFlyer SDK getting started", source="https://dev.appsflyer.com/hc/docs/getting-started"),

        dict(num=9, name="Branch", purpose="Branded direct/deferred deep links, campaign links, sharing links, and cross-platform routing.", priority="P1", disposition="Decision gate after MMP selection", current="Absent; app has named routes but no universal/app links, AASA/DAL files, or link routing service.", dependency="High - link domain/DNS, store IDs, Branch plan; 3-10 client days.", effort="L", layer="Flutter + web/domain + backend", recommendation="Do not add Branch automatically if AppsFlyer OneLink or Adjust deep links meet the requirement. If selected, define a narrow route allow-list for profiles/events/referrals and never encode personal data in URLs.", sdk="flutter_branch_sdk, Branch dashboard/link API, Android App Links, iOS Universal Links, web fallback.", frontend="Initialize link listener before routing, map canonical link payloads to safe named routes, handle cold/warm/deferred opens, logout identity reset, and invalid-link fallback.", backend="Optionally create signed/authorized share links and store referral/link ownership; validate any reward claim independently of link parameters.", database="DeepLinks/ReferralLinks (linkId, ownerUserId, targetType/id, campaign, status, expiresAt) and AttributionTouches only if programmatic links/referrals are used.", env="BRANCH_KEY per environment, BRANCH_SECRET server-only, BRANCH_LINK_DOMAIN, iOS team/bundle/App Store IDs, Android package/signing fingerprints.", webhooks="Optional Branch webhook/export callbacks; signed/authenticated and deduplicated.", team=["Decide overlap with MMP deep-link capability and document the reason for Branch.", "Configure link domains, AASA/assetlinks, route allow-list, and fallbacks.", "Implement Flutter cold/warm/deferred routing and backend link creation if needed.", "Add link abuse controls and campaign/share QA."], client=["Approve Branch as an additional vendor and plan; provide account owner/admin access.", "Provide final bundle/package/store IDs, Android release signing SHA-256, Apple Team ID, and DNS access.", "Approve branded link domain, link lifespan, referral rules, and fallback pages.", "Provide Privacy Policy/DPA approval and campaign naming owner."], steps=["Confirm vendor necessity.", "Configure apps/domains/store destinations.", "Implement safe link schema/routing.", "Implement API-generated links/referrals if required.", "Test every install/open/fallback matrix."], tests=["Installed, not-installed, cold-start, warm-start, logged-out, expired, and invalid links.", "Android App Links and iOS Universal Links validate against production signing/team IDs.", "No route permits unauthorized access to private profiles/chats.", "Referral/link reward cannot be replayed or self-awarded."], security=["Never put phone, email, Aadhaar, tokens, or sensitive profile fields in URLs.", "Backend re-authorizes target access after navigation.", "Use separate test/live keys and clear Branch identity on logout."], source_label="Branch Flutter SDK reference", source="https://help.branch.io/developer-hub/docs/flutter-sdk-full-reference"),

        dict(num=10, name="Firebase Cloud Messaging", purpose="Push notifications for transactional events, safety, re-engagement, messages, matches, events, and subscription status.", priority="P0", disposition="Complete existing partial integration", current="Backend has FCM HTTP v1 adapter, UserDevices and NotificationDeliveries; Flutter has no firebase_messaging, token registration, APNs capability, or Firebase config.", dependency="High - Firebase project, APNs key, app IDs/signing; 2-10 client days.", effort="M/L", layer="Flutter + backend", recommendation="Complete FCM because the backend delivery ledger already exists. Add OneSignal only if the client explicitly buys its campaign UI/segmentation and accepts replacing the current provider adapter.", sdk="firebase_messaging + firebase_core; Firebase Admin SDK/HTTP v1; APNs authentication key uploaded to Firebase.", frontend="Request permission contextually, obtain/rotate token, POST /api/devices, unregister on logout, handle foreground/background/terminated messages, route deep links, and show local foreground notifications where required.", backend="Replace raw credential handling with a shared Firebase Admin initialization where practical, support topic-free per-user sends, scheduled retry worker, invalid-token cleanup, quiet hours, collapse keys, and environment separation.", database="Existing UserDevices, Notifications, NotificationPreferences, NotificationDeliveries are suitable; add provider, appEnvironment, locale/timezone, token fingerprint, nextAttemptAt if needed.", env="FIREBASE_PROJECT_ID plus service-account/workload identity; GOOGLE_APPLICATION_CREDENTIALS or secret-manager fields; APNs key configured in Firebase; no service-account secret in Flutter.", webhooks="FCM delivery is outbound; token refresh is a client callback. Delivery metrics export is optional and privacy-reviewed.", team=["Add FlutterFire/FCM configuration and lifecycle handlers.", "Register/rotate/unregister tokens using the existing device endpoints.", "Harden backend retries, invalidation, quiet hours, deep-link payload schema, and observability.", "Create notification copy/category matrix and operational runbook."], client=["Provide owner access to a client-controlled Firebase project and approve dev/staging/prod app registrations.", "Provide final app identifiers and Apple Developer access/APNs .p8 key, Key ID, and Team ID; never email the private key in plaintext.", "Approve notification categories, opt-in copy, quiet hours, retention, and re-engagement policy.", "Provide physical Android/iOS test devices and production signing/store access."], steps=["Finalize app IDs and Firebase apps.", "Configure Android/iOS/APNs and Flutter SDK.", "Connect token lifecycle to /api/devices.", "Harden backend delivery/retry and payload routes.", "Run device matrix and production test push."], tests=["Foreground/background/terminated delivery on physical Android/iOS devices.", "Permission denied/later-granted, token rotation, reinstall, logout, and account-switch cases.", "Invalid tokens deactivate devices; retries do not duplicate user-visible alerts.", "Quiet hours/category preferences/mute and deep-link authorization work."], security=["Service-account/APNs secrets stay in managed secrets and client-owned consoles.", "Notification bodies reveal no sensitive dating/KYC content on lock screens unless explicitly approved.", "Backend verifies user/device ownership and rate limits registration."], source_label="FCM for Flutter", source="https://firebase.google.com/docs/cloud-messaging/flutter/get-started"),

        dict(num=11, name="OneSignal (alternative)", purpose="Managed push/in-app messaging, segmentation, campaigns, and marketer-facing engagement workflows.", priority="P1", disposition="Choose instead of direct FCM, not in addition", current="Absent; current backend is designed for direct FCM.", dependency="Very High - vendor selection/contract plus Firebase/APNs accounts; 1-3 weeks.", effort="L", layer="Flutter + backend + OneSignal", recommendation="Stay with direct FCM unless the client needs non-developer campaign creation, journeys, in-app messages, or advanced segmentation enough to justify migration and subscription cost.", sdk="onesignal_flutter, OneSignal REST API, FCM/APNs credentials.", frontend="Initialize OneSignal, request permission, set AMORAA user ID as External ID after login, clear identity at logout, handle click/deep links, tags, and consent.", backend="Create OneSignal provider adapter and API client; prevent duplicate delivery from current FCM adapter; map NotificationDeliveries to OneSignal message IDs.", database="Reuse UserDevices/NotificationDeliveries; add provider subscription ID/external ID and campaign metadata if selected.", env="ONESIGNAL_APP_ID public; ONESIGNAL_REST_API_KEY server-only; FCM service account and APNs .p8 configured in OneSignal.", webhooks="Delivery/click events and identity subscription callbacks as selected; verify authenticity and deduplicate.", team=["Prepare FCM-vs-OneSignal decision memo and migration plan.", "Implement SDK identity, token/subscription lifecycle, deep links, and backend adapter.", "Configure segments/templates/journeys with governance and no duplicate FCM sends.", "Map delivery outcomes to the existing ledger and preferences."], client=["Approve OneSignal selection, commercial plan, DPA, organization owner, and billing.", "Provide OneSignal admin access plus Firebase/APNs credentials through secure transfer.", "Approve who can send campaigns, audience rules, copy templates, frequency caps, and in-app messages.", "Nominate marketing operations and support owners."], steps=["Approve vendor and disable direct-send migration conflict.", "Configure apps/platform credentials.", "Integrate Flutter identity/consent and backend adapter.", "Migrate templates/segments and test.", "Cut over with duplicate-delivery monitoring."], tests=["OneSignal test subscriptions for Android/iOS and authenticated user mapping.", "Logout/account switch cannot receive another user's messages.", "Direct FCM path is disabled during OneSignal delivery.", "Click/deep-link, opt-out, frequency cap, and delivery callback tests pass."], security=["REST API key never ships in Flutter.", "External IDs are non-secret internal IDs; tags contain no sensitive traits.", "Least-privilege senders, MFA, approvals, and campaign audit logs."], source_label="OneSignal Flutter setup", source="https://documentation.onesignal.com/docs/en/flutter-sdk-setup"),

        dict(num=12, name="Twilio Verify v2", purpose="Provider-managed phone/email OTP verification lifecycle, anti-fraud controls, and delivery across approved channels.", priority="P0", disposition="Refactor existing Twilio Messaging OTP to Verify v2", current="twilio package and SMS adapter exist, but the backend generates/hashes OTPs locally and calls Programmable Messaging; no Verify Service SID.", dependency="Very High - Twilio account, identity/business verification, India messaging compliance/funding; 1-4 weeks.", effort="M/L", layer="Backend + existing Flutter OTP UI", recommendation="Use Verify v2 Verifications and VerificationCheck for phone verification. Retain local OtpTokens only for non-Verify email recovery/fallback if approved; do not maintain two authoritative phone-OTP lifecycles.", sdk="Twilio Node helper library; Verify v2 Services/Verifications/VerificationCheck APIs.", frontend="Keep existing account verification UI but normalize E.164 numbers, surface provider-neutral resend/attempt errors, and add cooldown/accessibility behavior.", backend="Replace createOtp/deliverOtp/verifyOtp for phone with Twilio Verify calls, map provider statuses/errors, rate limit by IP/phone/user, and use restricted API keys/subaccounts by environment.", database="Add VerificationAttempts/provider fields or repurpose OtpTokens for audit only (providerSid, channel, status, attempts, expiresAt); never store the provider OTP. Keep email reset tokens separate.", env="TWILIO_ACCOUNT_SID, TWILIO_API_KEY, TWILIO_API_SECRET (preferred), TWILIO_VERIFY_SERVICE_SID; optional messaging service/template configuration.", webhooks="Optional Verify status callbacks/fraud events; validate Twilio signatures and deduplicate.", team=["Refactor phone OTP start/check endpoints to Verify v2 while preserving API response contracts.", "Implement rate limits, generic responses, provider error mapping, retry/backoff, and environment isolation.", "Add audit/metrics without storing codes.", "Prepare cutover/fallback/runbook and update tests/Postman collection."], client=["Create a company-owned Twilio organization/subaccount, complete identity/business verification, fund it, and invite least-privilege team members.", "Create Verify Services for non-production/production and provide SIDs/credentials through a secret manager.", "Provide legal entity, registered address, tax/billing, authorized signatory, support contact, and India DLT/sender/template details if required.", "Approve OTP brand text, channels, country allow-list, language, fraud thresholds, resend/attempt policy, and monthly budget alerts."], steps=["Complete Twilio/compliance onboarding.", "Create Verify services and restricted credentials.", "Implement start/check adapter and database audit.", "Run test-number and real-device cases.", "Gradually cut over and monitor conversion/fraud/cost."], tests=["Approved, pending, expired, wrong, max-attempt, resend-cooldown, blocked-country, and provider-outage cases.", "E.164 normalization and India test devices/networks.", "Rate limiting and generic responses prevent enumeration/OTP abuse.", "No OTP is logged or stored in production."], security=["Prefer API keys over the primary Auth Token and rotate credentials.", "Do not log codes, full phone numbers, or provider request bodies.", "Use Twilio request signature verification for callbacks and separate subaccounts by environment."], source_label="Twilio Verify v2", source="https://www.twilio.com/docs/verify"),

        dict(num=13, name="Google OAuth / Sign in with Google", purpose="Google social sign-in with backend ID-token verification and AMORAA JWT session issuance.", priority="P0", disposition="Complete existing partial integration", current="Flutter google_sign_in and backend google-auth-library verification exist; GOOGLE_CLIENT_IDS is not configured and platform OAuth files/identifiers are incomplete.", dependency="High - Google Cloud OAuth ownership, app IDs/signing, consent screen; 2-10 client days.", effort="M", layer="Flutter/web + backend", recommendation="Keep the existing backend ID-token verification pattern, but finalize identifiers, current SDK configuration, account linking policy, nonce/state, and web Google Identity Services behavior before enabling production.", sdk="google_sign_in/latest compatible Flutter plugin; Google Identity Services; google-auth-library on Express.", frontend="Configure Android/iOS/web client IDs, platform URL schemes, consent/cancel/error UX, and send ID token over HTTPS to /api/auth/google.", backend="Verify issuer/audience/expiry/nonce where applicable, validate email verification, handle provider collisions/linking, log security events, and issue existing access/refresh tokens.", database="Existing Users.googleId/authProvider works; add ExternalIdentities (provider, providerSubject unique, userId, emailAtLink, linkedAt, revokedAt) for multi-provider/Apple-ready linking.", env="GOOGLE_CLIENT_IDS or per-platform GOOGLE_*_CLIENT_ID; web client secret only if an authorization-code flow is chosen; allowed origins/redirect URIs in Google Cloud.", webhooks="No routine sign-in webhook. Optional Cross Account Protection can be evaluated later.", team=["Update Google sign-in configuration for Android/iOS/web and preserve backend token verification.", "Implement safe account linking/collision policy and ExternalIdentities migration.", "Add telemetry/error mapping, logout/revocation behavior, and provider test coverage.", "Document production OAuth consent-screen/credential ownership."], client=["Provide a corporate Google Cloud/Firebase project and OAuth administrator access.", "Finalize Android package, iOS bundle, web domain, Android release SHA-1/SHA-256, and store identifiers.", "Configure/approve OAuth consent screen, app name/logo, support email, privacy/terms URLs, authorized domains/origins/redirects.", "Approve whether matching-email local accounts may be linked, blocked, or require re-authentication."], steps=["Finalize identifiers/signing.", "Create platform OAuth clients and consent screen.", "Configure Flutter/web and backend audiences.", "Implement/linking migration and errors.", "Test internal then production OAuth verification."], tests=["New user, existing Google user, local-email collision, deleted/deactivated account, cancellation, expired/wrong-audience token.", "Android release signing, iOS URL return, and web origin tests.", "Account switch/logout clears provider/session state.", "Backend rejects client-supplied subject/email without a valid ID token."], security=["ID tokens travel only over HTTPS and are verified server-side; never trust profile fields alone.", "Client secrets stay server-side; OAuth console access uses MFA/least privilege.", "Require explicit secure linking rules to prevent account takeover."], source_label="Google backend authentication", source="https://developers.google.com/identity/sign-in/ios/backend-auth"),

        dict(num=14, name="Sign in with Apple", purpose="Apple social sign-in for iOS/web with privacy-preserving email relay and backend session issuance.", priority="P0", disposition="Implement before iOS social-login release", current="Absent; no entitlement, Service ID, key, Apple callback, nonce, or ExternalIdentities table.", dependency="Very High - Apple Developer Account Holder/Admin and verified domain; 3-15 client days.", effort="M/L", layer="iOS Flutter + backend (+ web if offered)", recommendation="Implement alongside Google login before iOS release when social login is offered. Persist the Apple subject and first-return profile immediately; users may hide their email and Apple usually returns name/email only on first authorization.", sdk="sign_in_with_apple Flutter package/native AuthenticationServices; Apple OpenID Connect token endpoints/JWKS.", frontend="Add Apple button per platform guidelines, random nonce/state, iOS capability/entitlement, cancellation/error UX, and send identity token/authorization code to backend.", backend="Verify Apple JWS issuer/audience/nonce, exchange/revoke codes where required, support private relay email, link through ExternalIdentities, and issue AMORAA tokens.", database="ExternalIdentities plus Users email/relay flags; store Apple subject and authorization state, not raw identity tokens.", env="APPLE_TEAM_ID, APPLE_BUNDLE_ID/CLIENT_ID, APPLE_SERVICE_ID, APPLE_KEY_ID, APPLE_PRIVATE_KEY secret, APPLE_REDIRECT_URI.", webhooks="Apple server-to-server account events/revocation endpoint if configured; verify signed tokens and update link status.", team=["Add iOS capability/entitlement and Flutter flow with nonce/state.", "Implement Apple token verification, code exchange/revocation, relay-email handling, and account linking.", "Add ExternalIdentities migration and error/audit events.", "Prepare App Review test instructions and deletion/revocation behavior."], client=["Provide Apple Developer organization Account Holder/Admin access and App Store Connect access.", "Register final App ID/bundle, enable Sign in with Apple, create Service ID if web use is needed, and provide verified domain/return URL.", "Create .p8 key/Key ID/Team ID and transfer via an approved secret channel; approve rotation owner.", "Approve private relay support, account-linking policy, privacy/terms/support URLs, and App Review account/instructions."], steps=["Finalize Apple organization/app identifiers.", "Create capability, Service ID, key, domains/redirects.", "Implement Flutter and backend verification/linking.", "Test sandbox/TestFlight and revocation.", "Submit with App Review notes."], tests=["First login, repeat login without returned name, Hide My Email, cancellation, nonce mismatch, revoked credential, account collision.", "Physical-device/TestFlight return flow and relay email delivery.", "Deletion revokes provider tokens where required and invalidates AMORAA sessions.", "Wrong audience/issuer/expired JWS is rejected."], security=["Private key stays in secrets manager and never enters the app/repository.", "Use nonce/state and server verification; do not trust client profile claims.", "Minimize relay-email exposure and document account-linking/recovery paths."], source_label="Sign in with Apple", source="https://developer.apple.com/sign-in-with-apple/"),

        dict(num=15, name="Google Play Billing", purpose="Android digital subscription purchase, verification, acknowledgement, renewal/cancel lifecycle, and entitlement sync.", priority="P0", disposition="Implement; replace mobile Razorpay for digital subscriptions", current="Absent. Current payment UI uses razorpay_flutter and MySQL subscriptions/payments; no Play products, purchase token fields, Developer API, or RTDN.", dependency="Very High - verified Play Console organization, merchant/payments profile, app signing/release track, products; 1-6 weeks.", effort="XL", layer="Android Flutter + backend + MySQL", recommendation="Use Google Play Billing for Android in-app digital membership. Keep Razorpay only for eligible web/physical/off-app use after policy/legal review. MySQL remains the unified entitlement source after server verification.", sdk="Flutter in_app_purchase (or approved maintained billing wrapper), Google Play Billing Library, Google Play Developer API, Cloud Pub/Sub RTDN.", frontend="Query localized ProductDetails, launch purchase, handle pending/cancelled/completed results, send purchase token/product ID to backend, restore/query purchases, and never grant premium locally.", backend="Verify token with Developer API, bind obfuscated account ID, acknowledge initial purchase, process RTDN via Pub/Sub push/pull, handle renewal, grace, hold, pause, cancel, expire, refund/void, and reconcile.", database="Extend SubscriptionPlans with playProductId/basePlanId/offerId; Payments/Subscriptions with purchaseTokenHash, orderId, linkedPurchaseTokenHash, acknowledgementState, region, provider timestamps; add StoreEvents unique messageId.", env="GOOGLE_PLAY_PACKAGE_NAME, service account/workload identity credentials, GCP project/topic/subscription, RTDN auth audience, product mapping; credentials server-only.", webhooks="Cloud Pub/Sub RTDN endpoint authenticated with OIDC. Each RTDN triggers a Developer API lookup; messages are idempotent by messageId.", team=["Refactor membership checkout into platform-specific purchase adapters and one backend entitlement service.", "Implement Play verification/acknowledgement, RTDN ingestion, lifecycle mapping, reconciliation, and migration.", "Create products/base plans/offers and safe mapping to existing SubscriptionPlans.", "Add Play license-test, closed-track, refund/cancel, finance/support runbooks."], client=["Provide verified Google Play Console organization owner/admin access, final package ID, app signing configuration, and a closed test release.", "Create/approve merchant payments profile, tax/bank/settlement details, subscriptions/base plans/offers, prices/countries, grace/hold policy, and tester accounts.", "Provide GCP project/Pub/Sub access and authorize service account/Developer API linkage.", "Provide Terms, Privacy, subscription disclosure, refund/cancellation/support policies and pricing approval."], steps=["Finalize package/signing/store organization and product catalog.", "Implement Flutter purchase adapter and backend verification schema.", "Configure RTDN/Pub/Sub and lifecycle reconciliation.", "Run license/closed-track testing including pending/refund/cancel.", "Migrate checkout and enable production gradually."], tests=["Successful, pending, cancelled, duplicate, already-owned, restore, upgrade/downgrade, renewal, grace, hold, expire, refund, and revoked cases.", "RTDN duplicate/out-of-order/authentication tests plus Developer API reconciliation.", "Entitlement is granted only after verified PURCHASED state and acknowledgement.", "Cross-account purchase-token replay is rejected."], security=["Never trust client purchase result alone or store raw service credentials in app.", "Hash/encrypt purchase tokens and enforce unique ownership/idempotency.", "Authenticate RTDN, verify product/package/amount/state, and maintain immutable event audit."], source_label="Google Play backend integration", source="https://developer.android.com/google/play/billing/backend"),

        dict(num=16, name="Apple StoreKit / App Store Server API", purpose="iOS digital subscriptions, signed transaction verification, renewals/cancellations/refunds, and unified entitlements.", priority="P0", disposition="Implement", current="Absent; no IAP capability/products, transaction IDs, App Store keys, or server notification endpoint.", dependency="Very High - Apple Account Holder, paid agreements, banking/tax, products/review; 1-6 weeks.", effort="XL", layer="iOS Flutter + backend + MySQL", recommendation="Use StoreKit/App Store purchase flow for iOS membership and App Store Server Notifications V2. MySQL grants access only after verified signed transaction/server status.", sdk="Flutter in_app_purchase or approved StoreKit 2 bridge; App Store Server API library; signed JWS transaction verification; Notifications V2.", frontend="Load localized products, purchase/restore, finish transactions after backend confirmation, show pending/cancelled states, and open subscription management.", backend="Verify signed transaction/JWS and bundle/environment/product, query transaction/subscription history/status, process V2 notifications, map renewal/grace/billing retry/expire/refund/revoke, and reconcile.", database="SubscriptionPlans appleProductId/groupId; Payments/Subscriptions originalTransactionId, transactionId, webOrderLineItemId, signedDate, environment, ownershipType; StoreEvents unique notificationUUID.", env="APPLE_ISSUER_ID, APPLE_KEY_ID, APPLE_IAP_PRIVATE_KEY, APPLE_BUNDLE_ID, APPLE_APP_ID, environment URLs, product mapping.", webhooks="HTTPS App Store Server Notifications V2 for sandbox/production; verify signedPayload/JWS and notificationUUID idempotency; support test-notification API.", team=["Implement platform purchase adapter and unified entitlement mapping.", "Build JWS verification, Server API client, V2 notification ingestion, reconciliation, and support lookup.", "Configure products/subscription group/offers and map to MySQL plans.", "Add StoreKit local, sandbox, TestFlight, refund/revoke, and restore runbooks."], client=["Provide Apple Developer/App Store Connect Account Holder/Admin access and final bundle/App ID.", "Accept paid-app agreements and provide legal entity, banking, tax, pricing/country, subscription group/products/offers, localization, and review metadata.", "Create App Store Server API key (.p8), Issuer ID, Key ID and share securely; configure notification URLs.", "Provide subscription disclosures, privacy/terms, refund/support contacts, and App Review credentials/instructions."], steps=["Complete agreements/store product setup.", "Implement Flutter StoreKit purchase/restore and backend schema.", "Implement signed transaction verification and Server API.", "Configure/test Notifications V2 sandbox/TestFlight.", "Enable production and daily reconciliation."], tests=["StoreKit local, sandbox, TestFlight purchase/restore/cancel/renew/billing retry/grace/refund/revoke/upgrade/downgrade.", "V2 TEST notification, duplicates, wrong bundle/environment, invalid JWS, and out-of-order events.", "Entitlement follows verified original transaction lifecycle across devices.", "Account switch cannot claim another user's transaction."], security=[".p8 key is server-only with rotation/owner record.", "Verify Apple certificate/JWS chain, bundle, environment, product, and dates.", "Store minimal transaction identifiers and immutable event hashes, not receipts in logs."], source_label="App Store Server API", source="https://developer.apple.com/documentation/appstoreserverapi"),

        dict(num=17, name="Razorpay (Web payments/subscriptions)", purpose="Eligible web checkout and INR payment/subscription processing outside mandatory mobile-store billing scope.", priority="P1", disposition="Harden and reposition existing integration", current="Flutter mobile checkout, backend order/signature/payment fetch, webhook HMAC, Payments/PaymentEvents/Subscriptions already exist; credentials absent.", dependency="Very High - Razorpay KYC/activation, bank/tax/domain/policies; 1-4 weeks.", effort="M/L", layer="Flutter Web/web checkout + backend", recommendation="Retain the strong existing backend verification/idempotency path, but limit production Razorpay use to policy-eligible web/physical/off-app cases. Do not use it to bypass Play/App Store billing for in-app digital membership.", sdk="Razorpay Orders/Payments/Subscriptions APIs and Web Checkout; existing razorpay_flutter only where platform policy permits.", frontend="Use web checkout/hosted flow with server-created order and existing success/failure/pending UI; never embed key secret. Add return/recovery polling for interrupted checkout.", backend="Keep order creation, HMAC checkout verification, provider payment fetch, raw-body webhook validation, event dedupe; add refunds/subscription lifecycle and environment/merchant mapping as selected.", database="Existing Payments, PaymentEvents, Subscriptions are suitable; add providerSubscriptionId, invoiceId, refund/dispute details, checkoutPlatform, tax/receipt references and reconciliation status.", env="Existing RAZORPAY_KEY_ID, KEY_SECRET, WEBHOOK_SECRET, API_BASE_URL; separate test/live secrets and webhook URLs.", webhooks="payment.captured/failed, order.paid, refund, dispute/chargeback, subscription/invoice events if Razorpay Subscriptions is used; raw body + signature + event ID dedupe.", team=["Complete web checkout adapter and remove/disable non-compliant mobile digital checkout paths.", "Expand webhook lifecycle/refund/reconciliation and preserve existing idempotency/verification.", "Add policy-based product/channel gating, finance exports, receipts, and support tools.", "Update integration tests using test mode and production go-live checklist."], client=["Create/activate company Razorpay account and complete KYC with legal entity, PAN/GST/CIN as applicable, registered address, website/domain, business proof, bank account/cancelled cheque, authorized signatory, and settlement details.", "Provide owner/admin/developer access, test/live keys via secret manager, webhook secret, and approved webhook domain.", "Approve eligible products/platforms, prices/taxes/refunds/cancellations, receipt/invoice text, and settlement/reconciliation owner.", "Provide Terms, Privacy, Refund/Cancellation Policy, Contact/Support page, and billing contact."], steps=["Complete activation/KYC and policy scope.", "Configure test keys/webhook and web checkout.", "Harden lifecycle/refund/reconciliation and channel gating.", "Run test-mode payment/refund/dispute cases.", "Rotate to live keys and monitor settlements."], tests=["Success, failure, cancel, pending/external wallet, duplicate order/idempotency, invalid signature, amount mismatch, refund/dispute, and webhook replay/out-of-order.", "Test-mode dashboard/payment records reconcile to MySQL.", "Live smoke payment/refund uses client-approved minimal amount and finance sign-off.", "Mobile digital products cannot select Razorpay when store billing is required."], security=["Secrets are backend-only; checkout key ID is the only public credential.", "Use raw-body HMAC and unique event IDs as the current code already does.", "Avoid storing card/payment instrument data; rely on Razorpay-hosted checkout and audit access."], source_label="Razorpay webhook validation", source="https://razorpay.com/docs/webhooks/validate-test/"),

        dict(num=18, name="Google Maps Platform", purpose="Maps, location selection, place autocomplete/details, distance, and venue/date-spot discovery.", priority="P0", disposition="Implement", current="Android location permissions exist; DateSpotsMapScreen explicitly says map unavailable and uses static venue data. No Maps SDK/package/API keys/place schema.", dependency="Very High - Google Cloud billing, final app IDs/signing/domain, budget; 2-10 client days.", effort="L", layer="Flutter/web + backend + MySQL", recommendation="Use platform-restricted keys for map rendering and a backend-restricted key/proxy for Places/Routes data where business logic or secret controls are needed. Store Google place_id plus selected business fields, not copied provider data indefinitely beyond terms.", sdk="google_maps_flutter, Maps JavaScript API for web if supported, Places API (New), Geocoding/Routes/Distance Matrix equivalent as approved.", frontend="Replace map preview with live map, permission rationale, approximate-location fallback, autocomplete session tokens, selected place details, markers, and graceful quota/network fallback.", backend="Provide venue/place search proxy if needed, restrict fields, cache only permitted data, compute/store internal distance/location indexes, and enforce quotas/rate limits.", database="Add VenueLocations/Places (googlePlaceId unique, name/address, lat/lng POINT or decimals, categories, source, lastRefreshedAt) and Event placeId/lat/lng references; add spatial indexes if query volume warrants.", env="MAPS_ANDROID_API_KEY, MAPS_IOS_API_KEY, MAPS_WEB_API_KEY public but restricted; MAPS_SERVER_API_KEY server-only; GOOGLE_MAP_ID; API allow-lists and budget alerts.", webhooks="None for standard Maps/Places. Scheduled refresh may be used within provider terms.", team=["Enable only required APIs and create per-platform restricted keys.", "Implement map/autocomplete/place details with session tokens and fallback UX.", "Add venue/location schema, geospatial query strategy, rate limits, caching, and cost telemetry.", "Test permission, key restrictions, quota, billing, and accessibility."], client=["Provide client-owned Google Cloud project, billing profile/payment method, Maps admin/billing access, and monthly budget/alert thresholds.", "Provide final package/bundle/domain, Android release signing fingerprints, iOS bundle restrictions, and production web referrers.", "Approve location consent wording, approximate/precise use, retention, venue curation, countries/cities, and cost ceiling.", "Provide production/staging domains and legal/privacy policy coverage for location data."], steps=["Finalize identifiers/billing and enable minimal APIs.", "Create restricted keys and map IDs/styles.", "Implement Flutter UI and backend/schema.", "Run quota/cost/security tests.", "Launch city-by-city with monitoring."], tests=["Physical-device permissions: denied, approximate, precise, revoked, location off.", "Autocomplete session, place details, map tiles, route/distance, invalid/expired key, and quota failures.", "API key restriction tests from unauthorized app/domain/server.", "Venue coordinates/distance and privacy deletion behavior are verified."], security=["Restrict every key by API and application; server key never ships in app.", "Do not expose exact user location to other users or marketing tools.", "Use minimum precision/retention and explicit location-purpose consent."], source_label="Places Autocomplete (New)", source="https://developers.google.com/maps/documentation/places/web-service/place-autocomplete"),

        dict(num=19, name="Cloudinary", purpose="Profile/event/chat media storage, resizing, optimization, delivery, and optional moderation workflow.", priority="P0", disposition="Implement for public/non-regulated media", current="Profile photos are written under public /uploads; chat/KYC media use local private disk. No durable object storage/CDN/media asset model.", dependency="High - vendor plan/billing, media policy, domain; 2-10 client days.", effort="L", layer="Flutter/web + backend + MySQL", recommendation="Use signed backend-mediated uploads for profile/event/chat media and store asset IDs, not only URLs. Do not place raw Aadhaar/KYC images in public delivery; use private/authenticated assets or the separately approved KYC/S3 path.", sdk="Cloudinary Node SDK/Upload API; cloudinary_flutter for delivery; signed upload parameters; webhooks/moderation add-ons if selected.", frontend="Request upload signature/session or upload through Express, show compression/progress/retry, then render approved transformations with fallback and placeholder.", backend="Validate magic bytes/size/ownership, sign/upload, set folders/tags/context/access type, receive moderation callbacks, delete/replace old assets, and issue signed URLs for private media.", database="Create MediaAssets (id, userId, provider, publicId, resourceType, version, secureUrl, accessType, purpose, bytes, dimensions, moderationStatus, checksum, deletedAt); migrate OnboardingProfiles.photos and MessageMedia references.", env="CLOUDINARY_CLOUD_NAME and API_KEY; CLOUDINARY_API_SECRET server-only; upload preset/transformation/moderation webhook secret; CDN/custom-domain settings.", webhooks="Upload/moderation/analysis callbacks if enabled; verify signature/timestamp and map by public_id/version.", team=["Design MediaAssets schema and purpose-based access policy.", "Implement signed uploads, transformations, private delivery, deletion, migration, and provider abstraction.", "Integrate moderation state without exposing unapproved media.", "Add lifecycle cleanup, backup/export, cost/quota alerts, and media QA."], client=["Create a company Cloudinary product environment/organization, choose plan/region, add billing, and invite roles.", "Provide cloud name/account access and approve signed-upload/private delivery architecture.", "Approve media limits, formats, transformations, retention/deletion, moderation provider/thresholds, and acceptable-use policy.", "Provide custom media domain/DNS access if desired and legal/privacy consent for media processing."], steps=["Approve media classification/provider plan.", "Create MediaAssets schema and provider adapter.", "Implement signed upload/delivery/moderation/deletion.", "Migrate local development paths in staging.", "Load/security/cost test then cut over."], tests=["Valid/invalid MIME and magic bytes, oversize, multi-upload, retry, duplicate/checksum, replace/delete, authorization, signed URL expiry.", "Transformation quality/orientation/format and slow-network behavior.", "Moderation pending/approved/rejected callback and replay cases.", "No KYC/private chat asset is accidentally publicly addressable."], security=["API secret is backend-only; prefer signed uploads and restricted presets.", "Use authenticated/private delivery for sensitive media and short-lived URLs.", "Strip metadata as approved, malware/content-scan uploads, and audit deletions/access."], source_label="Cloudinary Upload API", source="https://cloudinary.com/documentation/image_upload_api_reference"),

        dict(num=20, name="AWS S3 + CloudFront (alternative)", purpose="Infrastructure-controlled object storage/CDN with private buckets, presigned uploads, lifecycle policies, and signed delivery.", priority="P1", disposition="Alternative/complement for regulated private media", current="Absent; local disk is not horizontally durable and no cloud IAM/bucket/CDN exists.", dependency="Very High - AWS organization/billing/IAM/DNS/security ownership; 1-3 weeks.", effort="XL", layer="Backend/infrastructure + Flutter upload client", recommendation="Choose S3/CloudFront instead of Cloudinary for the full media stack, or use it narrowly for KYC/private evidence when infrastructure control is required. Do not run duplicate public-media systems without an explicit storage-classification decision.", sdk="AWS SDK for JavaScript v3 (S3, CloudFront); presigned URL flow; CloudFront Origin Access Control and signed URLs/cookies.", frontend="Upload via short-lived presigned PUT/multipart instructions with checksum, progress, retry, and completion callback; consume short-lived signed delivery URLs.", backend="Authorize object key/purpose, generate presigned upload/download, verify size/checksum/content type after upload, trigger scanning/processing, delete/lifecycle, and sign CloudFront URLs.", database="MediaAssets with bucket/key/version/checksum/access class; UploadSessions; optional scan/moderation state. Never make S3 object keys the sole business record.", env="AWS_REGION, S3_* bucket names, CLOUDFRONT_DISTRIBUTION/KEY_PAIR, KMS key IDs; prefer workload roles over long-lived AWS keys.", webhooks="S3 Event Notifications to SQS/Lambda/HTTPS worker for scan/process; idempotent by event/object version.", team=["Produce Cloudinary-vs-AWS storage decision and data classification.", "Provision private buckets, OAC CloudFront, IAM roles/KMS/lifecycle/logging/backup.", "Implement upload sessions, checksums, scan pipeline, signed delivery, and MediaAssets mapping.", "Add infrastructure-as-code, disaster recovery, cost alerts, and runbook."], client=["Provide client-owned AWS Organization/account, billing/payment, preferred region, security/IAM administrator, and budget.", "Approve data residency, encryption/KMS ownership, retention/legal hold, backup/DR, CDN custom domain, and media classification.", "Provide DNS/certificate access and security/compliance requirements.", "Nominate operations owner for AWS cost, incidents, keys, and access reviews."], steps=["Approve architecture/classification/region.", "Provision infrastructure as code and least-privilege roles.", "Implement MediaAssets/upload/scan/signed delivery.", "Migrate/test staging data.", "Security/load/DR test then production cutover."], tests=["Unauthorized bucket/object access fails; presigned URL method/key/expiry/checksum limits are enforced.", "Multipart retry, corrupted checksum, oversized file, malware/quarantine, delete/lifecycle, and versioning.", "CloudFront OAC blocks direct public S3 access; signed URLs expire.", "Backup/restore, cross-zone failure, and cost alarms are exercised."], security=["Private-by-default buckets, Block Public Access, OAC, KMS, least-privilege roles, CloudTrail/access logging.", "No long-lived AWS access keys in Flutter or repository.", "Separate KYC/private/public buckets and lifecycle policies."], source_label="S3 presigned URLs", source="https://docs.aws.amazon.com/AmazonS3/latest/userguide/using-presigned-url.html"),

        dict(num=21, name="Stream Chat OR Sendbird", purpose="Managed real-time chat, delivery/read state, presence, moderation hooks, and scalable messaging operations.", priority="P0", disposition="Architectural decision - replace existing chat or retain it", current="AMORAA already has authenticated Socket.IO chat, MySQL conversations/messages/media/read state/presence/mute and tests. No managed chat SDK.", dependency="Very High - vendor selection/contract/data residency/migration; 2-6 weeks.", effort="XL", layer="Flutter + backend + managed chat", recommendation="Do not layer a managed provider on top of current Socket.IO. Choose either (a) retain/harden the working in-house chat, or (b) migrate to one vendor with a defined cutover, history, moderation, and source-of-truth plan. Stream is a natural Flutter candidate; Sendbird remains a viable alternative.", sdk="stream_chat_flutter + server SDK/webhooks, or sendbird_chat_sdk + Platform API/webhooks.", frontend="Replace ChatRepository with provider adapter, connect using short-lived backend-generated token, map users/channels/messages/read/presence/typing/media, and preserve current UI/access rules.", backend="Generate expiring provider tokens, synchronize user/channel lifecycle after active match/block/delete, verify webhooks, enforce bans/moderation, and maintain only required shadow metadata.", database="Decide system of record. Add ChatProviderMappings and ProviderEvents; either migrate/archive existing Messages or keep an immutable compliance/audit projection without dual-writes indefinitely.", env="CHAT_PROVIDER, STREAM_API_KEY/SECRET or SENDBIRD_APP_ID/API_TOKEN, webhook secret, region, moderation config; secret server-only.", webhooks="Message/channel/user/moderation delivery events; verify signatures, deduplicate, and handle retries/outages.", team=["Prepare build-vs-buy decision including current chat evidence, scale/SLA, moderation, data residency, pricing, and migration cost.", "If selected, build provider abstraction, backend token endpoint, lifecycle/webhook handlers, and migration tooling.", "Preserve block/match/account deletion authorization and media policy.", "Run dual-read/cutover strategy without long-term duplicate sending."], client=["Approve retain-in-house versus Stream/Sendbird and, if vendor, select commercial plan/region/SLA/DPA with billing and owner access.", "Approve chat retention, moderation rules, escalation, data export/deletion, support access, and prohibited-content policy.", "Provide expected MAU/concurrency/message/media volumes and incident/support requirements.", "Nominate trust-and-safety, legal/privacy, support, and procurement owners."], steps=["Decision workshop and vendor proof of concept.", "Approve data model/migration/source of truth.", "Implement token/lifecycle/webhook/provider repository.", "Migrate pilot cohort and validate moderation/history.", "Cut over with rollback and reconciliation."], tests=["Only active matched users can connect/create/read; block/unmatch/delete revokes access quickly.", "Offline/reconnect/order/read/presence/typing/media/push/webhook duplicate and provider outage cases.", "History migration checksums/counts and user isolation.", "Load, rate-limit, moderation, deletion/export, and rollback tests."], security=["Provider secret/API token stays backend-only; use short-lived user tokens.", "Verify webhook signatures and re-authorize lifecycle events against AMORAA state.", "Encrypt/minimize chat data, restrict support access, and define retention/legal hold."], source_label="Stream Flutter go-live checklist", source="https://getstream.io/chat/docs/sdk/flutter/guides/go-live-checklist/"),

        dict(num=22, name="Identity / KYC provider", purpose="ID document and selfie/liveness verification with auditable provider results and minimal regulated-data exposure.", priority="P0", disposition="Provider/legal decision required before production KYC", current="AMORAA accepts Aadhaar+selfie into private local storage and records pending status; no provider, liveness/OCR, admin review, or external decision callback.", dependency="Critical - legal review, vendor contract/KYB, approved Aadhaar method and data flow; 2-8+ weeks.", effort="XL", layer="Flutter + backend + MySQL + KYC provider", recommendation="Pause production use of raw Aadhaar-image capture until the client selects a legally approved provider/method and counsel confirms purpose, consent, storage, masking, retention, and UIDAI/DPDP compliance. Prefer provider-hosted capture/tokenized results; never use Aadhaar number as AMORAA's user identifier.", sdk="Selected India-capable KYC provider SDK/hosted flow/API for document OCR, face match/liveness and permitted Aadhaar/offline/DigiLocker flow; signed webhooks.", frontend="Replace direct upload with provider session/SDK or hosted flow, explicit consent/purpose notice, camera/liveness UX, retry/manual fallback, and status screen.", backend="Create provider session, bind user/reference, verify webhook signature/status, map normalized result to pending/verified/rejected/manual_review, fetch only necessary fields, and enforce review/admin authorization.", database="Extend IdentityVerifications with provider, providerReference unique, method, consentVersion/time, maskedDocumentLast4 or token only, liveness/face result summary, decision codes, expiresAt; remove/expire raw document paths per approved policy and add immutable audit events.", env="KYC_PROVIDER, client/account IDs, API keys/certificates, webhook secret, encryption key, region/base URLs; no KYC secrets or raw responses in Flutter/logs.", webhooks="Signed verification.completed/failed/manual-review callbacks with timestamp/replay protection/idempotency; poll only as fallback.", team=["Run vendor/legal architecture review and data-protection impact assessment before coding production flow.", "Implement provider session, Flutter capture/hosted flow, signed webhook, status mapping, and minimal schema.", "Implement admin/manual-review boundaries, retention/deletion, redaction, access logging, and incident response.", "Migrate or purge current raw development submissions according to approved policy."], client=["Select/contract the provider and complete its KYB: legal entity, CIN/GST/PAN as applicable, registered address, authorized signatory, directors/beneficial ownership, bank/billing, website/app, use-case volumes, security/privacy contacts, and any provider/UIDAI eligibility documents.", "Obtain written legal advice approving the exact Aadhaar/offline/DigiLocker/document flow, purpose, consent, retention, cross-border processing, age policy, and grievance/DPO contact.", "Provide client-owned provider account/admin access, sandbox/live credentials/certificates via secret manager, webhook allow-list, and production approval.", "Approve accepted documents, retry/manual-review/rejection policy, retention/deletion timetable, support copy, and escalation SLA."], steps=["Legal/DPIA and provider selection/KYB.", "Approve data-flow diagram and retention/minimization.", "Implement sandbox session/capture/webhook/status.", "Security/privacy/manual-review penetration and operational testing.", "Limited pilot then production approval; purge legacy raw files as directed."], tests=["Happy path, document mismatch, liveness fail, poor image, duplicate session, timeout, user cancellation, provider outage, manual review, rejection/resubmission.", "Forged/replayed/out-of-order webhook and cross-user reference tests.", "Access logs, encryption, retention expiry, deletion, data export/redaction, and incident drill.", "No raw Aadhaar number/biometric/OTP appears in logs, analytics, support, or marketing systems."], security=["UIDAI states Aadhaar biometric/OTP authentication data must not be stored permanently and Aadhaar must not be a domain identifier; counsel/provider must confirm the selected flow.", "Prefer tokenized/masked results and provider-hosted capture; segregate KYC data and keys from general media.", "Strict role-based access, encryption, audit, redaction, retention, and breach-response controls."], source_label="UIDAI requesting-entity security requirements", source="https://uidai.gov.in/index.php?id=1008&itemid=2420&lang=en&option=com_content&view=article"),

        dict(num=23, name="Sentry (optional)", purpose="Cross-stack exception/performance monitoring, release correlation, traces, and backend operational diagnostics.", priority="P1", disposition="Optional; recommended primarily for Express backend", current="Absent. Flutter has no Crashlytics/Sentry; Express logs to console only and has no release-aware APM.", dependency="Medium - organization/project/billing/data governance; 1-3 client days.", effort="M", layer="Flutter + Node/Express", recommendation="Use Firebase Crashlytics as P0 Flutter crash reporting and Sentry for Node/Express APM. Add Sentry Flutter only if the client explicitly wants one cross-stack tool and accepts overlap/cost with Crashlytics.", sdk="@sentry/node for Express; sentry_flutter optional; source-map/symbol upload CLI/integration.", frontend="If enabled, initialize before app start, scrub context, set release/environment/user ID, capture unhandled errors/traces, and upload symbols without PII.", backend="Initialize before Express routes, add tracing/error middleware in correct order, tag request/user/release, scrub bodies/headers, and exclude expected business errors.", database="No Sentry business table; optional IntegrationIncident references for operational workflows. Do not store DSN/auth token in MySQL.", env="SENTRY_DSN, SENTRY_ENVIRONMENT, SENTRY_RELEASE, traces/profile sample rates; SENTRY_AUTH_TOKEN only in CI for symbol/source-map upload.", webhooks="Optional issue alert webhook to support/incident tooling; verify secret and avoid exposing event payloads.", team=["Create backend Sentry project and privacy scrubber/allow-list.", "Instrument Express errors, performance, release/deploy, and health alerts.", "Optionally instrument Flutter if duplication is approved.", "Configure symbol/source-map upload and incident runbook."], client=["Approve Sentry use/plan/region/DPA and create client-owned organization/projects.", "Invite development/operations roles and provide DSNs; provide CI token securely if symbol upload is approved.", "Approve retention, sampling, alert recipients, on-call workflow, and data scrubbing.", "Nominate production incident owner."], steps=["Approve scope (backend only vs full stack).", "Create projects/roles and scrub rules.", "Instrument non-production and release pipeline.", "Generate controlled errors/performance traces.", "Enable production sampling/alerts."], tests=["Controlled backend exception and slow transaction appear with correct release/environment.", "Request bodies, Authorization, cookies, OTP, KYC, messages, email/phone are absent/redacted.", "Source maps/symbols resolve stack traces.", "Alert routing and quota/sampling behavior are exercised."], security=["DSN is not treated as authorization; CI auth token is secret and scoped.", "Use deny-before-send/allow-list scrubbing and low sampling for sensitive routes.", "Restrict issue access and set appropriate retention/deletion."], source_label="Sentry Flutter package", source="https://pub.dev/packages/sentry_flutter"),

        dict(num=24, name="Firebase Crashlytics", purpose="Production Flutter crash, non-fatal, and Android ANR visibility with release/symbol mapping.", priority="P0", disposition="Implement", current="Absent; main() does not initialize Firebase or install fatal/asynchronous error handlers.", dependency="High - Firebase access and production signing/release pipeline; 1-3 client days.", effort="M", layer="Flutter Android/iOS", recommendation="Make Crashlytics the canonical mobile crash tool. Do not attach sensitive user/profile/KYC/chat content to custom keys or logs.", sdk="firebase_crashlytics + firebase_core; Crashlytics Gradle/Xcode build phases and symbol upload.", frontend="Initialize Firebase, wire FlutterError.onError and PlatformDispatcher.onError, set release-safe user identifier and bounded non-sensitive keys, opt-in/collection behavior, and controlled non-fatal capture.", backend="None. Backend APM is handled separately by Sentry or another server tool.", database="None; optionally record client app version/build in support tickets, not raw crash reports.", env="Firebase app config per environment; FIREBASE_APP_ID for symbol upload; CI service identity/CLI auth.", webhooks="Optional alert integrations from Firebase/Google Cloud; no application webhook required.", team=["Install/configure Crashlytics and top-level error handlers.", "Configure build IDs, obfuscation/symbol upload, release/environment tagging, and privacy scrubbing.", "Add controlled crash/non-fatal QA and alert routing.", "Document triage, ownership, severity, and release regression workflow."], client=["Provide Firebase project owner/admin access and production app registrations.", "Provide release signing/store/CI access needed to create builds and upload symbols.", "Approve crash collection/opt-out, retention, alert recipients, and privacy disclosure.", "Nominate mobile crash triage and release owners."], steps=["Configure Firebase apps and plugin/build scripts.", "Add handlers/keys/privacy controls.", "Upload symbols for test release.", "Force test crash/non-fatal on physical devices.", "Enable production alerts and release process."], tests=["Forced fatal/non-fatal on Android/iOS appears with correct build and readable symbols.", "Background/async Flutter errors are captured.", "User ID is pseudonymous and cleared/changed on account switch.", "No PII/sensitive values are attached to reports."], security=["Use only pseudonymous user ID and allow-listed keys.", "Protect Firebase roles/CI tokens and separate environments.", "Document collection consent/retention and account-deletion expectations."], source_label="Crashlytics for Flutter", source="https://firebase.google.com/docs/crashlytics/flutter/get-started"),

        dict(num=25, name="Firebase App Check", purpose="Attestation that requests originate from genuine AMORAA app instances and reduction of scripted/API abuse.", priority="P0", disposition="Implement after final signing/app IDs", current="Absent. Express accepts JWT-authenticated traffic without app attestation; Firebase resources/config are absent.", dependency="High - Firebase/Play/Apple app setup and production signing; 2-10 client days.", effort="L", layer="Flutter + custom Express backend", recommendation="Use App Check as defense in depth, not as user authentication. Start in monitoring mode, verify tokens on high-risk Express routes, then enforce gradually with documented web/debug behavior and emergency bypass.", sdk="firebase_app_check; Play Integrity (Android), App Attest/DeviceCheck (Apple), reCAPTCHA provider for web; Firebase Admin SDK token verification on custom backend.", frontend="Activate provider after Firebase initialization, fetch/refresh App Check token, attach X-Firebase-AppCheck to API requests, and handle attestation failure without blocking account recovery indiscriminately.", backend="Verify App Check token/project/app ID before selected routes, cache JWKS/admin verification, log only result/fingerprint, rate limit, and support monitor/enforce/kill-switch modes.", database="No token storage. Add SecurityEvent/abuse counters if not already present; store verification outcome/app ID, not raw token.", env="FIREBASE_PROJECT_ID/admin credentials; APP_CHECK_MODE monitor|enforce; allowed Firebase app IDs; debug tokens only in non-production secrets.", webhooks="None. App Check token verification is request-time.", team=["Configure providers for dev/staging/prod and Flutter token header injection.", "Add Express verification middleware and route risk tiers.", "Run monitoring to measure false rejects, then staged enforcement.", "Add abuse dashboards, debug/test strategy, emergency rollback, and documentation."], client=["Provide Firebase/GCP owner access, final package/bundle/web domain, Play Console/Apple Developer access, and release signing.", "Approve which routes require enforcement, rollout risk, supported device/OS policy, and customer-support fallback.", "Provide physical test devices including older/non-standard devices within support policy.", "Nominate security/operations owner for false-positive review and emergency disable."], steps=["Finalize IDs/signing and configure attestation providers.", "Implement Flutter token transport and backend monitor middleware.", "Analyze metrics and fix compatibility gaps.", "Enforce high-risk routes, then broaden.", "Run quarterly keys/provider review."], tests=["Valid release, debug/staging, emulator/simulator, rooted/jailbroken or unsupported, replayed/expired/wrong-app token.", "JWT without App Check is monitored/rejected according to route/mode.", "Outage/kill switch and clock/JWKS refresh behavior.", "Web provider and accessibility/account-recovery flows do not deadlock users."], security=["App Check does not replace JWT, authorization, rate limits, or server validation.", "Never log raw attestation tokens or ship debug tokens in release builds.", "Enforce expected project/app IDs and least-privilege Firebase admin access."], source_label="Firebase App Check Flutter setup", source="https://firebase.google.com/docs/app-check/flutter/default-providers"),

        dict(num=26, name="Firebase Remote Config", purpose="Remote feature flags, controlled rollouts, kill switches, copy/configuration, and experiments.", priority="P1", disposition="Implement after analytics/governance", current="Absent; features and thresholds are compiled into Flutter/backend code.", dependency="Medium - Firebase roles and product approval/flag owners; 1-3 client days.", effort="M", layer="Flutter + optional backend mirror", recommendation="Use for non-secret presentation/behavior flags only. Authorization, price, entitlement, KYC, security limits, and provider secrets remain server-authoritative.", sdk="firebase_remote_config; optional Remote Config REST/Admin access for release automation.", frontend="Set safe in-code defaults, fetch/activate after consent/network policy, typed flag registry, minimum fetch interval by environment, exposure events, and graceful offline fallback.", backend="Optionally mirror a small approved flag set through server config; never trust client Remote Config for access control.", database="FlagExposure or ExperimentAssignment only if needed for server reconciliation; otherwise analytics event outbox captures exposure. No secrets.", env="Firebase config and REMOTE_CONFIG_MIN_FETCH_INTERVAL; CI credentials if templates are versioned/deployed.", webhooks="None for standard use. CI publishes versioned templates with review/rollback.", team=["Define typed flag registry, owners, default/fail-safe values, expiry dates, and naming.", "Implement fetch/activate/exposure and offline/error behavior.", "Create review/publish/rollback process and template export in source control.", "Add guardrails preventing security/price/entitlement decisions in client flags."], client=["Provide Firebase Remote Config access and nominate product/engineering publishers/approvers.", "Approve initial flags, targeting attributes, rollout cohorts, experiment metrics, and privacy constraints.", "Approve emergency kill-switch owners and after-hours escalation.", "Provide release calendar and business acceptance criteria."], steps=["Establish governance and defaults.", "Implement typed service and exposure events.", "Create non-production template/flags.", "Test staged percentage/audience rollout and rollback.", "Promote to production with expiry review."], tests=["Defaults before fetch/offline, fetch failure, throttling, stale config, invalid type, rollback.", "Targeting and exposure event match expected cohort without sensitive attributes.", "Security/entitlement remains enforced when a flag is tampered with.", "Kill switch takes effect within documented fetch behavior."], security=["All Remote Config values are public to app users; never store secrets.", "Do not target by sensitive dating/KYC traits.", "Use least-privilege publishing, review, audit, and flag expiry."], source_label="Firebase Remote Config for Flutter", source="https://firebase.google.com/docs/remote-config/get-started?platform=flutter"),

        dict(num=27, name="OpenAI API / approved LLM provider", purpose="Roadmap-dependent assistant, matching explanations, icebreakers, bio/prompt analysis, and safety-aware AI features.", priority="P1", disposition="Decision gate; backend-only implementation", current="No provider SDK/API. Active matching explanation is deterministic; AI coach/icebreaker UI exists but must not claim external AI results without integration.", dependency="Very High - approved use cases, budget, safety/privacy/legal review, vendor account; 1-3 weeks.", effort="XL", layer="Backend + Flutter UI + MySQL audit", recommendation="Start with one narrow, low-risk, user-initiated use case (for example, profile-based icebreakers) and a non-AI fallback. Never send KYC, raw private messages, precise location, contact details, or full profile history; require product/legal approval before automated match decisions.", sdk="Official OpenAI Node SDK and Responses API (or approved equivalent), moderation/safety controls, structured outputs, store:false where appropriate.", frontend="Call AMORAA backend only, show loading/cancel/retry/disclaimer, allow report/feedback, preserve deterministic fallback, and avoid representing generated text as objective compatibility.", backend="Authorize/rate limit, assemble minimized/redacted prompt, version templates, request structured output, validate schema/content, apply safety filters, cache if allowed, record cost/latency/outcome, and provide circuit breaker/fallback.", database="AiRequests (userId, useCase, promptVersion, model, inputHash, outputHash or approved text, safetyStatus, tokens/cost, latency, providerRequestId, consentVersion, createdAt) and AiFeedback; no raw KYC/messages by default.", env="OPENAI_API_KEY, OPENAI_PROJECT_ID/ORG if used, model IDs, prompt version, timeouts, budgets, store flag; key server-only in secret manager.", webhooks="Usually none for synchronous Responses; background/batch callbacks/polling only if selected and compatible with retention policy.", team=["Define approved use case, data-flow/privacy/safety threat model, quality rubric, and non-AI fallback.", "Implement backend-only provider adapter, prompts/structured output, moderation, rate/cost limits, and audit.", "Integrate Flutter UX, user feedback/reporting, and transparency copy.", "Run offline evals/red-team tests before small cohort rollout and monitor regressions."], client=["Create client-owned OpenAI organization/project, add billing/usage limits, invite least-privilege developers, and provide project API key via secret manager.", "Approve exact AI use cases, model/provider, budget, retention/data controls, regions, DPA/legal terms, age/safety policy, and human escalation.", "Provide product-approved prompt tone, prohibited outputs, disclaimers, quality examples, evaluation dataset stripped of sensitive data, and fallback behavior.", "Nominate AI product, safety, privacy/legal, finance, and incident owners."], steps=["Use-case/legal/safety approval and provider project.", "Build redacted prompt/eval dataset and baseline.", "Implement backend adapter/audit/rate/cost controls.", "Integrate UI and run red-team/evals.", "Limited Remote Config rollout with rollback/fallback."], tests=["Golden-set quality/structure, hallucination, bias/fairness, prompt injection, abuse, sexual/minor safety, PII leakage, timeout/rate limit/provider outage.", "Verify store/retention setting and that prohibited fields never enter requests/logs.", "Cost/latency/concurrency limits and circuit breaker.", "User report/feedback and deterministic fallback end-to-end."], security=["API key never ships in Flutter; use project-scoped key and spend limits.", "Data minimization/redaction before the provider; no KYC, biometrics, contact info, private messages, or precise location.", "Treat outputs as untrusted: validate, moderate, label, and avoid autonomous entitlement/safety decisions."], source_label="OpenAI API quickstart", source="https://platform.openai.com/docs/quickstart/make-your-first-api-request"),

        dict(num=28, name="Transactional email provider (SES / SendGrid / approved)", purpose="Account recovery, verification, receipts, membership lifecycle, safety alerts, support, and operational email delivery.", priority="P0", disposition="Choose provider and complete existing SMTP abstraction", current="Nodemailer/SMTP adapter exists; production refuses to start without SMTP, but credentials are empty and only password-reset email is implemented.", dependency="Very High - domain/DNS, sender verification, provider onboarding/billing, legal templates; 2 days to several weeks.", effort="M/L", layer="Backend + MySQL delivery ledger", recommendation="Keep the provider-neutral mail service but choose SES or SendGrid. Use domain authentication (SPF/DKIM/DMARC), provider webhooks, templated versioning, suppression handling, and separate transactional/marketing streams.", sdk="AWS SES v2 SDK or SMTP; SendGrid Mail API/SMTP; Nodemailer transport retained behind adapter.", frontend="No SDK. Ensure email preferences, resend UX, support links, and user-visible delivery failure states are truthful.", backend="Provider adapter, template renderer/versioning, queue/retry, idempotency, bounce/complaint/suppression webhooks, unsubscribe/preferences by category, and reconciliation.", database="EmailMessages/NotificationDeliveries (template, recipientHash/userId, providerMessageId, status, attempts, lastError, sent/delivered/bounced/complained timestamps) plus EmailSuppressions; avoid storing rendered sensitive bodies.", env="EMAIL_PROVIDER plus SES region/from/configuration set or SENDGRID_API_KEY; EMAIL_FROM/REPLY_TO; webhook signing secret; SMTP fields only for selected transport.", webhooks="Delivery, bounce, complaint, deferral, open/click only if approved; verify SNS/signature or Event Webhook signature and deduplicate provider event ID.", team=["Select provider and implement adapter/queue/templates/delivery ledger.", "Configure authenticated sending domain and production access/reputation controls.", "Implement webhooks, suppression, retry/idempotency, preference categories, and monitoring.", "Create transactional template catalog and accessibility/client rendering QA."], client=["Choose/create company provider account, billing and owner/admin access; for SES provide AWS account/region and approve production-access request.", "Provide DNS access to add SPF/DKIM/DMARC/custom return-path records and approve From/Reply-To/support addresses.", "Provide business website, legal address, privacy/terms, unsubscribe/contact requirements, brand assets/copy/locales, and transactional vs marketing classification.", "Nominate deliverability, support, privacy/legal, and finance owners; provide test inboxes across Gmail/Outlook/Apple."], steps=["Provider selection/domain authentication.", "Implement adapter/templates/queue/ledger.", "Configure webhooks/suppression and non-production stream.", "Run deliverability/render/security tests.", "Request/approve production sending and gradual warm-up."], tests=["Send/deliver/bounce/complaint/defer/suppress/retry/duplicate and provider outage.", "SPF/DKIM/DMARC pass; links/from/reply-to work; templates render accessibly on major clients.", "Forgot-password enumeration protection and token expiry remain intact.", "No secrets/OTP/password reset tokens appear in logs; webhook replay rejected."], security=["API/SMTP credentials in secret manager, rotated and least privilege.", "Suppress bounced/complaining recipients and separate marketing consent.", "Minimize stored body/recipient data; signed links and short-lived recovery tokens."], source_label="SendGrid sender identity", source="https://www.twilio.com/docs/sendgrid/for-developers/sending-email/sender-identity"),

        dict(num=29, name="Transactional SMS / WhatsApp provider", purpose="Non-OTP transactional alerts, compliant fallback communications, reminders, and support notifications.", priority="P1", disposition="Conditional; separate from Twilio Verify OTP", current="Generic sendSms sends OTP through Twilio Programmable Messaging only; no messaging service, templates, WhatsApp sender, consent ledger, delivery callbacks, or category routing.", dependency="Critical for India/WhatsApp - business verification, number/sender/DLT/WABA templates and legal consent; 2-8+ weeks.", effort="L", layer="Backend + MySQL delivery ledger", recommendation="Use Twilio Messaging/WhatsApp or another approved India provider only for specifically approved transactional categories. Keep OTP in Verify. Do not use SMS/WhatsApp for sensitive dating content or marketing without explicit consent and compliant templates.", sdk="Provider Messaging API/Node SDK; WhatsApp Business Platform via BSP; India DLT/template/sender integrations as applicable.", frontend="Add channel preferences/consent and clear opt-out/help text; do not expose provider keys. Deep links must re-authenticate before showing private content.", backend="Channel router, template IDs/variables, opt-in/out/quiet hours, messaging service/sender selection, queue/retry, delivery callbacks, status ledger, cost/frequency caps, and fallback rules.", database="MessagingConsents, MessageDeliveries, ApprovedTemplates, ChannelSuppressions; store body/template variables only as minimally required and avoid sensitive content.", env="MESSAGING_PROVIDER, API credentials, messaging service/sender/WhatsApp number, template/content SIDs, callback secret; DLT entity/header/template IDs where required.", webhooks="Delivery status and inbound opt-out/reply callbacks; verify provider signature, deduplicate, and honor STOP immediately.", team=["Separate Verify OTP from general messaging and define approved category/channel matrix.", "Implement template-based provider adapter, consent/preferences, queue/status callbacks, rate/cost caps, and fallback.", "Configure WhatsApp/DLT/sender/template approvals and operational dashboards.", "Add opt-out, deep-link privacy, incident, and content QA."], client=["Select provider/BSP and complete company/business verification, billing, approved phone number ownership, WABA/Meta Business verification for WhatsApp, and India DLT Principal Entity/sender/header/template registrations where applicable.", "Provide legal entity/tax/address/authorized signatory/website/privacy details and secure account invitations/credentials.", "Approve channels, use cases, templates/locales, opt-in evidence, opt-out/help copy, quiet hours, frequency caps, fallback, and monthly budget.", "Nominate messaging compliance, marketing/operations, support, and finance owners."], steps=["Legal/provider/channel approval and registrations.", "Approve templates/consent/fallback matrix.", "Implement adapter/ledger/callbacks/preferences.", "Sandbox and real-number/template testing.", "Limited production rollout and delivery/cost monitoring."], tests=["Approved/unapproved template, opt-in/out/STOP, quiet hours, invalid number, carrier/provider failure, duplicate/retry, fallback, and account switch.", "Signed callback and replay rejection.", "WhatsApp template window/status and SMS DLT mapping where applicable.", "Lock-screen/message content reveals no sensitive relationship activity."], security=["Provider secrets backend-only and callbacks signature-verified.", "Explicit consent/opt-out and minimum necessary message content.", "No KYC, exact location, private message text, sexual/religious preferences, or sensitive match details."], source_label="Twilio WhatsApp Business Platform", source="https://www.twilio.com/docs/whatsapp"),

        dict(num=30, name="Customer support / helpdesk provider", purpose="Tickets, help center, escalation, identity-aware support, safety triage, and operational reporting.", priority="P1", disposition="Choose provider and support operating model", current="Flutter has local FAQ/support UI; no ticket API, external help center, support accounts, SLA workflow, or SupportTickets table.", dependency="Very High - tool procurement, agents/workflows/SLAs/legal access; 1-4 weeks.", effort="L", layer="Flutter + backend/admin + helpdesk", recommendation="Select Zendesk, Freshdesk, Intercom, or approved equivalent based on support/safety workflow. Use backend-mediated OAuth/service integration; if Zendesk is chosen, design for OAuth rather than long-lived API tokens because token authentication is being retired.", sdk="Selected helpdesk Tickets/Users/Help Center API, OAuth 2.0/service credentials, webhooks, and optional Flutter/web messenger only if privacy-approved.", frontend="Add authenticated contact form, category/severity, attachment consent, ticket reference/status, help-center deep links, and emergency/safety routing; do not expose provider credentials.", backend="Create/update requester/ticket, attach sanitized files, map user/ticket IDs, verify provider webhooks, enforce rate limits, support deletion/anonymization, and route critical safety cases.", database="SupportTickets (userId, provider, externalTicketId unique, category, priority, status, created/updated/resolved timestamps), SupportTicketEvents, attachment MediaAssets; store minimal duplicate content.", env="SUPPORT_PROVIDER, subdomain/base URL, OAuth client/secret/refresh token or service credentials, webhook signing secret, help center URL.", webhooks="Ticket/status/comment/user events; verify signature/timestamp, deduplicate, and avoid feedback loops.", team=["Run provider/workflow decision and map support/safety categories, SLAs, escalation, roles, and data access.", "Implement backend proxy, ticket mappings, sanitized attachments, signed webhooks, and status sync.", "Integrate Flutter contact/status/help-center UX and operational alerts.", "Add retention/deletion/export, audit, rate-limit, and incident runbooks."], client=["Select/purchase client-owned helpdesk plan and provide organization owner/admin invitations, billing, support domain/subdomain, and brand assets.", "Provide support agent emails/roles, business hours/time zones/languages, categories/forms/macros, SLAs, escalation matrix, safety/emergency procedure, and approval boundaries.", "Provide help-center content, Privacy/Terms/Community Guidelines/Refund/Cancellation links, support/grievance/DPO contacts, and attachment/retention policy.", "Approve OAuth app/service account, webhook configuration, SSO/customer identity, reporting, and who may access sensitive tickets."], steps=["Select provider and design operating model.", "Configure sandbox/forms/roles/help center/OAuth.", "Implement backend ticket/webhook mapping and Flutter UX.", "Run agent, privacy, abuse, attachment, and escalation tests.", "Train agents and launch with SLA monitoring."], tests=["Anonymous/authenticated create, duplicate/rate limit, attachment type/size/scan, status/comment sync, agent close/reopen, provider outage.", "Signed/replayed webhook, wrong user/ticket access, deletion/anonymization/export.", "Critical safety category reaches correct on-call path; non-emergency copy is clear.", "Support users cannot see unrelated account/KYC/chat data."], security=["Backend-mediated OAuth/least privilege; no provider secret in app.", "Strict role-based ticket access, audit, redaction, and minimal replicated content.", "Separate emergency/safety escalation from ordinary support and document lawful disclosure process."], source_label="Zendesk webhook security", source="https://developer.zendesk.com/documentation/webhooks/webhook-security-and-authentication/"),
    ]


def build_architecture_diagram(path: Path):
    width, height = 1800, 1000
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 46)
        head_font = ImageFont.truetype("arialbd.ttf", 28)
        body_font = ImageFont.truetype("arial.ttf", 22)
        small_font = ImageFont.truetype("arial.ttf", 18)
    except OSError:
        title_font = ImageFont.load_default()
        head_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.text((60, 34), "AMORAA target third-party integration architecture", fill="#19283F", font=title_font)
    draw.text((62, 92), "Authoritative business state remains in Express + MySQL; providers receive the minimum approved data.", fill="#667085", font=body_font)

    boxes = [
        (70, 180, 470, 420, "Flutter clients", ["Android / iOS / Web", "Consent + SDK adapters", "JWT + App Check", "No provider secrets"]),
        (700, 160, 1110, 450, "Node.js / Express", ["Authentication + policy", "Integration adapters", "Webhook verification", "Outbox / retries", "Unified entitlements"]),
        (700, 590, 1110, 870, "MySQL", ["Users + profiles", "Payments + subscriptions", "Provider mappings", "Event outbox / audit", "Minimal sensitive data"]),
        (1340, 170, 1730, 865, "Third-party services", ["Firebase / GA4 / Ads", "OAuth / Verify / email", "Play / App Store / Razorpay", "Maps / media / chat / KYC", "MMP / AI / support"]),
    ]
    for x1, y1, x2, y2, title, lines in boxes:
        fill = "#F8E9EE" if "Flutter" in title else "#E8EEF5" if "Express" in title else "#E8F5EE" if title == "MySQL" else "#FFF4D6"
        outline = "#B54A66" if "Flutter" in title else "#2E74B5" if "Express" in title else "#217A57" if title == "MySQL" else "#946200"
        draw.rounded_rectangle((x1, y1, x2, y2), radius=26, fill=fill, outline=outline, width=4)
        draw.text((x1 + 26, y1 + 22), title, fill=outline, font=head_font)
        yy = y1 + 75
        for line in lines:
            draw.ellipse((x1 + 28, yy + 9, x1 + 38, yy + 19), fill=outline)
            draw.text((x1 + 52, yy), line, fill="#26313D", font=body_font)
            yy += 42

    def arrow(start, end, color="#667085", both=False):
        draw.line((start, end), fill=color, width=6)
        x2, y2 = end
        x1, y1 = start
        import math
        angle = math.atan2(y2 - y1, x2 - x1)
        for direction, point in [(angle, end)] + ([(angle + math.pi, start)] if both else []):
            px, py = point
            size = 20
            pts = [
                (px, py),
                (px - size * math.cos(direction - 0.55), py - size * math.sin(direction - 0.55)),
                (px - size * math.cos(direction + 0.55), py - size * math.sin(direction + 0.55)),
            ]
            draw.polygon(pts, fill=color)

    arrow((470, 290), (700, 290), both=True)
    arrow((905, 450), (905, 590), both=True)
    arrow((1110, 300), (1340, 300), both=True)
    arrow((1110, 720), (1340, 720), both=True)
    draw.text((505, 235), "HTTPS / Socket", fill="#667085", font=small_font)
    draw.text((915, 505), "Sequelize", fill="#667085", font=small_font)
    draw.text((1150, 245), "SDK/API", fill="#667085", font=small_font)
    draw.text((1150, 665), "Webhooks", fill="#667085", font=small_font)
    image.save(path, "PNG", optimize=True)


def configure_styles(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.86)
    sec.right_margin = Inches(0.86)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.38)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 16, NAVY, 14, 8),
        ("Heading 2", 12.5, ROSE, 10, 5),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        st._element.rPr.rFonts.set(qn("w:ascii"), st.font.name)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), st.font.name)
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Aptos"
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15


def add_running_furniture(doc: Document):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("AMORAA  |  Third-Party API Implementation Report")
        set_font(r, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        add_page_number(fp)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(88)
    r = p.add_run("AMORAA")
    set_font(r, size=11, color=ROSE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Third-Party API\nImplementation Report")
    set_font(r, name="Aptos Display", size=29, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("Flutter frontend | Node.js/Express backend | MySQL")
    set_font(r, size=13.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(78)
    r = p.add_run("Client-ready implementation roadmap, account/document checklist, security plan, and testing reference")
    set_font(r, size=11, color=ROSE, bold=True)
    add_key_table(doc, [
        ("Prepared for", "AMORAA project stakeholders and client account owners"),
        ("Repository reviewed", "D:\\Projects\\amora_ai | commit 2cb0ed59 | 14 August 2026"),
        ("Document date", "14 August 2026"),
        ("Status", "Final implementation planning reference - credentials not included"),
        ("Classification", "Client confidential"),
    ], widths=(2200, 7160))
    add_callout(doc, "Important", "This report separates engineering work from client-owned prerequisites. Vendor pricing, SDK versions, product policies, and approval requirements must be revalidated at the start of each integration.", "info")
    doc.add_page_break()


def add_front_matter(doc: Document, records):
    add_toc(doc)
    doc.add_heading("Document purpose and decision rules", level=1)
    doc.add_paragraph("This report converts the supplied 30-service inventory into an implementation plan grounded in the current AMORAA repository. It identifies what already exists, what must be completed or replaced, how each integration connects to Flutter, Express, and MySQL, and exactly what the client must provide before work can start.")
    add_callout(doc, "Core decision rule", "‘All required integrations’ does not mean installing every alternative. The project must choose one MMP (AppsFlyer or Adjust), one managed push approach (direct FCM or OneSignal), one managed chat path (Stream or Sendbird, or retain the existing chat), and one primary public-media stack (Cloudinary or AWS).", "decision")
    add_bullets(doc, [
        "P0 integrations are launch-critical only after the relevant channel/product is approved and the required client accounts are ready.",
        "P1 integrations are staged, roadmap-dependent, optional, or require a commercial decision before implementation.",
        "Backend/MySQL remains authoritative for user identity, payment verification, entitlement, KYC status, support mapping, and consent/audit records.",
        "No live secret values were copied into this report. The existing Backend/.env was checked only for presence/configuration state.",
    ])

    doc.add_heading("Executive assessment of the existing project", level=1)
    add_key_table(doc, [
        ("Flutter application", "202 Dart source files; authenticated HTTP client with refresh-token rotation; secure token storage; feature repositories; Socket.IO chat; Razorpay UI; Google sign-in package."),
        ("Express backend", "Express 4 + Sequelize + MySQL; JWT/refresh tokens; OTP/email/SMS adapters; Google token verification; FCM HTTP v1 adapter; Razorpay verification/webhook; media/KYC local storage; Socket.IO."),
        ("MySQL", "23 ordered migrations in the latest internal audit; users, OTP/refresh tokens, profiles, discover, matches, chat, events, subscriptions/payments, notifications/devices/delivery, and identity verification."),
        ("Existing verification evidence", "Internal audit dated 12 August 2026 reports Flutter analysis/build/tests, backend tests, applied migrations, restart persistence, and multi-user isolation passing for active internal flows; provider E2E remained credential-blocked."),
        ("Provider credentials", "Firebase, Razorpay, Twilio, Google OAuth, and SMTP production credentials are absent/not configured in the checked environment. CORS is configured."),
    ])
    add_callout(doc, "Release blockers before provider onboarding", "Android uses com.example.amora_ai and debug signing; iOS has no entitlements file; no google-services.json, GoogleService-Info.plist, or firebase_options.dart exists; the web shell has no tag/consent layer. Final identifiers, signing, store ownership, domains, privacy/legal documents, and secret management must be completed first.", "risk")

    doc.add_heading("Target integration architecture", level=1)
    doc.add_paragraph("Provider SDKs handle collection or user experience, while Express validates identity/business rules and MySQL stores the durable, auditable state. Webhooks are untrusted inputs until signature, timestamp, environment, ownership, and idempotency checks pass.")
    add_picture_with_alt(
        doc,
        DIAGRAM_FILE,
        Inches(6.45),
        "AMORAA target integration architecture",
        "Architecture diagram showing Flutter and web clients connected through Express authentication, consent, and integration adapters to MySQL ledgers, then to third-party providers and verified webhooks.",
    )
    cap = doc.add_paragraph("Figure 1. Target integration boundaries and systems of record.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_font(run, size=8.5, color=MUTED, italic=True)

    doc.add_heading("Canonical cross-provider data contracts", level=1)
    add_bullets(doc, [
        "Event ID: globally unique UUID generated at the authoritative action; reused for browser/server deduplication.",
        "User identity: internal numeric user ID; external provider subjects live in ExternalIdentities; email/phone are never general analytics identifiers.",
        "Consent snapshot: category, version, jurisdiction, granted/withdrawn timestamps, and collection source available to every marketing delivery.",
        "Entitlement: MySQL Subscription/Payment state; store/provider callbacks trigger verification, not direct grants.",
        "Provider event ledger: provider event ID, signature result, payload hash, processing status, attempts, and immutable timestamps.",
        "Media: MediaAssets row points to provider asset ID/key, purpose/access class, checksum, moderation and deletion state.",
    ])
    add_callout(doc, "Prohibited cross-provider data", "Aadhaar/KYC material, biometrics, OTPs, private message content, exact location, contact details, sexual orientation, religion, and sensitive dating preferences must not be sent to analytics, ads, MMPs, error tools, AI, or support by default.", "risk")

    doc.add_heading("Integration disposition matrix", level=1)
    matrix = doc.add_table(rows=1, cols=5)
    matrix.style = "Table Grid"
    matrix.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = matrix.rows[0].cells
    for i, value in enumerate(["#", "Integration", "P", "Disposition", "Current state"]):
        hdr[i].text = value
        shade(hdr[i]._tc, NAVY)
        for run in hdr[i].paragraphs[0].runs:
            set_font(run, size=8.5, color=WHITE, bold=True)
    set_repeat_table_header(matrix.rows[0])
    for item in records:
        row = matrix.add_row()
        set_cant_split(row)
        cells = row.cells
        values = [str(item["num"]), item["name"], item["priority"], item["disposition"], item["current"]]
        for i, value in enumerate(values):
            cells[i].text = value
            if item["num"] % 2 == 0:
                shade(cells[i]._tc, "FAFBFC")
            for run in cells[i].paragraphs[0].runs:
                set_font(run, size=7.8, color=INK, bold=(i in (0, 1)))
    set_table_geometry(matrix, [580, 2280, 520, 2500, 3480])
    table_borders(matrix, color="C9CFD7", size=4)

    doc.add_heading("Client-owned prerequisite package", level=1)
    doc.add_paragraph("The fastest route is to assemble one controlled onboarding package before implementation begins. Accounts should be owned by the client, with development access granted by invitation; credentials should not be created under personal developer ownership.")
    add_key_table(doc, [
        ("Recommended account aliases", "accounts@, engineering@, billing@, privacy@, security@, support@, marketing@ on the client domain; individual named admins with MFA rather than shared passwords."),
        ("Corporate identity", "Legal entity name, registration/CIN, GST/PAN/tax details as applicable, registered/operating address, website, authorized signatory, directors/beneficial ownership if requested, bank/settlement proof, billing card/contact."),
        ("App ownership", "Google Play Console organization; Apple Developer organization/App Store Connect; final package/bundle IDs; release signing/certificates; store listing IDs; test tracks/TestFlight."),
        ("Web/domain", "Production and staging domains, DNS access, TLS/hosting, support/privacy/terms/refund URLs, subdomains for API, links, media, analytics, email return path."),
        ("Legal/privacy", "Privacy Policy, Terms, Cookie Policy/CMP decision, Community Guidelines, refund/cancellation/subscription disclosures, consent text/versioning, DPDP/Aadhaar/KYC legal approval, DPO/grievance contact, retention/deletion schedule."),
        ("Operations", "Marketing/ad owners, support agents/SLA, trust & safety/KYC reviewers, on-call/security contacts, finance/reconciliation owner, monthly budget/alerts, production approval matrix."),
    ])

    doc.add_heading("Account and document checklist by provider family", level=1)
    checklist_rows = [
        ("Google/Firebase/Maps/Ads/GTM", "Corporate Google accounts, organization/billing, verified domains, final app IDs/signing fingerprints, OAuth consent assets, Ads advertiser/billing, GA/GTM publishers, Maps budgets."),
        ("Apple", "Organization membership, D-U-N-S/legal enrollment if applicable, Account Holder, certificates/keys, App Store agreements, banking/tax, products/pricing, APNs and Server API keys."),
        ("Meta/TikTok", "Verified business portfolios, ad accounts, payment methods, domains, pixels/datasets, campaign owners, data-sharing/consent approval."),
        ("Twilio/email/WhatsApp", "Business verification, funding, number/sender ownership, India DLT/WABA/templates where applicable, sending domain DNS, from/reply-to, opt-in/out copy."),
        ("Payments", "Razorpay KYC/legal/tax/bank/website/policies; Play/Apple merchant agreements and digital product catalog."),
        ("KYC", "Vendor contract/KYB, legal/DPIA, exact Aadhaar/document method approval, retention/deletion, authorized signatory, webhook/live approval."),
        ("MMP/deep links", "Vendor contract/billing, store IDs, ad-network access, ATT/privacy strategy, link domain/DNS, campaign taxonomy."),
        ("Cloud/media/chat/support/AI/observability", "Vendor owner/billing/region/DPA, data residency/retention, role matrix, volumes/SLA, API projects/keys through secret manager, incident and deletion owners."),
    ]
    add_key_table(doc, checklist_rows)

    doc.add_heading("Pre-implementation technical gates", level=1)
    add_numbered(doc, [
        "Replace Android applicationId com.example.amora_ai with the final reverse-domain ID and configure production signing; record SHA-1/SHA-256 fingerprints.",
        "Finalize iOS bundle ID/App ID, provisioning, entitlements, APNs, associated domains, and StoreKit/Sign in with Apple capabilities.",
        "Create dev/staging/prod domains and HTTPS API/webhook endpoints; lock CORS to explicit origins.",
        "Adopt a managed secret store and environment promotion process; remove provider secrets from local .env distribution and CI logs.",
        "Approve the event taxonomy, consent model, data retention/deletion, and provider data classification before analytics/ads/MMP SDK start.",
        "Create generic integration outbox/webhook-event/mapping migrations and a worker/retry/dead-letter pattern before adding many adapters.",
        "Define mobile-store versus web-payment product rules and one cross-platform entitlement state machine.",
        "Freeze choose-one vendor decisions and avoid duplicate SDKs/destinations until the prior option is deliberately replaced.",
    ])

    doc.add_heading("Indicative implementation roadmap", level=1)
    roadmap = [
        ("Phase 0 - ownership and governance", "Weeks 0-2+", "Identifiers/signing/domains, secret management, accounts, legal/privacy/consent, vendor decisions, event and entitlement contracts. External approvals can extend this phase."),
        ("Phase 1 - foundation/observability", "Weeks 2-4", "Firebase Core/Analytics/Crashlytics/App Check monitor mode/Remote Config foundation; GA4/GTM consent layer; Sentry backend; integration outbox and webhook ledger."),
        ("Phase 2 - identity and communications", "Weeks 3-6", "Google/Apple sign-in, Twilio Verify, email, FCM token/delivery completion; transactional messaging decision."),
        ("Phase 3 - monetization", "Weeks 5-9", "Unified entitlement migration; Play Billing/RTDN; StoreKit/Server API/Notifications V2; Razorpay web hardening; finance reconciliation."),
        ("Phase 4 - product infrastructure", "Weeks 7-12", "Maps/Places, Cloudinary or AWS media, KYC provider, chat build-vs-buy decision/migration if selected."),
        ("Phase 5 - acquisition stack", "Weeks 9-13", "MMP, deep links, Meta/Google Ads, TikTok if active, server-side GTM if approved; test campaigns and reconciliation."),
        ("Phase 6 - roadmap services", "Weeks 11-16+", "AI pilot, support/helpdesk, advanced campaigns/experiments, hardening, penetration/load/DR, runbooks, production cutovers."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, v in enumerate(["Workstream", "Indicative window", "Exit outcome"]):
        table.rows[0].cells[i].text = v
        shade(table.rows[0].cells[i]._tc, NAVY)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            set_font(run, size=9, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(roadmap):
        cells = table.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = v
            if i % 2:
                shade(cells[j]._tc, "FAFBFC")
            for run in cells[j].paragraphs[0].runs:
                set_font(run, size=8.6, color=INK, bold=(j == 0))
    set_table_geometry(table, [2500, 1600, 5260])
    table_borders(table)
    add_callout(doc, "Schedule note", "The windows assume parallel Flutter/backend/infrastructure work after client prerequisites are ready. Store, ad, KYC, WhatsApp/DLT, payment, and vendor procurement approvals are external lead times and are not guaranteed by engineering.", "info")

    doc.add_page_break()
    doc.add_heading("Cross-provider testing and release gates", level=1)
    add_bullets(doc, [
        "Contract/unit tests for every adapter, signature verifier, event mapper, consent gate, and error translation.",
        "Sandbox/test-environment integration tests with provider-owned test tools and idempotent webhook fixtures.",
        "Physical Android/iOS device matrix for push, OAuth, deep links, ATT, location, purchases, media, KYC, and account switching.",
        "Reconciliation from provider dashboard/API to MySQL and outbox; document expected reporting delay/discrepancy.",
        "Security tests: secret scanning, wrong-audience/wrong-environment tokens, signature/replay/out-of-order webhooks, object access, rate abuse, and least privilege.",
        "Privacy tests: consent denied/withdrawn, data deletion/export, log scrubbing, SDK identity reset, retention expiry, and prohibited-field scans.",
        "Operational tests: provider outage, timeouts/retry/dead-letter, quota/cost alarms, credential rotation, webhook backlog, rollback/kill switch, and incident notifications.",
        "Production smoke tests use client-approved test accounts/amounts/content and require business, finance, legal/privacy, and engineering sign-off where relevant.",
    ])
    doc.add_page_break()

    doc.add_heading("Detailed integration implementation plans", level=1)
    doc.add_paragraph("The following sections are the implementation reference for each item in the supplied master inventory. Each section includes current-state evidence, cross-layer changes, configuration, webhooks, ownership, sequence, testing, and security controls.")
    doc.add_page_break()


def add_appendices(doc: Document, records):
    doc.add_heading("Appendix A - Environment variable register", level=1)
    doc.add_paragraph("Exact names may be normalized during implementation. Values classified as secrets must be stored in a managed secret store and injected at runtime; they must not be committed, embedded in Flutter, or included in this document.")
    env_rows = [
        ("Firebase/Google", "FIREBASE_PROJECT_ID; workload identity/service account; Firebase app config; GA4_MEASUREMENT_ID (public); GA4_API_SECRET; GTM_CONTAINER_ID; SGTM_BASE_URL; Google Ads OAuth/action IDs; Maps restricted keys."),
        ("Auth/messaging", "GOOGLE_*_CLIENT_ID; APPLE_TEAM/KEY/SERVICE/CLIENT IDs and private key; TWILIO API key/secret + VERIFY_SERVICE_SID; email provider keys; FCM/APNs config; OneSignal IDs if selected."),
        ("Payments", "Play package/service identity/PubSub; Apple Store API issuer/key/private key; Razorpay key ID/secret/webhook secret."),
        ("Acquisition", "MMP dev key/app token/S2S token; Branch keys/domain; Meta pixel/dataset/access token; TikTok pixel/access token."),
        ("Product providers", "Maps server key; Cloudinary cloud/key/secret; AWS region/buckets/distribution/KMS; chat provider key/secret; KYC client/key/cert/webhook; support OAuth/webhook."),
        ("Observability/AI", "SENTRY_DSN/AUTH_TOKEN; OPENAI_API_KEY/PROJECT/model; release/environment/sampling/budget values."),
    ]
    add_key_table(doc, env_rows)

    doc.add_heading("Appendix B - Recommended new/extended MySQL structures", level=1)
    db_rows = [
        ("IntegrationOutbox / DeliveryAttempts", "Durable authoritative events, destinations, consent snapshot, retries, provider response, dead-letter."),
        ("ProviderWebhookEvents", "Provider/event ID unique, payload hash, signature result, environment, received/processed status and timestamps."),
        ("ExternalIdentities", "Google/Apple/provider subject unique, user mapping, linked/revoked timestamps, relay flags."),
        ("AttributionInstalls/Touches", "MMP and web campaign first/last touch, approved click IDs, consent and expiry."),
        ("StoreEvents + subscription extensions", "Play/Apple product/purchase/original transaction identifiers, lifecycle/acknowledgement, unique notification/message IDs."),
        ("MediaAssets/UploadSessions", "Provider asset ID/key, purpose/access class, checksum, sizes, moderation, version/deletion."),
        ("VenueLocations", "Google place ID, name/address/coordinates/categories/source/refresh time; references from events/date spots."),
        ("IdentityVerification extensions", "Provider reference, method, consent, masked/tokenized result, liveness/decision, expiry; no raw Aadhaar identifier."),
        ("AiRequests/AiFeedback", "Use case, prompt/model version, hashes, safety/cost/latency/provider request, feedback; minimized content."),
        ("SupportTickets/Events", "User/provider ticket mapping, category/priority/status/SLA timestamps and minimal sync state."),
        ("MessagingConsents/Deliveries/Suppressions", "Channel/category consent, template, provider message state, opt-out/bounce/complaint."),
    ]
    add_key_table(doc, db_rows)

    doc.add_heading("Appendix C - Client handover and acceptance checklist", level=1)
    add_bullets(doc, [
        "All production accounts are owned by the client organization; at least two client admins have MFA and recovery access.",
        "Development access is by named invitation and least privilege; shared passwords and emailed private keys are prohibited.",
        "A credential register records owner, purpose, environment, location, creation/rotation/expiry, and emergency revocation procedure.",
        "Final package/bundle/store IDs, signing fingerprints, domains, redirect URIs, and webhook URLs are documented and tested.",
        "Legal/privacy documents and consent versions cover every enabled SDK/destination; data retention/deletion and DPO/grievance contacts are approved.",
        "Provider sandboxes and production dashboards have test evidence, alert owners, budgets, quotas, and support contacts.",
        "Reconciliation, incident, outage, credential rotation, webhook backlog, refund, KYC escalation, support, and rollback runbooks are accepted.",
        "No unresolved choose-one conflict remains; duplicate SDKs/destinations are disabled before launch.",
        "Production smoke tests and final go-live sign-off are recorded by engineering plus the relevant client business/legal/finance owner.",
    ])

    doc.add_heading("Appendix D - Source and review basis", level=1)
    doc.add_paragraph("Repository evidence reviewed includes Flutter pubspec/main/config/auth and feature repositories; Android/iOS/web release configuration; Express server/config/routes/controllers/providers/storage/realtime; Sequelize models/migrations; environment example/configuration presence; and the internal production-readiness audit dated 12 August 2026. The report was prepared against commit 2cb0ed59 on 14 August 2026.")
    doc.add_paragraph("Official vendor documentation links are included in each integration section. Key regulatory references include UIDAI requesting-entity security requirements and MeitY's Digital Personal Data Protection Rules, 2025. This engineering report is not a substitute for legal, tax, payment-policy, or Aadhaar/KYC counsel.")
    p = doc.add_paragraph()
    p.add_run("MeitY DPDP Rules 2025: ").bold = True
    add_hyperlink(p, "Official document page", "https://www.meity.gov.in/documents/act-and-policies/digital-personal-data-protection-rules-2025-gDOxUjMtQWa")
    p = doc.add_paragraph()
    p.add_run("UIDAI regulations index: ").bold = True
    add_hyperlink(p, "Official updated regulations", "https://uidai.gov.in/en/about-uidai/legal-framework/updated-regulation")

    doc.add_heading("Appendix E - Final decisions required from the client", level=1)
    decisions = [
        "Final Android/iOS application identifiers, release signing ownership, store organizations, and production/staging domains.",
        "Consent/CMP, analytics/ads/MMP data-sharing, retention, ATT, audiences, and prohibited-data policy.",
        "AppsFlyer or Adjust; whether its deep links are sufficient or Branch is additionally justified.",
        "Direct FCM or OneSignal; no dual production send path.",
        "Retain current Socket.IO/MySQL chat or replace it with Stream/Sendbird; if replacing, choose vendor/region/migration policy.",
        "Cloudinary for public media, AWS for full stack, or a documented hybrid with KYC/private-media segregation.",
        "Approved KYC provider/method and written legal/data-retention decision before raw Aadhaar production processing.",
        "Play/App Store digital subscription catalog and Razorpay-eligible web/physical product boundary.",
        "Email provider, transactional SMS/WhatsApp scope, India DLT/WABA path, and support/helpdesk provider.",
        "OpenAI/LLM use case, model, safety/privacy/budget/retention and launch cohort; or explicit deferral.",
        "Server-side GTM hosting/budget and TikTok activation based on funded campaign plans.",
    ]
    add_numbered(doc, decisions)
    add_callout(doc, "Ready-to-start definition", "An integration is ready only when the client-owned account is verified and funded if needed, final IDs/domains are available, credentials are transferred securely, legal/consent decisions are approved, test users/devices exist, and a named client approver accepts the provider-specific checklist.", "success")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    build_architecture_diagram(DIAGRAM_FILE)
    records = integration_records()
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_front_matter(doc, records)
    for item in records:
        add_integration_section(doc, item)
    add_appendices(doc, records)
    add_running_furniture(doc)
    props = doc.core_properties
    props.title = "AMORAA Third-Party API Implementation Report"
    props.subject = "Client-ready implementation roadmap and prerequisites for 30 third-party service integrations"
    props.author = ""
    props.keywords = ""
    props.comments = ""
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
