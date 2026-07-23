from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path('output/docx/AMORA_AI_Stage_3_Final_Frontend_Readiness_Report.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = '0D1B3E'
ROSE = 'C8174A'
GOLD = 'B8882A'
BLUE = '2E74B5'
LIGHT_BLUE = 'EAF2F8'
LIGHT_GRAY = 'F2F4F7'
MID_GRAY = '6B7280'
GREEN = '107C41'
AMBER = '9A6700'
RED = 'B42318'
WHITE = 'FFFFFF'
BLACK = '1A1A2E'
TABLE_WIDTH = 9360


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, name='Calibri', size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), color)


def cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def keep_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn('w:w'), str(widths[index]))
            tc_w.set(qn('w:type'), 'dxa')
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths, status_column=None):
    result = doc.add_table(rows=1, cols=len(headers))
    result.style = 'Table Grid'
    header = result.rows[0]
    repeat_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(value))
        set_run_font(run, size=9, color=WHITE, bold=True)
    for values in rows:
        row = result.add_row()
        keep_row(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            if len(result.rows) % 2 == 1:
                shade(cell, 'FAFBFC')
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(value))
            color = BLACK
            bold = False
            if status_column == index:
                upper = str(value).upper()
                color = GREEN if upper.startswith('PASS') else AMBER if upper.startswith('PARTIAL') else RED
                bold = True
            set_run_font(run, size=8.5, color=color, bold=bold)
    set_table_geometry(result, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return result


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instruction
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, end])


def add_para(doc, text='', bold_lead=None, italic=False, color=BLACK, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=11, color=color, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=color, italic=italic)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.add_run(text)
    return p


def callout(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    result = doc.add_table(rows=1, cols=1)
    result.style = 'Table Grid'
    cell = result.cell(0, 0)
    shade(cell, fill)
    cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    set_table_geometry(result, [TABLE_WIDTH])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 12, '1F4D78', 8, 4),
]:
    style = styles[name]
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = header.add_run('AMORA AI  |  STAGE 3 FRONTEND LAUNCH READINESS')
set_run_font(hr, size=8.5, color=MID_GRAY, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run('Confidential - AMORA AI  |  ')
set_run_font(fr, size=8, color=MID_GRAY)
add_field(footer, 'PAGE')

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(18)
title.paragraph_format.space_after = Pt(4)
run = title.add_run('FINAL FRONTEND\nLAUNCH READINESS REPORT')
set_run_font(run, size=24, color=NAVY, bold=True)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
run = subtitle.add_run('Stage 3 - SOW Verification, Production QA and Release Assessment')
set_run_font(run, size=14, color=ROSE, bold=True)

metadata = [
    ('Project', 'AMORA AI Flutter Application'),
    ('Source of truth', 'Amora_AI_SOW.pdf - Version 1.0, 43 pages, dated 28 June 2026'),
    ('Assessment date', '14 July 2026'),
    ('Scope', 'Existing Flutter frontend only; business logic and backend preserved'),
    ('Release status', 'CONDITIONAL RELEASE CANDIDATE'),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    a = p.add_run(f'{label}: ')
    set_run_font(a, size=10.5, color=BLACK, bold=True)
    b = p.add_run(value)
    set_run_font(b, size=10.5, color=BLACK)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
table(doc, ['READINESS', 'SOW CORE', 'AUTOMATED QA', 'RELEASE BUILD'], [
    ('92%', '32 / 32 present', '11 / 11 passed', 'APK built'),
], [2340, 2340, 2340, 2340])
callout(
    doc,
    'Executive decision',
    'The existing frontend is a conditional release candidate. Analyzer, route rendering, responsive widget tests, text-only chat controls, accessibility touch targets and the Android release build all pass. Final production approval remains gated by physical-device QA and APK-size optimization against the SOW target.',
    fill='FFF7ED',
    accent=AMBER,
)
add_para(doc, 'Prepared as the requested professional Word deliverable. No production business logic, API, repository, service, provider, controller, model, payment, authentication, AI, route, or navigation behavior was changed during Stage 3.', italic=True, color=MID_GRAY)

doc.add_page_break()
heading(doc, '1. Executive Summary', 1)
add_para(doc, 'The complete 43-page SOW was reviewed as the governing baseline. All 32 SOW core screens have corresponding existing Flutter implementations, and the broader registered route surface was audited through an expanded 88-route widget test matrix. The prior Stage 1 design system and Stage 2 screen migrations were preserved.')
bullet(doc, 'PASS - flutter analyze completed with zero issues.')
bullet(doc, 'PASS - all 11 automated tests passed after flutter clean and flutter pub get.')
bullet(doc, 'PASS - release APK built successfully with icon tree-shaking enabled.')
bullet(doc, 'PASS - all 88 registered route names render without exceptions across 320, 360, 375, 390, 412, 414, 768 and 1024dp widths.')
bullet(doc, 'PARTIAL - SOW performance target is less than 30 MB; the universal release APK is 76.77 MiB.')
bullet(doc, 'NOT EXECUTED - physical Android device, gesture navigation, keyboard, status/navigation bar and real-device animation review require connected hardware.')
callout(doc, 'Approved scope reconciliation', 'The SOW describes media-rich chat, while the current approved AMORA frontend scope is text-only. Stage 3 preserves text messaging, emoji, timestamps, read status and typing UI, and does not reintroduce call, recording or attachment controls. The legacy shared-media route remains registered because route removal is prohibited.', fill='FDF2F8', accent=ROSE)

heading(doc, '2. SOW Coverage', 1)
add_para(doc, 'Coverage is measured against Section 6 of the SOW (32 distinct screens). “Pass” means an existing registered Flutter destination corresponds to the SOW screen. “Pass - scoped” identifies an approved frontend deviation without changing routes or backend code.')

sow_rows = [
    (1, 'Splash Screen', '/splash', 'PASS - animated brand launch exists'),
    (2, 'Onboarding', '/onboarding', 'PASS - existing onboarding flow'),
    (3, 'Login', '/login, /phone-login', 'PASS - existing auth entry points'),
    (4, 'Signup', '/signup', 'PASS - existing registration UI'),
    (5, 'Compatibility Onboarding', '/compatibility', 'PASS - existing quiz UI'),
    (6, 'Profile Setup', '/profile-setup', 'PASS - existing profile setup UI'),
    (7, 'KYC Verification', '/kyc', 'PASS - existing verification UI'),
    (8, 'Home / Discovery Feed', '/home', 'PASS - existing home/discovery UI'),
    (9, 'Browse Grid', '/browse, /discover', 'PASS - aliases preserved'),
    (10, 'Advanced Filters', '/filters', 'PASS - existing filters UI'),
    (11, 'Profile Detail', '/profile-detail', 'PASS - existing profile UI'),
    (12, 'Match Screen', '/match', 'PASS - existing match UI'),
    (13, 'Chat / Messages', '/chats, /chat-detail', 'PASS - SCOPED: text-only UI'),
    (14, 'Super Like', '/super-like', 'PASS - existing confirmation UI'),
    (15, 'Send Gift', '/send-gift', 'PASS - existing commerce UI'),
    (16, 'Events Browse', '/events', 'PASS - existing event browse UI'),
    (17, 'Event Detail', '/event-detail', 'PASS - existing event detail UI'),
    (18, 'Ticket Booking', '/ticket-booking', 'PASS - existing booking UI'),
    (19, 'My Events', '/my-events', 'PASS - existing event management UI'),
    (20, 'Subscriptions', '/subscription', 'PASS - existing plan UI'),
    (21, 'Payment', '/payment', 'PASS - existing payment presentation'),
    (22, 'AI Dating Coach', '/ai-coach', 'PASS - existing coach UI'),
    (23, 'AI Icebreakers', '/ai-icebreakers', 'PASS - existing icebreaker UI'),
    (24, 'Amora Wallet', '/wallet', 'PASS - existing wallet UI'),
    (25, 'Refer & Earn', '/refer-earn', 'PASS - existing referral UI'),
    (26, 'Profile Settings', '/profile-settings', 'PASS - existing settings UI'),
    (27, 'Safety & Privacy', '/safety-privacy', 'PASS - existing safety UI'),
    (28, 'FAQ & Support', '/faq-support', 'PASS - existing support UI'),
    (29, 'Host Dashboard', '/host-dashboard', 'PASS - existing dashboard UI'),
    (30, 'Admin Panel', '/admin-panel', 'PASS - existing admin UI'),
    (31, 'Date Spots Map', '/date-spots', 'PASS - existing date-spots UI'),
    (32, 'Notifications Hub', '/notifications', 'PASS - existing notification UI'),
]
table(doc, ['#', 'SOW SCREEN', 'EXISTING DESTINATION', 'RESULT / NOTES'], sow_rows, [500, 2100, 2700, 4060], status_column=3)

doc.add_page_break()
heading(doc, '3. Screens Reviewed', 1)
add_para(doc, 'The audit covered all 32 SOW core screens plus the existing extended production and roadmap route set. The automated matrix now contains 88 registered route names, including aliases and preserved legacy destinations.')
for group, screens in [
    ('Launch and authentication', 'Splash, Landing, Onboarding, Auth, Login, Phone OTP, Signup, Compatibility Onboarding'),
    ('Discovery and matching', 'Home, Browse Grid, Discover alias, Advanced Filters, Profile Detail, Matches, Match, Super Like, Why We Matched, Liked You, Boost'),
    ('Messaging', 'Chat List, Chat Detail, preserved shared-media route alias, Event Group Chat'),
    ('Events and commerce', 'Events Browse, Event Detail, Ticket Booking, My Events, Waitlist, Feedback, Date Spots, Gift Catalog, Send Gift'),
    ('Premium and AI', 'Subscription, Payment, Wallet, Referral, AI Dating Coach, AI Icebreakers, Dating Recap'),
    ('Profile, settings and safety', 'Profile, Profile Setup, Photo Manager, Bio Builder, Dealbreakers, KYC, Settings, Profile Settings, Safety & Privacy, Notifications, Language, Offline Mode, Accessibility, Data Export, Report, SOS, FAQ & Support'),
    ('Operational and roadmap', 'Host Dashboard, Admin Panel and all registered Phase 2/3 roadmap destinations'),
]:
    bullet(doc, f'{group}: {screens}.')

heading(doc, '4. Screens Updated', 1)
add_para(doc, 'No production screen implementation required a Stage 3 code change after the final audit. Existing Stage 1 and Stage 2 presentation work was preserved to minimize regression risk and comply with the prohibition on behavioral redesign.')
bullet(doc, 'Updated QA surface: test/widget_test.dart now includes all registered Phase 2/3 routes and the actual /video-speed-dating-room path.')
bullet(doc, 'Removed false-positive route coverage: the previous /video-speed-dating test value exercised onUnknownRoute instead of the intended screen.')
bullet(doc, 'Production screen files modified during Stage 3: none.')

heading(doc, '5. UI Improvements', 1)
add_para(doc, 'Stage 3 validated the completed visual migration rather than duplicating earlier screen-level redesign. The retained frontend consistently uses the AMORA semantic palette, shared spacing/radius/elevation foundations, Material 3 theming, premium card treatments, branded badges, shared search/input patterns and responsive navigation.')
bullet(doc, 'Text-only chat UI remains free of voice, video, camera, attachment, audio-recording and shared-media controls.')
bullet(doc, 'Premium and verified profile presentation, zodiac chip, event/premium cards, settings groupings and shared semantic icons remain intact.')
bullet(doc, 'No missing Material icons were detected; the release build tree-shook the bundled font by 97.4%.')

heading(doc, '6. Responsive Improvements', 1)
add_para(doc, 'The responsive route matrix renders every registered route at all requested widths: 320, 360, 375, 390, 412, 414, 768 and 1024dp. Tests assert that screen entry and route exit remain exception-free, detecting RenderFlex and layout exceptions surfaced by Flutter.')
bullet(doc, 'All tested widths passed with no captured layout exceptions.')
bullet(doc, 'Chat has dedicated 320dp accessible icon-target coverage and 375dp overflow coverage.')
bullet(doc, 'Landscape, foldables and keyboard-in-motion require physical/emulator interaction testing beyond static widget surface sizing.')

doc.add_page_break()
heading(doc, '7. Navigation Verification', 1)
add_para(doc, 'Navigation verification was strengthened without changing application routes. Every registered destination is constructed through the current MaterialApp route table and exercised by name.')
table(doc, ['CHECK', 'RESULT', 'EVIDENCE'], [
    ('Core SOW destinations', 'PASS', '32 of 32 corresponding screens/routes present'),
    ('Registered route construction', 'PASS', '88 route names build without exceptions'),
    ('Responsive route entry/exit', 'PASS', '88 routes x 8 widths'),
    ('Unknown-route masking issue', 'FIXED IN TEST', 'Corrected /video-speed-dating-room and added omitted routes'),
    ('Route additions/removals', 'NONE', 'Production route table unchanged'),
    ('Physical tap-through', 'NOT EXECUTED', 'Requires device/emulator and backend-ready fixtures'),
], [2500, 1600, 5260], status_column=1)

heading(doc, '8. Accessibility Improvements', 1)
bullet(doc, 'Shared AppPrimaryButton and IconButton controls meet the 48dp minimum touch-target test.')
bullet(doc, 'Semantic Material 3 color mapping is covered by design-system tests.')
bullet(doc, 'Chat actions expose tooltips and maintain accessible targets at 320dp.')
bullet(doc, 'Large-width responsive tests passed, but TalkBack announcements, focus order, dynamic text at extreme scaling and color-contrast instrumentation remain physical/manual audit items.')

heading(doc, '9. Performance Improvements', 1)
bullet(doc, 'Release compilation completed with icon tree-shaking; MaterialIcons-Regular.otf reduced from 1,645,184 bytes to 42,232 bytes.')
bullet(doc, 'Route tests identify construction and layout regressions early across the complete navigation surface.')
bullet(doc, 'No business logic or state architecture was changed in pursuit of micro-optimizations, avoiding release risk.')
bullet(doc, 'APK size is the primary measured performance gap: 76.77 MiB versus the SOW target below 30 MB. Produce ABI-split APKs or an Android App Bundle and inspect asset/dependency composition before store submission.')

heading(doc, '10. Design System Compliance', 1)
add_para(doc, 'Automated tests confirm that semantic AMORA colors map into Material 3, shared controls enforce touch sizing, Material icon assets are enabled, and no undeclared Cupertino or raw IconData fonts are used. A source scan found no direct raw Colors.* usage in feature code outside AppColors references.')
callout(doc, 'Remaining token-debt note', 'Static scanning still finds local TextStyle, EdgeInsets, BorderRadius, BoxShadow and primitive Card/TextField constructors in feature code. Many wrap or derive shared tokens, but the Stage 1 rule “everything only from shared components” is not mechanically enforceable today. Add custom lint rules and migrate remaining primitives in a separate non-behavioral cleanup before claiming literal 100% token compliance.', fill='FFF7ED', accent=AMBER)

heading(doc, '11. QA Results', 1)
table(doc, ['QA ACTIVITY', 'RESULT', 'DETAIL'], [
    ('SOW document review', 'PASS', '43 of 43 pages extracted, read and visually inspected'),
    ('Core screen coverage', 'PASS', '32 of 32 SOW screens mapped'),
    ('Registered routes', 'PASS', '88 route names build at 8 widths'),
    ('Text chat controls', 'PASS', 'Text send works; unsupported call/media UI absent'),
    ('Design-system tests', 'PASS', 'Semantic colors and accessible controls'),
    ('Icon-system tests', 'PASS', 'Material asset and semantic icon validation'),
    ('Physical Android QA', 'NOT EXECUTED', 'No attached physical Android device available'),
], [2500, 1600, 5260], status_column=1)

doc.add_page_break()
heading(doc, '12. flutter analyze Result', 1)
callout(doc, 'PASS - 0 issues', 'Command: flutter analyze\nResult: No issues found (completed in 3.1 seconds after clean dependency resolution).', fill='ECFDF3', accent=GREEN)

heading(doc, '13. flutter test Result', 1)
callout(doc, 'PASS - 11 of 11 tests', 'Command: flutter test\nResult: All tests passed in 65.5 seconds. Coverage includes launch smoke, 88-route construction, eight responsive widths, text-only chat send behavior, chat icon layout, semantic colors, touch targets and icon-system integrity.', fill='ECFDF3', accent=GREEN)

heading(doc, '14. APK Build Result', 1)
table(doc, ['ATTRIBUTE', 'VALUE'], [
    ('Command', 'flutter build apk --release'),
    ('Result', 'PASS - release APK built successfully'),
    ('Build duration', '107.6 seconds total; Gradle assembleRelease 104.0 seconds'),
    ('Artifact', 'build/app/outputs/flutter-apk/app-release.apk'),
    ('Size', '80,494,089 bytes (76.77 MiB; Flutter reported 76.8 MB)'),
    ('SHA-256', 'DCEB0093B80AF9D77E427DA77C9E0A03533FAFA13E524969AAEDF6A4582D9452'),
    ('SOW size target', 'PARTIAL - exceeds the <30 MB APK/IPA target'),
], [2500, 6860])

heading(doc, '15. Remaining Frontend Improvements', 1)
table(doc, ['PRIORITY', 'REMAINING GATE', 'RECOMMENDED ACTION'], [
    ('P0', 'Physical Android device QA', 'Validate gesture navigation, system bars, keyboard, scroll, touch feedback, images and animations on the target device matrix.'),
    ('P0', 'APK size target', 'Use app bundle/ABI splits, audit asset payloads and dependencies, then remeasure against the SOW target.'),
    ('P1', 'Manual accessibility audit', 'Run TalkBack, focus-order, contrast and extreme text-scale checks on primary journeys.'),
    ('P1', 'Integration/UAT evidence', 'Add integration tests for full match, chat, subscribe, event and referral flows in a backend-ready staging environment.'),
    ('P2', 'Token enforcement', 'Introduce static lint rules and migrate remaining local style/spacing/radius constructors to shared tokens.'),
    ('P2', 'Performance profiling', 'Capture DevTools cold-start, frame pacing, memory and navigation transition metrics on representative hardware.'),
], [900, 3000, 5460])

heading(doc, '16. Final Frontend Readiness Score', 1)
callout(doc, '92% - CONDITIONAL RELEASE CANDIDATE', 'The frontend satisfies the SOW core screen-presence baseline, compiles cleanly, passes all current automated tests, renders every registered route at the required widths, and produces a release APK. The eight-point holdback reflects unexecuted physical-device/accessibility/performance validation and the measured APK-size miss. Release approval should follow closure of the two P0 gates.', fill='FFF7ED', accent=AMBER)

heading(doc, 'Readiness Score Basis', 2)
table(doc, ['AREA', 'WEIGHT', 'SCORE', 'RATIONALE'], [
    ('SOW screen coverage', '20', '20', '32 of 32 core screens mapped'),
    ('Build and analyzer health', '20', '20', '0 analyzer issues; release build passes'),
    ('Automated UI/navigation QA', '20', '20', '11 tests; 88 routes x 8 widths'),
    ('Design-system/accessibility', '15', '13', 'Core tests pass; manual audit and token lint remain'),
    ('Performance/package readiness', '15', '10', 'Build passes; APK exceeds SOW size target'),
    ('Physical-device/UAT readiness', '10', '9', 'Automated surfaces pass; physical execution pending'),
    ('TOTAL', '100', '92', 'Conditional release candidate'),
], [2800, 1000, 1000, 4560])

heading(doc, 'Change Control Confirmation', 2)
add_para(doc, 'Stage 3 changed only test/widget_test.dart to close route-coverage gaps. No production frontend screen, route, backend interface, business logic, service, repository, provider, controller, model, authentication, payment, AI, state-management or navigation behavior was changed.')

doc.core_properties.title = 'AMORA AI Stage 3 Final Frontend Launch Readiness Report'
doc.core_properties.subject = 'SOW verification, frontend QA and release readiness'
doc.core_properties.author = 'AMORA AI Frontend QA'
doc.core_properties.keywords = 'AMORA AI, Flutter, SOW, QA, launch readiness'
doc.save(OUT)
print(OUT.resolve())
