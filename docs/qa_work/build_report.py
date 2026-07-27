from __future__ import annotations

import json
import math
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "docs"
WORK = DOCS / "qa_work"
EVIDENCE = DOCS / "qa_evidence"
OUTPUT = DOCS / "Amora_QA_Testing_Report.docx"
SUMMARY = DOCS / "Amora_QA_Testing_Report_Summary.txt"
DATA = json.loads((WORK / "qa_dataset.json").read_text(encoding="utf-8"))
META = DATA["metadata"]

PRIMARY = "3D0B3F"
SECONDARY = "EC5FA8"
TERTIARY = "F4A9CE"
BACKGROUND = "FDF1F7"
SURFACE = "FFFFFF"
TEXT = "2B2B2B"
MUTED = "6D5A6E"
FONT = "Aptos"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name: str = FONT, size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=65, start=80, bottom=65, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str = TEXT,
    size: float = 8.5,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size)
    run.bold = bold
    run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color: str = TERTIARY, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths: list[float] | None = None,
    font_size: float = 8.2,
    repeat_header: bool = True,
) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, color=SURFACE, size=font_size)
        shade_cell(cell, PRIMARY)
        if widths:
            cell.width = Inches(widths[index])
    if repeat_header:
        set_repeat_header(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            set_cell_text(cell, value, size=font_size)
            if row_index % 2 == 0:
                shade_cell(cell, BACKGROUND)
            if widths:
                cell.width = Inches(widths[col_index])
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, instruction: str, display: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8)
    add_field(paragraph, "PAGE", "1")
    run = paragraph.add_run(" of ")
    set_run_font(run, size=8)
    add_field(paragraph, "NUMPAGES", "1")


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        first.bold = True
        set_run_font(first, size=10.2)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=10.2)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.2)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2.5)
        paragraph.paragraph_format.line_spacing = 1.12
        run = paragraph.add_run(item)
        set_run_font(run, size=9.8)


def add_callout(doc: Document, title: str, body: str, *, fill=BACKGROUND) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_run_font(run, size=11)
    run.bold = True
    run.font.color.rgb = rgb(PRIMARY)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(body)
    set_run_font(run, size=9.6)
    run.font.color.rgb = rgb(TEXT)
    set_table_borders(table, color=TERTIARY, size=6)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, image: str, caption: str, width: float) -> None:
    path = EVIDENCE / image
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=8.5)
    caption_run.italic = True
    caption_run.font.color.rgb = rgb(MUTED)


def configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.25)


def apply_headers_and_footers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header = section.header
        if header.paragraphs:
            header.paragraphs[0].text = ""
        table = header.add_table(rows=1, cols=2, width=Inches(9.6 if section.orientation == WD_ORIENT.LANDSCAPE else 6.7))
        table.autofit = False
        set_cell_text(table.cell(0, 0), "AMORA • QA TESTING REPORT", bold=True, color=PRIMARY, size=8)
        set_cell_text(
            table.cell(0, 1),
            "RELEASE DECISION: NOT READY",
            bold=True,
            color=PRIMARY,
            size=8,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        set_table_borders(table, color=TERTIARY, size=4)
        footer = section.footer
        footer.paragraphs[0].text = ""
        footer_table = footer.add_table(
            rows=1,
            cols=2,
            width=Inches(9.6 if section.orientation == WD_ORIENT.LANDSCAPE else 6.7),
        )
        set_cell_text(
            footer_table.cell(0, 0),
            "Internal QA / Submission • 27 July 2026",
            color=MUTED,
            size=8,
        )
        add_page_number(footer_table.cell(0, 1).paragraphs[0])
        set_table_borders(footer_table, color=TERTIARY, size=4)
        if index == 0:
            section.different_first_page_header_footer = True
            cover_footer = section.first_page_footer
            cover_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cover_footer.paragraphs[0].add_run(
                "AMORA AI • INTERNAL QA / SUBMISSION"
            )
            set_run_font(run, size=8)
            run.font.color.rgb = rgb(MUTED)


def module_count_row(module: str) -> list[str]:
    counts = DATA["module_summary"].get(module, {})
    return [
        module,
        str(sum(counts.values())),
        str(counts.get("Pass", 0)),
        str(counts.get("Fail", 0)),
        str(counts.get("Blocked", 0)),
        str(counts.get("Not Run", 0)),
        str(counts.get("Not Applicable", 0)),
    ]


doc = Document()
configure_section(doc.sections[0])
doc.core_properties.title = META["title"]
doc.core_properties.subject = "Comprehensive frontend QA assessment"
doc.core_properties.author = "Amora QA"
doc.core_properties.keywords = "Flutter, QA, Amora, testing, release readiness"
doc.core_properties.comments = "Generated from observed project and command evidence."

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.2)
normal.font.color.rgb = rgb(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.2
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

for level, size in ((1, 18), (2, 14), (3, 11.5)):
    style = styles[f"Heading {level}"]
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(PRIMARY if level < 3 else TEXT)
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    style.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    style.paragraph_format.keep_with_next = True
styles["Heading 1"].paragraph_format.page_break_before = True
styles["Caption"].font.name = FONT
styles["Caption"].font.color.rgb = rgb(MUTED)
styles["Caption"]._element.rPr.rFonts.set(qn("w:ascii"), FONT)
styles["Caption"]._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)

settings = doc.settings._element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

# Editorial cover.
band = doc.add_table(rows=1, cols=1)
band.autofit = False
band.cell(0, 0).width = Inches(6.7)
shade_cell(band.cell(0, 0), PRIMARY)
set_cell_text(
    band.cell(0, 0),
    "AMORA AI",
    bold=True,
    color=SURFACE,
    size=13,
)
band.rows[0].height = Inches(0.72)
band.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
doc.add_paragraph().paragraph_format.space_after = Pt(34)

eyebrow = doc.add_paragraph()
eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = eyebrow.add_run("QUALITY ASSURANCE • SUBMISSION REPORT")
set_run_font(run, size=10)
run.bold = True
run.font.color.rgb = rgb(SECONDARY)
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(8)
run = title.add_run("Amora QA\nTesting Report")
set_run_font(run, size=32)
run.bold = True
run.font.color.rgb = rgb(PRIMARY)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(24)
run = subtitle.add_run(
    "Comprehensive frontend verification, automation review, "
    "build evidence and production release assessment"
)
set_run_font(run, size=14)
run.font.color.rgb = rgb(TEXT)

status_table = doc.add_table(rows=1, cols=2)
status_table.autofit = False
set_cell_text(status_table.cell(0, 0), "RELEASE DECISION", bold=True, color=PRIMARY, size=9)
shade_cell(status_table.cell(0, 0), TERTIARY)
set_cell_text(status_table.cell(0, 1), "NOT READY", bold=True, color=SURFACE, size=13)
shade_cell(status_table.cell(0, 1), PRIMARY)
status_table.cell(0, 0).width = Inches(2.2)
status_table.cell(0, 1).width = Inches(4.5)
set_table_borders(status_table, color=SECONDARY, size=8)
doc.add_paragraph().paragraph_format.space_after = Pt(18)

add_table(
    doc,
    ["Document", "Value"],
    [
        ["Project", "Amora AI Flutter Application"],
        ["Report version", "1.0"],
        ["Assessment date", "27 July 2026"],
        ["Prepared for", "Submission and production-readiness review"],
        ["Evidence basis", "Inspected source, automated tests, builds and rendered widgets"],
        ["Classification", "Internal QA / Submission"],
    ],
    widths=[1.75, 4.95],
    font_size=9.2,
)

add_callout(
    doc,
    "Outcome at a glance",
    "Static analysis, web compilation and the Android debug APK succeeded. "
    "The full test suite did not pass, the project is frontend/local-data led, "
    "and production auth, backend, payment, notification and release-signing "
    "readiness are not established.",
)
doc.add_page_break()

# Document control and revision.
add_heading(doc, "Document Control", 1)
add_table(
    doc,
    ["Control", "Detail"],
    [
        ["Document ID", "AMORA-QA-2026-07-27"],
        ["Owner", "Amora QA"],
        ["Version", "1.0"],
        ["Status", "Final assessment artifact"],
        ["Review cadence", "Update after blocking defects and integrations are resolved"],
        ["Source boundary", r"D:\Projects\amora_ai (no Git/GitHub used)"],
    ],
    widths=[1.8, 4.9],
)
add_heading(doc, "Revision History", 2)
add_table(
    doc,
    ["Version", "Date", "Author", "Change"],
    [["1.0", "27 Jul 2026", "Amora QA", "Initial comprehensive QA submission report"]],
    widths=[0.8, 1.2, 1.4, 3.3],
)
add_heading(doc, "Approval Record", 2)
add_table(
    doc,
    ["Role", "Name", "Decision", "Date"],
    [
        ["QA", "Amora QA", "Assessment completed", "27 Jul 2026"],
        ["Engineering", "Pending", "Open defects require review", "—"],
        ["Product / Release", "Pending", "Release decision is Not Ready", "—"],
    ],
    widths=[1.35, 1.55, 2.7, 1.1],
)
doc.add_page_break()

add_heading(doc, "Table of Contents", 1)
toc_paragraph = doc.add_paragraph()
add_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u', "")
add_body(
    doc,
    "This document contains an automatic Word TOC field covering Heading "
    "levels 1–3. Word updates page references when fields are refreshed.",
)
add_heading(doc, "Contents Guide", 2)
section_titles = [
    "1 Executive Summary",
    "2 Project Understanding",
    "3 Testing Scope",
    "4 Test Strategy",
    "5 Test Environment",
    "6 Requirement Traceability Matrix",
    "7 Test Case Design",
    "8 Authentication Testing",
    "9 Onboarding Testing",
    "10 Discover Page Testing",
    "11 Filters Testing",
    "12 Profile Detail Testing",
    "13 User Profile Testing",
    "14 Edit Profile Testing",
    "15 Profile Preview Testing",
    "16 AI Matches Testing",
    "17 Chat Testing",
    "18 Events Page Testing",
    "19 Event Details Testing",
    "20 Notifications Testing",
    "21 Subscription Testing",
    "22 Payment Testing",
    "23 FAQ & Support Testing",
    "24 Settings and Account Testing",
    "25 UI and Visual Testing",
    "26 Responsive Testing",
    "27 Accessibility Testing",
    "28 Performance Testing",
    "29 Security and Privacy Validation",
    "30 Compatibility Testing",
    "31 Build and Static Analysis",
    "32 Automation Test Review",
    "33 Defect Management",
    "34 Test Execution Summary",
    "35 Risk Assessment",
    "36 Release Readiness",
    "37 Recommendations",
    "38 Final Conclusion",
    "39 Appendices",
]
half = math.ceil(len(section_titles) / 2)
contents_rows = []
for index in range(half):
    left = section_titles[index]
    right = section_titles[index + half] if index + half < len(section_titles) else ""
    contents_rows.append([left, right])
add_table(doc, ["Core assessment", "Quality and evidence"], contents_rows, widths=[3.35, 3.35], font_size=8.5)

# 1 Executive Summary.
add_heading(doc, "1. Executive Summary", 1)
add_callout(
    doc,
    "Release decision — NOT READY",
    "The application compiles for web and Android debug and has meaningful "
    "widget coverage, but production release criteria are not met. The full "
    "suite reports 23 failures, eight confirmed defects remain open, the "
    "Android release configuration uses template identity/debug signing, and "
    "core product flows are local simulations without production services.",
    fill=TERTIARY,
)
add_table(
    doc,
    ["Metric", "Observed", "Interpretation"],
    [
        ["Source reviewed", "137 Dart files / 49,921 lines", "Large presentation-led prototype"],
        ["Catalogue", "180 cases", "125 pass; 26 fail; 10 blocked; 16 not run; 3 N/A"],
        ["Automated suite", "143 invocations", "115 pass; 23 fail; 5 skipped"],
        ["Analyzer", "Pass", "No analyzer issues"],
        ["Web build", "Pass", "68,034,713-byte output"],
        ["Debug APK", "Pass", "172,299,833 bytes; minSdk 24 / targetSdk 36"],
        ["Confirmed defects", "8 open", "Includes two P1 release/security blockers"],
        ["Backend readiness", "Not established", "No production backend/network SDK declared"],
    ],
    widths=[1.5, 2.0, 3.2],
)
add_heading(doc, "1.1 Major strengths", 2)
add_bullets(
    doc,
    [
        "Feature-oriented Flutter presentation structure with a centralized design system.",
        "Strong widget coverage for current Discover, Events, FAQ Support, Chats, Profile, Notifications and onboarding flows.",
        "Analyzer-clean source, successful web build and successful Android debug packaging.",
        "Local assets are bundled and testable; Events uses five distinct local images.",
        "FAQ Support enforces the intended email-only direct support presentation.",
    ],
)
add_heading(doc, "1.2 Primary limitations", 2)
add_bullets(
    doc,
    [
        "Authentication, data, chat, RSVP, subscription and notifications are frontend/local-state simulations.",
        "No integration_test suite, coverage baseline, physical Android execution or assistive-technology session was available.",
        "The complete suite fails; obsolete regression contracts coexist with genuine UI/runtime defects.",
        "Release package identity/signing is not production safe and a release APK was intentionally not produced.",
    ],
)

# 2 Project Understanding.
add_heading(doc, "2. Project Understanding", 1)
add_body(
    doc,
    "Amora AI is a Flutter dating-product prototype with named-route navigation, "
    "a five-destination main shell, shared Amora design primitives and extensive "
    "presentation screens. The inspected state is not a production-connected "
    "application: profiles, matches, chats, events, onboarding and session state "
    "are driven by local repositories, dummy datasets or in-memory notifiers.",
)
add_heading(doc, "2.1 Architecture observed", 2)
add_table(
    doc,
    ["Layer / concern", "Observed implementation", "QA implication"],
    [
        ["App shell", "MaterialApp named routes and MainShell IndexedStack", "Route-build and tab-state tests applicable"],
        ["Presentation", "36 feature directories; 66 screen files", "Broad UI surface, high regression cost"],
        ["Shared UI", "Core theme, spacing, widgets, navigation", "Semantic tests can detect cross-app drift"],
        ["State", "ChangeNotifier/local ValueNotifier patterns", "Deterministic widget tests; no persistence guarantee"],
        ["Data", "Dummy/static data and asset repositories", "Frontend behavior only; no server correctness claim"],
        ["External integration", "url_launcher for email", "Native client behavior requires device test"],
        ["Backend/payment", "No production SDK declared", "End-to-end cases blocked"],
    ],
    widths=[1.25, 2.75, 2.7],
)
add_heading(doc, "2.2 Product surface reviewed", 2)
add_bullets(
    doc,
    [
        "Auth and profile onboarding, Discover and filters, profile detail and identity editing.",
        "AI Matches, chat inbox/detail, Notifications, Events/detail/My Events.",
        "Subscription/payment test isolation, FAQ email support, settings, safety and privacy presentation.",
        "Cross-cutting navigation, responsive behavior, accessibility-oriented assertions and build configuration.",
    ],
)

# 3 Scope.
add_heading(doc, "3. Testing Scope", 1)
add_heading(doc, "3.1 In scope", 2)
add_bullets(
    doc,
    [
        "Complete Flutter frontend source, routes, assets, tests and Android/web build configuration.",
        "Non-invasive static security/privacy review and dependency inventory.",
        "Automated unit/source and widget tests, including responsive golden evidence.",
        "Buildability for Chrome web and Android debug APK.",
        "Test-case catalogue, RTM, defect log, risk register and release recommendation.",
    ],
)
add_heading(doc, "3.2 Out of scope or unavailable", 2)
add_bullets(
    doc,
    [
        "Production backend, APIs, database rules, Firebase, socket transport or payment processing.",
        "Real OTP/SMS/email delivery, real financial transactions and production credentials.",
        "Penetration testing, destructive security activity and production-data modification.",
        "Physical Android/iOS device testing, store signing and stakeholder UAT.",
    ],
)
add_heading(doc, "3.3 Assumptions", 2)
add_bullets(
    doc,
    [
        "Local fixtures are intentionally fictional prototype data.",
        "AMORA_MEMBERSHIP_TEST must remain disabled unless the isolated test journey is explicitly requested.",
        "Named route strings are treated as compatibility contracts even where screens are presentation-only.",
        "A build success does not imply production service readiness.",
    ],
)

# 4 Strategy.
add_heading(doc, "4. Test Strategy", 1)
strategy_rows = [
    ["Code/source review", "Performed", "Routes, dependencies, local/mock boundaries, assets, signing and sensitive behavior"],
    ["Static analysis", "Performed", "flutter analyze"],
    ["Unit/source tests", "Performed", "Controller, design-system and repository calculations"],
    ["Widget tests", "Performed", "Complete test folder; 143 invocations"],
    ["Responsive rendering", "Performed (targeted)", "Golden evidence at 390x844 and Events at 1280x900"],
    ["Regression", "Performed", "Full suite; obsolete tests and genuine defects separated"],
    ["Smoke/build", "Performed", "Analyzer, web build, Chrome launch and debug APK"],
    ["Manual Chrome journey", "Partially performed", "Debug service connected; browser control isolation prevented a completed manual pass"],
    ["Integration/device", "Not performed", "No integration_test suite or Android device"],
    ["Accessibility manual", "Not performed", "No screen reader session"],
    ["Performance profiling", "Not performed", "No profiling-mode timing/memory evidence"],
    ["Security validation", "Review only", "No penetration or destructive testing"],
    ["UAT", "Not performed", "Requires stakeholders"],
]
add_table(doc, ["Test type", "Disposition", "Evidence / boundary"], strategy_rows, widths=[1.65, 1.45, 3.6])
add_heading(doc, "4.1 Entry and exit criteria", 2)
add_bullets(
    doc,
    [
        "Entry: dependencies resolve, source parses and requested build targets are available.",
        "Exit for this assessment: evidence captured, all commands attempted honestly, defects/risks traced and artifacts validated.",
        "Production exit is not met while P1 defects, failing tests and integration blockers remain.",
    ],
)

# 5 Environment.
add_heading(doc, "5. Test Environment", 1)
environment_rows = [
    ["Operating system", "Windows 11 Home Single Language 64-bit, 25H2"],
    ["Flutter", "3.44.6 stable"],
    ["Dart / DevTools", "3.12.2 / 2.57.0"],
    ["Android SDK", "36.1.0; platform/build-tools 36.1.0"],
    ["Built APK SDK", "minSdk 24; targetSdk 36"],
    ["Java", "OpenJDK 21.0.8 via Android Studio; shell Java 22.0.1"],
    ["Chrome", "150.0.7871.182"],
    ["Connected targets", "Windows, Chrome, Edge; no Android device/emulator"],
    ["Build modes", "Web release bundle; Android debug APK"],
    ["Internet", "Flutter doctor network resources passed"],
    ["Backend environment", "None declared; frontend/local fixtures"],
    ["IDE/tooling", "Android Studio detected; Codex desktop QA workflow"],
]
add_table(doc, ["Property", "Observed value"], environment_rows, widths=[1.85, 4.85])
add_callout(
    doc,
    "Environment limitation",
    "Windows desktop doctor reports an incomplete C++ workload. This did not "
    "block requested Android/web work, but Windows release support is not validated.",
)

# 6 RTM.
add_heading(doc, "6. Requirement Traceability Matrix", 1)
add_body(
    doc,
    "Twenty consolidated requirements were derived from the QA brief and mapped "
    "to current source/test evidence. The companion workbook contains the full "
    "editable RTM; Appendix B reproduces it in this report.",
)
add_table(
    doc,
    ["Requirement", "Module", "Implementation", "Test status"],
    [
        [
            item["id"],
            item["module"],
            item["implementation"],
            item["test_status"] or "No direct case; covered by broader review",
        ]
        for item in DATA["requirements"]
    ],
    widths=[1.25, 1.55, 1.7, 2.2],
    font_size=7.8,
)
add_heading(doc, "6.1 Traceability interpretation", 2)
add_bullets(
    doc,
    [
        "Implemented locally means presentation/local-state behavior exists; it does not imply server integration.",
        "Frontend fixture identifies screens backed by static or dummy data.",
        "Not production ready denotes security/integration requirements with blocking evidence.",
    ],
)

# 7 Test design.
add_heading(doc, "7. Test Case Design", 1)
add_table(
    doc,
    ["Field", "Definition"],
    [
        ["ID", "Stable AUT/BLD/REV/MAN/NA identifier"],
        ["Module and scenario", "Product surface and behavior under assessment"],
        ["Objective and preconditions", "Why the check exists and what must be true first"],
        ["Test data", "Safe fictional/local data only"],
        ["Execution steps", "Reproducible command or interaction sequence"],
        ["Expected / actual", "Desired contract versus observed evidence"],
        ["Status", "Pass, Fail, Blocked, Not Run or Not Applicable"],
        ["Priority / severity", "Business importance and P0–P4 impact if failed"],
        ["Environment / evidence / remarks", "Trace to the host, artifact and interpretation"],
    ],
    widths=[1.7, 5.0],
)
add_heading(doc, "7.1 Catalogue accounting", 2)
add_table(
    doc,
    ["Status", "Count", "Rule"],
    [
        ["Pass", "125", "Observed expected result"],
        ["Fail", "26", "Observed deviation; not all map one-to-one to product defects"],
        ["Blocked", "10", "External service/device unavailable"],
        ["Not Run", "16", "No execution claim"],
        ["Not Applicable", "3", "Platform/signing precondition intentionally absent"],
        ["Total", "180", "Formula-backed in companion workbook"],
    ],
    widths=[1.5, 1.0, 4.2],
)

module_sections = [
    (
        8,
        "Authentication Testing",
        "Authentication",
        [
            "App launch and unauthenticated routing",
            "Login, signup, phone/country code and password visibility",
            "Six-cell OTP input, paste/backspace, empty/incomplete/invalid states",
            "Loading, button state, keyboard inset, logout and guarded actions",
        ],
        "Current auth widgets are responsive and well tested, but verification and session state are local simulations. A demo OTP is exposed in the UI.",
        "EV-UI-001_auth_mobile.png",
        "Figure 1: Authentication Entry — Mobile Widget Render",
    ),
    (
        9,
        "Onboarding Testing",
        "Onboarding",
        [
            "Name, gender, interests, relationship goal and Gujarat city catalogue",
            "Birth-date day/month/year wheels, keyboard input and minimum-age boundary",
            "Progress, back/resume behavior, completion and returning users",
            "Photo/prompt/permission behavior where implemented",
        ],
        "Age calculation and minimum-age inline validation have dedicated passing tests. Three legacy end-to-end widget locators are ambiguous after redesign.",
        None,
        None,
    ),
    (
        10,
        "Discover Page Testing",
        "Discover",
        [
            "Single profile card, multiple local images and gallery progress",
            "Horizontal swipes, thresholds, cancellation and vertical-gesture rejection",
            "Like, pass, undo, Super Like, filters, empty/restart and detail route",
            "Five-tab bottom navigation, mobile/desktop constraints and keyboard shortcuts",
        ],
        "The current Discover suite passes broad gesture and responsive assertions. A separate obsolete four-tab suite causes fifteen failures and must be retired or updated.",
        "EV-UI-002_discover_mobile.png",
        "Figure 2: Discover Page — Mobile Widget Render",
    ),
    (
        11,
        "Filters Testing",
        "Filters",
        [
            "Open/close, search, age/distance and supported preference sections",
            "Selection, reset, apply, active indicator and route return",
            "Compact and desktop overflow checks",
        ],
        "All four focused filter tests pass. Server-side result correctness is not testable because filtering uses local profile data.",
        None,
        None,
    ),
    (
        12,
        "Profile Detail Testing",
        "Profile Detail",
        [
            "Gallery de-duplication, swipe, full-screen zoom and image tap",
            "Name, age, location, occupation, zodiac, prompts, interests and intentions",
            "Like, message, Super Like, report/block and sticky actions",
            "Compact and desktop presentation",
        ],
        "Focused profile-detail tests pass. Actions route or mutate local callbacks only; server persistence and safety escalation are blocked.",
        None,
        None,
    ),
    (
        13,
        "User Profile Testing",
        "User Profile",
        [
            "Identity-first hero, photo, name/age, zodiac, city and occupation",
            "Completion, missing items, gallery, prompts, voice intro and About Me",
            "Dating intentions, interests, personality and verification presentation",
            "Membership, settings, legal, logout and delete-account exposure",
        ],
        "The identity screen passes compact/desktop and local-save tests. Profile data is a singleton local draft and is not persisted across a process restart.",
        "EV-UI-005_profile_mobile.png",
        "Figure 3: Premium Profile Identity — Mobile Widget Render",
    ),
    (
        14,
        "Edit Profile Testing",
        "User Profile",
        [
            "Existing values, basic data, occupation/city/education and lifestyle",
            "Prompts, photos, interests, intentions and local save",
            "Validation, cancel/unsaved warning, double save and preview update",
        ],
        "Local save behavior is covered by passing profile tests. Network error, remote persistence, media upload and true concurrent save behavior are blocked.",
        None,
        None,
    ),
    (
        15,
        "Profile Preview Testing",
        "User Profile",
        [
            "Public-profile photo order and identity data",
            "Prompts, interests, voice, intentions, verification and About",
            "No Like, Message, Super Like or Gift actions",
            "Responsive rendering and return to editor",
        ],
        "Preview routing is exercised by regression/profile tests, but no dedicated device-level golden journey exists.",
        None,
        None,
    ),
    (
        16,
        "AI Matches Testing",
        "AI Matches",
        [
            "Featured match, compatibility score and available reasons",
            "Shared interests, filtering, profile/message routes and empty/error states",
            "Static/mock identification and unsupported-claim review",
            "Compact/desktop performance-oriented layout",
        ],
        "Six focused tests pass. The feature is fixture-driven and must not be represented as a validated AI recommendation service.",
        None,
        None,
    ),
    (
        17,
        "Chat Testing",
        "Chats",
        [
            "Conversation list, active users, search, unread counts and ordering",
            "Conversation route, text send, empty/long text and emoji",
            "Read/typing/online presentation and supported management actions",
            "Keyboard, scrolling, accessible icon targets and desktop layout",
        ],
        "Nine focused frontend tests pass and two real-delivery cases are blocked. No socket/Firebase/API transport or attachment backend is declared.",
        None,
        None,
    ),
    (
        18,
        "Events Page Testing",
        "Events",
        [
            "Direct Events access without subscription/payment redirect",
            "Featured event, category filters, recommended/This Week/Near You/Circles",
            "My Events, local unique images, host/attendee presentation and RSVP states",
            "Loading, empty/error, one bottom navigation and responsive layout",
        ],
        "Current Events tests pass direct access, local assets and desktop layout. A 35 px overflow occurs at 320x700 with 1.3x text scaling; server-backed RSVP is blocked.",
        "EV-UI-003_events_mobile.png",
        "Figure 4: Events Discovery — Mobile Widget Render",
    ),
    (
        19,
        "Event Details Testing",
        "Events",
        [
            "Immersive hero, back/save/share, identity and category",
            "Date/time, location, host, description and expectation content",
            "Attendee/safety presentation, join/leave/waitlist state and sticky CTA",
            "SafeArea, transition and desktop rendering",
        ],
        "Focused detail tests pass. The separate waitlist route fails Flutter's Material/ListTile contract and all RSVP state is local.",
        "EV-UI-004_event_detail_mobile.png",
        "Figure 5: Event Detail — Mobile Widget Render",
    ),
    (
        20,
        "Notifications Testing",
        "Notifications",
        [
            "Timeline, grouping, read/unread filters and badge presentation",
            "Swipe read/delete, long-press selection and related-route navigation",
            "Empty/error and desktop layout",
            "Deep-link and missing-target behavior where available",
        ],
        "Seven focused presentation tests pass. Push delivery and external deep links are blocked because no notification provider is configured.",
        None,
        None,
    ),
    (
        21,
        "Subscription Testing",
        "Subscription & Payment",
        [
            "Plan presentation, monthly/annual selection and benefits",
            "Current membership, upgrade/manage/restore where supported",
            "Events not gated by subscription",
            "Test-mode isolation and unavailable-plan/error behavior",
        ],
        "Five membership journey invocations are intentionally skipped with AMORA_MEMBERSHIP_TEST disabled. No production entitlement behavior was claimed.",
        None,
        None,
    ),
    (
        22,
        "Payment Testing",
        "Subscription & Payment",
        [
            "Selected plan/amount, simulated success/failure/cancel/pending",
            "Loading, retry, duplicate-tap prevention and confirmation",
            "Sensitive-data handling and external SDK boundary",
        ],
        "The payment screen is an isolated test fixture. No real or sandbox financial transaction was executed and no payment SDK is declared.",
        None,
        None,
    ),
    (
        23,
        "FAQ & Support Testing",
        "FAQ & Support",
        [
            "FAQ opening, local search, categories, accordion and empty results",
            "Email Support address, subject/body generation and launcher callback",
            "Absence of WhatsApp, call, support ticket, status and My Tickets",
            "Safety/legal links and desktop responsiveness",
        ],
        "Five support tests plus source review pass. Native email-client launch remains blocked without an Android device.",
        "EV-UI-006_faq_support_mobile.png",
        "Figure 6: FAQ & Email Support — Mobile Widget Render",
    ),
    (
        24,
        "Settings and Account Testing",
        "Cross-functional",
        [
            "Account, privacy, notifications and blocked users",
            "Safety Center, guidelines, Terms and Privacy Policy",
            "Logout, delete confirmation/cancel and restricted routes",
        ],
        "Routes build through broad tests, but production session clearing, remote deletion and authorization enforcement are unavailable.",
        None,
        None,
    ),
]

for number, title, module, coverage, finding, image, caption in module_sections:
    add_heading(doc, f"{number}. {title}", 1)
    add_heading(doc, f"{number}.1 Coverage assessed", 2)
    add_bullets(doc, coverage)
    add_heading(doc, f"{number}.2 Observed result", 2)
    add_body(doc, finding)
    add_table(
        doc,
        ["Module", "Total", "Pass", "Fail", "Blocked", "Not Run", "N/A"],
        [module_count_row(module)],
        widths=[1.65, 0.72, 0.72, 0.72, 0.82, 0.82, 0.7],
        font_size=8,
    )
    if image:
        add_figure(doc, image, caption, 3.25 if "desktop" not in image else 6.6)
    if number == 18:
        add_figure(
            doc,
            "EV-UI-007_events_desktop.png",
            "Figure 7: Events Discovery — 1280×900 Responsive Widget Render",
            6.6,
        )

# 25 Visual.
add_heading(doc, "25. UI and Visual Testing", 1)
add_table(
    doc,
    ["Area", "Observation", "Disposition"],
    [
        ["Palette", "Central palette matches the six requested colors", "Pass with one contrast defect"],
        ["Typography", "Consistent native/platform text system", "Code-review pass; device font rendering not certified"],
        ["Icons", "Bundled Material icon system", "Focused tests pass"],
        ["Spacing/radius", "Shared tokens plus feature-local values", "Visually coherent in evidence"],
        ["Images", "Local profile/event assets render in golden evidence", "Pass for sampled screens"],
        ["Overflow", "Events scaled-text overflow confirmed", "Fail — BUG-EVT-001"],
        ["Navigation", "Five tabs and single Events bottom bar tested", "Current tests pass; legacy suite stale"],
        ["States", "Loading/empty/error widgets exist across core areas", "Targeted coverage; not every route manually viewed"],
    ],
    widths=[1.15, 3.7, 1.85],
)
add_callout(
    doc,
    "Contrast finding",
    "White text on #EC5FA8 measured 3.10:1 in the project test. This does not "
    "meet the test's 4.5:1 normal-text threshold. Use #3D0B3F text on pink, "
    "or reserve white-on-pink for typography that meets the applicable large-text criterion.",
)

# 26 Responsive.
add_heading(doc, "26. Responsive Testing", 1)
viewport_rows = [
    ["320×568", "Automated route/compact tests; Events overflow at related 320×700 + 1.3x", "Partial / Fail"],
    ["360×640", "Covered indirectly by compact widget ranges", "Partial"],
    ["375×667", "Chat focused widget case", "Pass (targeted)"],
    ["390×844", "Seven evidence renders include core mobile screens", "Pass (sampled)"],
    ["412×915", "Not manually executed", "Not Run"],
    ["600×960", "No dedicated rendered evidence", "Not Run"],
    ["768×1024", "No tablet device", "Not Run"],
    ["1024×768", "No full manual landscape journey", "Not Run"],
    ["1280×900", "Events desktop golden", "Pass (sampled)"],
    ["1366×768", "Chrome process launched; interaction not completed", "Not Run"],
    ["1440×900", "Desktop constraints covered in selected widget tests", "Partial"],
    ["1920×1080", "Not manually executed", "Not Run"],
]
add_table(doc, ["Viewport", "Evidence", "Result"], viewport_rows, widths=[1.1, 4.3, 1.3], font_size=8.2)
add_heading(doc, "26.1 Responsive conclusion", 2)
add_body(
    doc,
    "Responsive confidence is moderate for sampled Flutter widget layouts, "
    "but low for full-device coverage. The confirmed scaled-text Events "
    "overflow prevents a clean responsive sign-off.",
)

# 27 Accessibility.
add_heading(doc, "27. Accessibility Testing", 1)
add_table(
    doc,
    ["Check", "Evidence", "Status"],
    [
        ["Semantic labels/tooltips", "Used throughout core screens; targeted assertions", "Partial pass"],
        ["48 dp touch targets", "Shared controls and chat icon tests", "Pass (targeted)"],
        ["Text scaling", "Targeted tests; Events failure at 1.3x", "Fail / incomplete"],
        ["Contrast", "Automated contrast utility", "Fail for white on secondary"],
        ["Keyboard support", "OTP and Discover shortcuts tested", "Partial pass"],
        ["Focus visibility", "Code review only", "Not fully verified"],
        ["Screen reader", "No TalkBack/NVDA/VoiceOver session", "Not Run"],
        ["Reduced motion", "No comprehensive runtime verification", "Not Run"],
    ],
    widths=[1.55, 3.6, 1.55],
)
add_body(
    doc,
    "These observations are WCAG-oriented engineering checks, not an official "
    "accessibility certification.",
)

# 28 Performance.
add_heading(doc, "28. Performance Testing", 1)
add_table(
    doc,
    ["Signal", "Observed", "Interpretation"],
    [
        ["Web compilation", "81.7 seconds", "Build-time measurement only"],
        ["Debug APK build", "54.5 seconds", "Build-time measurement only"],
        ["Web output size", "68,034,713 bytes", "Large, includes CanvasKit variants"],
        ["Debug APK size", "172,299,833 bytes", "High; measure release split/AAB after signing"],
        ["Asset payload", "24,408,115 bytes / 163 files", "Image optimization opportunity"],
        ["Scrolling/animation", "Widget tests complete without timing metrics", "Functional only"],
        ["CPU/memory/startup", "No profiler evidence", "Not Run"],
        ["Network/offline", "No production network stack", "Blocked"],
    ],
    widths=[1.55, 2.1, 3.05],
)
add_body(
    doc,
    "No 120 FPS, startup-time, frame-time, CPU or memory claim is made. "
    "A profiling-mode run on representative Android hardware is required.",
)

# 29 Security/privacy.
add_heading(doc, "29. Security and Privacy Validation", 1)
add_table(
    doc,
    ["Area", "Observation", "Result"],
    [
        ["Hardcoded secrets", "No API keys/tokens found in the inspected Dart dependency surface", "No confirmed secret"],
        ["OTP", "Locally generated demo OTP is displayed", "Fail — BUG-SEC-001"],
        ["Password fields", "Visibility controls and local validation exist", "Frontend pass"],
        ["Auth state", "In-memory ValueNotifier", "Production blocker"],
        ["Route protection", "Frontend guard sheets/named routing", "Not server authorization"],
        ["Payment data", "No real gateway/credentials", "Blocked / no financial test"],
        ["External links", "url_launcher mailto for support", "Device verification required"],
        ["Delete/logout", "Presentation/local session behavior", "Remote enforcement blocked"],
        ["Private data in evidence", "No passwords, real OTPs, messages, tokens or payment data included", "Pass"],
    ],
    widths=[1.35, 3.8, 1.55],
)
add_callout(
    doc,
    "Security boundary",
    "This was a non-invasive source/configuration review. No penetration "
    "testing, credential probing, backend authorization test or destructive "
    "action was attempted.",
)

# 30 Compatibility.
add_heading(doc, "30. Compatibility Testing", 1)
add_table(
    doc,
    ["Platform/input", "Configuration", "Tested status"],
    [
        ["Android", "minSdk 24; targetSdk 36; debug APK", "Build passed; install not run"],
        ["Chrome Web", "Chrome 150", "Build passed; debug process started; manual pass incomplete"],
        ["Edge Web", "Detected", "Not Run"],
        ["Windows", "Flutter target detected; C++ workload incomplete", "Not Applicable for this release"],
        ["iOS", "Project support theoretical on Windows", "Not Applicable on host"],
        ["Portrait/touch", "Widget test surface", "Targeted pass"],
        ["Landscape/tablet", "No physical device", "Not Run"],
        ["Mouse/keyboard", "Selected web/widget interactions", "Partial"],
        ["Dark mode", "No project-wide dark theme validation", "Not Run"],
    ],
    widths=[1.4, 3.4, 1.9],
)

# 31 Build.
add_heading(doc, "31. Build and Static Analysis", 1)
build_rows = [
    ["flutter clean", "Pass", "3.8 s", "Generated outputs removed"],
    ["flutter pub get", "Pass", "4.5 s", "Four newer incompatible versions noted"],
    ["dart format lib", "Pass", "3.23 s", "137 files; 0 changed"],
    ["flutter analyze", "Pass", "3.0 s", "No issues found"],
    ["flutter test", "Fail", "Full run", "115 pass; 23 fail; 5 skipped"],
    ["flutter build web", "Pass", "81.7 s", r"build\web; 68,034,713 bytes"],
    ["flutter run -d chrome", "Started", "50.3 s", "Debug service connected; interaction incomplete"],
    ["flutter build apk --debug", "Pass", "54.5 s", "app-debug.apk; 172,299,833 bytes"],
    ["flutter build apk --release", "Not attempted", "—", "Debug signing/template ID; unsafe to claim release"],
]
add_table(doc, ["Command", "Result", "Duration", "Output / warning"], build_rows, widths=[2.0, 1.05, 1.0, 2.65], font_size=8.3)
add_heading(doc, "31.1 APK evidence", 2)
add_bullets(
    doc,
    [
        r"Path: D:\Projects\amora_ai\build\app\outputs\flutter-apk\app-debug.apk",
        "Size: 172,299,833 bytes (164.3 MiB).",
        "SHA-256: A006530443D7EC2BFFCE01DA0DACBDBA5DE2ACD10E3212FB662EF5F1133A1988.",
        "Manifest: com.example.amora_ai; version 1.0.0 (1); minSdk 24; targetSdk 36.",
    ],
)

# 32 Automation review.
add_heading(doc, "32. Automation Test Review", 1)
add_table(
    doc,
    ["Type", "Observed", "Assessment"],
    [
        ["Test files", "23 after QA evidence addition", "Broad presentation coverage"],
        ["Widget test sources", "123 declarations", "Primary automation layer"],
        ["Unit/source test sources", "18 declarations", "Controller/theme/calculation checks"],
        ["Runtime invocations", "143", "115 pass; 23 fail; 5 skipped"],
        ["Integration tests", "0", "Critical gap"],
        ["Golden evidence", "7 QA captures", "Generated from real widgets"],
        ["Coverage report", "Not configured", "No percentage claim"],
        ["Mocks/fixtures", "Local repositories and dummy data", "Deterministic but not end-to-end"],
    ],
    widths=[1.55, 1.9, 3.25],
)
add_heading(doc, "32.1 Test-maintenance findings", 2)
add_bullets(
    doc,
    [
        "four_tab_discover_test.dart conflicts with the current five-tab shell and newer Discover suite.",
        "The Events locked-state test targets deliberately removed premium-gate copy.",
        "Legacy onboarding journeys use ambiguous text locators after UI redesign.",
        "All-routes tests correctly expose the current /event-waitlist Material assertion.",
    ],
)
add_heading(doc, "32.2 Recommended automated additions", 2)
add_bullets(
    doc,
    [
        "integration_test/auth_onboarding_journey_test.dart with a fake service contract.",
        "integration_test/discover_to_chat_journey_test.dart.",
        "integration_test/events_rsvp_journey_test.dart with a fake repository.",
        "test/accessibility/text_scale_route_matrix_test.dart.",
        "test/security/demo_behavior_compile_flag_test.dart.",
        "CI coverage generation and thresholds after obsolete tests are removed.",
    ],
)

# 33 Defects.
add_heading(doc, "33. Defect Management", 1)
add_body(
    doc,
    "Eight confirmed defects are logged. Related automated failures are grouped "
    "where one root cause produces many assertions; risks and unexecuted areas "
    "remain separate.",
)
add_table(
    doc,
    ["ID", "Module", "Title", "Severity", "Priority", "Status"],
    [
        [
            item["id"],
            item["module"],
            item["title"],
            item["severity"],
            item["priority"],
            item["status"],
        ]
        for item in DATA["defects"]
    ],
    widths=[1.1, 1.25, 2.8, 0.75, 0.8, 0.75],
    font_size=7.8,
)
add_heading(doc, "33.1 Severity model", 2)
add_table(
    doc,
    ["Level", "Definition"],
    [
        ["P0 / Blocker", "Application cannot be used or built"],
        ["P1 / Critical", "Core feature unavailable, crash, data/security or release risk"],
        ["P2 / Major", "Important feature works incorrectly"],
        ["P3 / Minor", "Limited functional/usability impact"],
        ["P4 / Cosmetic", "Visual or copy issue"],
    ],
    widths=[1.45, 5.25],
)

# 34 Execution summary.
add_heading(doc, "34. Test Execution Summary", 1)
counts = META["catalogue_counts"]
executed = counts["Pass"] + counts["Fail"]
pass_rate = counts["Pass"] / executed
fail_rate = counts["Fail"] / executed
add_table(
    doc,
    ["Metric", "Value"],
    [
        ["Catalogue total", str(META["catalogue_total"])],
        ["Passed", str(counts["Pass"])],
        ["Failed", str(counts["Fail"])],
        ["Blocked", str(counts["Blocked"])],
        ["Not Run", str(counts["Not Run"])],
        ["Not Applicable", str(counts["Not Applicable"])],
        ["Executed denominator", str(executed)],
        ["Executed pass percentage", f"{pass_rate:.1%}"],
        ["Executed fail percentage", f"{fail_rate:.1%}"],
    ],
    widths=[3.4, 3.3],
)
module_rows = []
for module, values in DATA["module_summary"].items():
    module_rows.append(
        [
            module,
            str(sum(values.values())),
            str(values.get("Pass", 0)),
            str(values.get("Fail", 0)),
            str(values.get("Blocked", 0)),
            str(values.get("Not Run", 0)),
            str(values.get("Not Applicable", 0)),
        ]
    )
add_heading(doc, "34.1 Module-wise result", 2)
add_table(
    doc,
    ["Module", "Total", "Pass", "Fail", "Blocked", "Not Run", "N/A"],
    module_rows,
    widths=[2.0, 0.7, 0.7, 0.7, 0.8, 0.8, 0.7],
    font_size=7.5,
)
severity_counts = Counter(item["severity"] for item in DATA["defects"])
add_heading(doc, "34.2 Confirmed defects by severity", 2)
add_table(
    doc,
    ["Severity", "Count"],
    [[key, str(value)] for key, value in sorted(severity_counts.items())],
    widths=[3.3, 3.4],
)

# 35 Risks.
add_heading(doc, "35. Risk Assessment", 1)
add_table(
    doc,
    ["ID", "Risk", "Probability", "Impact", "Level", "Module", "Owner", "Status"],
    [
        [
            item["id"],
            item["risk"],
            item["probability"],
            item["impact"],
            item["level"],
            item["module"],
            item["owner"],
            item["status"],
        ]
        for item in DATA["risks"]
    ],
    widths=[0.75, 2.2, 0.75, 0.7, 0.7, 1.05, 0.85, 0.7],
    font_size=7.1,
)
add_heading(doc, "35.1 Highest release risks", 2)
add_bullets(
    doc,
    [
        "No production backend/auth/payment integration.",
        "Template Android identity and debug signing.",
        "Demo OTP exposure and memory-only session state.",
        "No integration-test/device/accessibility validation.",
        "Large debug binary and static product data.",
    ],
)

# 36 Release readiness.
add_heading(doc, "36. Release Readiness", 1)
readiness_rows = [
    ["Functional stability", "Not Ready", "Full suite fails; core flows are local"],
    ["UI consistency", "Conditional", "Strong sampled UI; contrast/overflow defects"],
    ["Build success", "Conditional", "Web/debug APK pass; release artifact not safe"],
    ["Automated coverage", "Conditional", "Broad widgets; no integration suite/coverage baseline"],
    ["Security/privacy", "Not Ready", "Demo OTP and no server authorization"],
    ["Backend/data", "Not Ready", "No production integration declared"],
    ["Payment/entitlement", "Not Ready", "Test-only disabled simulation"],
    ["Notifications", "Not Ready", "Local presentation only"],
    ["Accessibility", "Not Ready", "Contrast/scale defects; no screen reader session"],
    ["Performance", "Not Ready", "No profiling evidence; large debug APK"],
    ["Documentation", "Ready for current assessment", "QA report/evidence/workbook produced"],
]
add_table(doc, ["Dimension", "Status", "Basis"], readiness_rows, widths=[1.7, 1.15, 3.85])
add_callout(
    doc,
    "Decision",
    "NOT READY for production or store submission. A development/demo build is "
    "available for controlled frontend review only.",
    fill=TERTIARY,
)
add_heading(doc, "36.1 Minimum release conditions", 2)
add_bullets(
    doc,
    [
        "Close both P1 defects and the major Events/UI runtime issues.",
        "Return the full automated suite to green after removing obsolete contracts.",
        "Integrate test-environment auth, backend data, chat, RSVP, notifications and sandbox payment.",
        "Configure approved Android identity and secure release signing.",
        "Add integration tests, physical-device compatibility, screen reader/text-scale and profiling evidence.",
    ],
)

# 37 Recommendations.
add_heading(doc, "37. Recommendations", 1)
recommendations = [
    ["Critical", "Immediate", "Isolate/remove demo OTP from production builds; integrate server verification."],
    ["Critical", "Before production", "Replace template application ID and debug release signing."],
    ["Critical", "Before production", "Integrate authenticated backend, authorization and persisted session handling."],
    ["High", "Immediate", "Fix Events scaled-text overflow and /event-waitlist Material hierarchy."],
    ["High", "Immediate", "Resolve secondary contrast and semantic button color mapping."],
    ["High", "Automation", "Remove obsolete four-tab/locked-Events tests; repair onboarding locators."],
    ["High", "Automation", "Add integration_test journeys and coverage thresholds."],
    ["High", "Accessibility", "Run 200% scaling plus TalkBack/NVDA/VoiceOver audits."],
    ["High", "Performance", "Profile representative Android hardware and optimize images/binary size."],
    ["Medium", "Before production", "Replace static event dates/fixtures with time-zone-aware service data."],
    ["Medium", "Documentation", "Publish API contracts, environment setup, test accounts and release runbook."],
    ["Low", "Future", "Install Windows desktop workload only if Windows becomes a supported target."],
]
add_table(doc, ["Priority", "Horizon", "Recommendation"], recommendations, widths=[1.0, 1.35, 4.35], font_size=8.3)

# 38 Conclusion.
add_heading(doc, "38. Final Conclusion", 1)
add_body(
    doc,
    "Amora demonstrates a substantial, visually cohesive Flutter frontend with "
    "good current-widget coverage and successful web/debug Android builds. "
    "The strongest areas are Discover interaction tests, Events asset/detail "
    "presentation, identity-focused Profile UI, FAQ email-only support and "
    "shared design primitives.",
)
add_body(
    doc,
    "The assessment also establishes that this state is a prototype rather "
    "than a production-complete dating platform. Service-backed authentication, "
    "data, real-time chat, RSVP, notifications and payment are not integrated; "
    "the complete suite is red; release signing/identity is unsafe; and "
    "accessibility/performance/device evidence is incomplete.",
)
add_callout(
    doc,
    "Recommended next action",
    "Treat the generated debug APK as an internal frontend-review artifact. "
    "Prioritize P1 fixes, restore a green suite, integrate a controlled backend "
    "test environment, then repeat device, accessibility, performance and release validation.",
)

# 39 Appendices.
add_heading(doc, "39. Appendices", 1)
add_body(
    doc,
    "The appendices provide the traceable records behind the assessment. "
    "The companion XLSX is the editable source for filtering and formula-backed summaries.",
)

# Appendix A in landscape.
landscape = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(landscape, landscape=True)
add_heading(doc, "Appendix A — Complete Test Case Catalogue", 2)
add_body(
    doc,
    "Each catalogue record contains the complete required field set. Multiple "
    "automated assertions may map to a shared confirmed defect when they have one root cause.",
)
table = doc.add_table(rows=1, cols=4)
table.autofit = False
table.style = "Table Grid"
headers = [
    "Control",
    "Scenario / objective / setup",
    "Execution",
    "Expected / actual / trace",
]
widths = [1.15, 2.85, 2.75, 3.45]
for index, header in enumerate(headers):
    set_cell_text(table.rows[0].cells[index], header, bold=True, color=SURFACE, size=7)
    shade_cell(table.rows[0].cells[index], PRIMARY)
    table.rows[0].cells[index].width = Inches(widths[index])
set_repeat_header(table.rows[0])
for row_index, item in enumerate(DATA["cases"]):
    row = table.add_row()
    prevent_row_split(row)
    severity = item["severity"] or "—"
    values = [
        (
            f"Test Case ID: {item['id']}\n"
            f"Module: {item['module']}\n"
            f"Status: {item['status']}\n"
            f"Priority: {item['priority']}\n"
            f"Severity if failed: {severity}"
        ),
        (
            f"Test Scenario: {item['scenario']}\n"
            f"Test Objective: {item['objective']}\n"
            f"Preconditions: {item['preconditions']}\n"
            f"Test Data: {item['test_data']}"
        ),
        f"Execution Steps:\n{item['steps']}",
        (
            f"Expected Result: {item['expected']}\n"
            f"Actual Result: {item['actual']}\n"
            f"Environment: {item['environment']}\n"
            f"Evidence Reference: {item['evidence']}\n"
            f"Remarks: {item['remarks'] or 'None'}"
        ),
    ]
    for col_index, value in enumerate(values):
        cell = row.cells[col_index]
        set_cell_text(cell, value, size=6.25)
        cell.width = Inches(widths[col_index])
        if row_index % 2 == 0:
            shade_cell(cell, BACKGROUND)
set_table_borders(table)

# Return to portrait.
portrait = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(portrait)
add_heading(doc, "Appendix B — Requirement Traceability Matrix", 2)
add_table(
    doc,
    ["ID", "Module", "Requirement", "Cases", "Implementation", "Test status"],
    [
        [
            item["id"],
            item["module"],
            item["requirement"],
            item["test_cases"],
            item["implementation"],
            item["test_status"],
        ]
        for item in DATA["requirements"]
    ],
    widths=[1.0, 1.15, 1.9, 1.4, 1.15, 1.1],
    font_size=6.9,
)

add_heading(doc, "Appendix C — Defect Log", 2)
for item in DATA["defects"]:
    add_heading(doc, f"{item['id']} — {item['title']}", 3)
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ["Module", item["module"]],
            ["Description", item["description"]],
            ["Steps to Reproduce", item["steps"]],
            ["Expected Result", item["expected"]],
            ["Actual Result", item["actual"]],
            ["Severity / Priority / Status", f"{item['severity']} / {item['priority']} / {item['status']}"],
            ["Evidence", item["evidence"]],
            ["Recommended Fix", item["fix"]],
        ],
        widths=[1.7, 5.0],
        font_size=7.9,
    )

add_heading(doc, "Appendix D — Build Logs Summary", 2)
add_table(doc, ["Command", "Result", "Duration", "Evidence"], build_rows, widths=[2.0, 1.0, 1.0, 2.7], font_size=7.9)
add_body(
    doc,
    "Detailed plain-text command/environment evidence is stored in "
    "docs/qa_evidence/EV-BUILD-001_build_and_environment.txt.",
)

add_heading(doc, "Appendix E — Test Environment", 2)
add_table(doc, ["Property", "Observed value"], environment_rows, widths=[1.85, 4.85], font_size=8.2)

add_heading(doc, "Appendix F — Screenshots and Evidence", 2)
add_body(
    doc,
    "Images below are deterministic Flutter widget renders at controlled "
    "dimensions. They contain local fictional data and no credentials, private "
    "messages, real OTPs or payment information.",
)
figures = [
    ("EV-UI-001_auth_mobile.png", "Figure F1: Authentication Entry", 3.2),
    ("EV-UI-002_discover_mobile.png", "Figure F2: Discover", 3.2),
    ("EV-UI-003_events_mobile.png", "Figure F3: Events", 3.2),
    ("EV-UI-004_event_detail_mobile.png", "Figure F4: Event Details", 3.2),
    ("EV-UI-005_profile_mobile.png", "Figure F5: Profile Identity", 3.2),
    ("EV-UI-006_faq_support_mobile.png", "Figure F6: FAQ & Email Support", 3.2),
    ("EV-UI-007_events_desktop.png", "Figure F7: Events Desktop", 6.55),
]
for image, caption, width in figures:
    add_figure(doc, image, caption, width)
add_table(
    doc,
    ["Evidence ID", "Artifact", "Purpose"],
    [
        ["EV-BUILD-001", "Build and environment TXT", "Command outcomes, SDKs, APK metadata"],
        ["EV-TEST-001", "Test run summary TXT", "Automated counts and failure families"],
        ["EV-SOURCE-001", "Project inventory TXT", "Architecture/dependency facts"],
        ["EV-LIMIT-001", "Limitations TXT", "Explicit non-claims and blockers"],
    ],
    widths=[1.25, 2.45, 3.0],
    font_size=8,
)

add_heading(doc, "Appendix G — Route Inventory", 2)
add_body(
    doc,
    f"The scan found {META['route_constants_discovered']} route-name constants/"
    f"aliases. {len(DATA['routes'])} literal route values are enumerated below; "
    "alias constants referencing another class are counted in the source total "
    "but do not duplicate a literal row.",
)
add_table(
    doc,
    ["Route", "Source"],
    [[item["route"], item["source"]] for item in DATA["routes"]],
    widths=[2.2, 4.5],
    font_size=7.2,
)

add_heading(doc, "Appendix H — Screen Inventory", 2)
screen_rows = [[str(index), path] for index, path in enumerate(DATA["screens"], 1)]
add_table(doc, ["#", "Screen source"], screen_rows, widths=[0.6, 6.1], font_size=7.3)

add_heading(doc, "Appendix I — Dependency Review", 2)
add_table(
    doc,
    ["Dependency / class", "Observed purpose", "QA note"],
    [
        ["flutter", "UI/runtime", "Core SDK"],
        ["cupertino_icons", "Bundled icons", "Declared"],
        ["url_launcher", "Email support launcher", "Requires native-client verification"],
        ["flutter_test", "Unit/widget testing", "Development only"],
        ["flutter_lints", "Static lint rules", "Development only"],
        ["No Firebase/HTTP/Dio/socket/database/payment SDK", "Not present", "Production integration blocker"],
    ],
    widths=[2.3, 1.8, 2.6],
    font_size=8,
)

add_heading(doc, "Appendix J — Recommended Automated Tests", 2)
recommended_tests = [
    ["integration_test/auth_onboarding_journey_test.dart", "Server-contract auth, OTP and onboarding resume"],
    ["integration_test/discover_to_chat_journey_test.dart", "Swipe/match/message golden path"],
    ["integration_test/events_rsvp_journey_test.dart", "Repository-backed join/leave/waitlist state"],
    ["integration_test/support_email_launch_test.dart", "Android/iOS mailto launch and fallback"],
    ["test/accessibility/text_scale_route_matrix_test.dart", "All primary routes at 1.3x/2.0x"],
    ["test/accessibility/semantics_snapshot_test.dart", "Labels, selected/expanded state and focus order"],
    ["test/security/demo_behavior_compile_flag_test.dart", "Prevent demo OTP/payment state in production"],
    ["test/build/release_configuration_test.dart", "Reject template ID/debug release signing"],
    ["test/events/time_filtering_test.dart", "Clock-injected event freshness/time zones"],
    ["test/navigation/route_contract_test.dart", "Canonical literal route uniqueness and shell behavior"],
    ["test/regression/current_discover_contract_test.dart", "Single maintained Discover contract"],
    ["CI coverage job", "Line/branch baseline, test sharding and artifact retention"],
]
add_table(doc, ["Recommended artifact", "Purpose"], recommended_tests, widths=[3.2, 3.5], font_size=8)

apply_headers_and_footers(doc)
doc.save(OUTPUT)

summary_text = f"""AMORA QA TESTING REPORT — SUMMARY
Generated: 2026-07-27

Report path: {OUTPUT}
Workbook path: {DOCS / 'Amora_QA_Testing_Report_Test_Data.xlsx'}
Evidence folder: {EVIDENCE}

Modules reviewed: {META['modules_reviewed']}
Feature directories discovered: {META['feature_directories']}
Screens discovered: {META['screens_discovered']}
Route constants/aliases discovered: {META['route_constants_discovered']}

Documented test cases: {META['catalogue_total']}
Passed: {META['catalogue_counts']['Pass']}
Failed: {META['catalogue_counts']['Fail']}
Blocked: {META['catalogue_counts']['Blocked']}
Not Run: {META['catalogue_counts']['Not Run']}
Not Applicable: {META['catalogue_counts']['Not Applicable']}

Automated suite: {META['automated_total']} invocations
Automated passed: {META['automated_passed']}
Automated failed: {META['automated_failed']}
Automated skipped: {META['automated_skipped']}

Confirmed defect count: {META['confirmed_defects']}
flutter analyze: PASS — no issues found
flutter test: FAIL — 23 failures
Web build: PASS — {META['web_bytes']} bytes
Debug APK: PASS — {META['apk_path']}
Debug APK size: {META['apk_bytes']} bytes
Release APK: NOT ATTEMPTED — template application ID and debug signing

Release readiness: {META['release_decision']}

Important limitations:
- No production backend, API, database, Firebase, socket, payment or push SDK is declared.
- Authentication/session/profile/chat/event state is local or fixture-driven.
- No Android device/emulator, iOS host, integration_test suite, screen-reader session,
  profiling run, stakeholder UAT or production signing was available.
- Chrome debug launch connected, but an interactive manual browser pass was not completed.
"""
SUMMARY.write_text(summary_text, encoding="utf-8")

# Structural validation.
check = Document(OUTPUT)
all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
for section_title in section_titles:
    expected_heading = section_title.replace(" ", ". ", 1) if False else section_title
    number, title = section_title.split(" ", 1)
    if f"{number}. {title}" not in all_text:
        raise RuntimeError(f"Missing major heading: {number}. {title}")
if "[[TOC]]" in all_text or "TODO" in all_text:
    raise RuntimeError("Unresolved placeholder found")
if len(check.inline_shapes) < 7:
    raise RuntimeError("Evidence images are missing")
if len(check.tables) < 30:
    raise RuntimeError("Expected report tables are missing")
with zipfile.ZipFile(OUTPUT) as archive:
    document_xml = archive.read("word/document.xml").decode("utf-8")
    settings_xml = archive.read("word/settings.xml").decode("utf-8")
    footer_xml = "\n".join(
        archive.read(name).decode("utf-8")
        for name in archive.namelist()
        if name.startswith("word/footer")
    )
    if "TOC \\o" not in document_xml:
        raise RuntimeError("Automatic TOC field missing")
    if "updateFields" not in settings_xml:
        raise RuntimeError("Field update setting missing")
    if "NUMPAGES" not in footer_xml or "PAGE" not in footer_xml:
        raise RuntimeError("Page-number fields missing")
print(OUTPUT)
print(SUMMARY)
